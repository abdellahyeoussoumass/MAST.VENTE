import customtkinter as ctk
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from database.db import session
from database.models import Facture, Client
from datetime import date
from utils.theme import get_theme
import os
import tempfile

# ══════════════════════════════════════════════════════════
#   IMPORTS OPTIONNELS
# ══════════════════════════════════════════════════════════
try:
    from fpdf import FPDF
    FPDF_OK = True
except ImportError:
    FPDF_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer, HRFlowable)
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_OK = True
except ImportError:
    EXCEL_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_OK = True
except ImportError:
    WORD_OK = False

import csv
import json


# ══════════════════════════════════════════════════════════
#   CONSTANTES STATUT / PAIEMENT
# ══════════════════════════════════════════════════════════
STATUTS   = ["Payee", "Non payee", "En attente", "Annulee"]
PAIEMENTS = ["Especes", "Virement", "Cheque", "Carte bancaire"]

STATUT_ICONE = {
    "Payee":      "✅",
    "Non payee":  "❌",
    "En attente": "⏳",
    "Annulee":    "🚫",
}
STATUT_COULEUR = {
    "Payee":      "#43A047",
    "Non payee":  "#E53935",
    "En attente": "#FFA000",
    "Annulee":    "#9E9E9E",
}
PAIEMENT_ICONE = {
    "Especes":        "💵",
    "Virement":       "🏦",
    "Cheque":         "📄",
    "Carte bancaire": "💳",
}


# ══════════════════════════════════════════════════════════
#   STYLE TREEVIEW
# ══════════════════════════════════════════════════════════
def _style_treeview(is_dark):
    style = ttk.Style()
    style.theme_use("clam")

    bg_tree = "#12122A" if is_dark else "#FFFFFF"
    bg_head = "#E65100"
    bg_sel  = "#FF6D00" if is_dark else "#FFB74D"
    fg_text = "#E8E0FF" if is_dark else "#1A1A1A"
    bg_row1 = "#1A1A35" if is_dark else "#FFFFFF"
    bg_row2 = "#12122A" if is_dark else "#FFF8F0"

    style.configure("FA.Treeview",
                    background=bg_tree, foreground=fg_text,
                    rowheight=32, fieldbackground=bg_tree,
                    borderwidth=0, font=("Arial", 10))
    style.configure("FA.Treeview.Heading",
                    background=bg_head, foreground="white",
                    font=("Arial", 10, "bold"),
                    relief="flat", borderwidth=0)
    style.map("FA.Treeview",
              background=[("selected", bg_sel)],
              foreground=[("selected", "#FFFFFF")])
    style.map("FA.Treeview.Heading",
              background=[("active", "#BF360C")])
    return bg_row1, bg_row2


# ══════════════════════════════════════════════════════════
#   HELPERS UI
# ══════════════════════════════════════════════════════════
def _fenetre_modale(titre, largeur=480, hauteur=600):
    t   = get_theme()
    win = tk.Toplevel()
    win.title(titre)
    win.geometry(f"{largeur}x{hauteur}")
    win.configure(bg=t["bg"])
    win.resizable(False, False)
    win.grab_set()

    hdr = tk.Frame(win, bg="#E65100", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=titre, bg="#E65100", fg="white",
             font=("Arial", 14, "bold")).pack(pady=12)

    body = tk.Frame(win, bg=t["bg"])
    body.pack(fill="both", expand=True, padx=24, pady=16)
    return win, body, t


def _lbl(body, text, t):
    tk.Label(body, text=text, bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(
             fill="x", pady=(10, 2))


def _entry(body, t, default=""):
    e = tk.Entry(body, font=("Arial", 11),
                 bg=t["card"], fg=t["text"],
                 insertbackground=t["text"],
                 relief="flat", bd=6)
    e.pack(fill="x", ipady=4)
    if default != "":
        e.insert(0, str(default))
    return e


def _combo(body, t, values, default=""):
    var = tk.StringVar(value=default)
    cb  = ttk.Combobox(body, textvariable=var, values=values,
                       font=("Arial", 11), state="readonly")
    cb.pack(fill="x", ipady=4)
    cb.var = var
    return cb


def _btn_ok(parent, text, cmd, color="#E65100", hover="#BF360C"):
    return tk.Button(
        parent, text=text, command=cmd,
        bg=color, fg="white",
        font=("Arial", 12, "bold"),
        relief="flat", cursor="hand2",
        padx=16, pady=9,
        activebackground=hover,
        activeforeground="white",
    )


def _sep(body, t):
    tk.Frame(body,
             bg="#2A2A4A" if t["bg"] in ("#1A1A2E", "#0F0F1A") else "#E8D8C8",
             height=1).pack(fill="x", pady=(10, 0))


# ══════════════════════════════════════════════════════════
#   TRI COLONNES
# ══════════════════════════════════════════════════════════
_tri_etat = {}

def _trier(table, col):
    rows = [(table.set(r, col), r) for r in table.get_children("")]
    rev  = _tri_etat.get(col, False)
    try:
        rows.sort(key=lambda x: float(
            x[0].replace(",", "").replace("%", "")), reverse=rev)
    except ValueError:
        rows.sort(key=lambda x: x[0].lower(), reverse=rev)
    for i, (_, r) in enumerate(rows):
        table.move(r, "", i)
    _tri_etat[col] = not rev


# ══════════════════════════════════════════════════════════
#   MENU DEROULANT EXPORTER
# ══════════════════════════════════════════════════════════
def _show_dropdown_export(btn_widget, table, t):
    is_dark  = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
    bg_menu  = "#1A1A35" if is_dark else "#FFFFFF"
    bg_hover = "#252545" if is_dark else "#FFF0E5"
    border_c = "#2A2A4A" if is_dark else "#E8D8C8"

    items = [
        ("📊  Excel (.xlsx)",  "#1565C0", "excel"),
        ("📝  Word (.docx)",   "#1976D2", "word"),
        ("📕  PDF (.pdf)",     "#C62828", "pdf"),
        ("📄  CSV (.csv)",     "#00838F", "csv"),
        ("📋  JSON (.json)",   "#6A1B9A", "json"),
    ]

    menu = tk.Toplevel()
    menu.overrideredirect(True)
    menu.configure(bg=border_c)
    menu.attributes("-topmost", True)

    btn_widget.update_idletasks()
    x = btn_widget.winfo_rootx()
    y = btn_widget.winfo_rooty() + btn_widget.winfo_height() + 2
    menu.geometry(f"210x{len(items) * 40 + 10}+{x}+{y}")

    inner = tk.Frame(menu, bg=bg_menu, bd=0)
    inner.pack(fill="both", expand=True, padx=1, pady=1)

    tk.Label(inner, text="Choisir le format :",
             bg=bg_menu, fg="#FF8F00",
             font=("Arial", 8, "bold"),
             anchor="w").pack(fill="x", padx=10, pady=(6, 2))

    def _fermer():
        try:
            menu.destroy()
        except Exception:
            pass

    for label, couleur, fmt in items:
        f = tk.Frame(inner, bg=bg_menu, cursor="hand2")
        f.pack(fill="x", padx=4, pady=1)
        lbl = tk.Label(f, text=label, bg=bg_menu, fg=couleur,
                       font=("Arial", 10, "bold"),
                       anchor="w", padx=12, pady=8)
        lbl.pack(fill="x")

        def _enter(e, fr=f, lb=lbl):
            fr.configure(bg=bg_hover); lb.configure(bg=bg_hover)
        def _leave(e, fr=f, lb=lbl):
            fr.configure(bg=bg_menu);  lb.configure(bg=bg_menu)
        def _click(e, fn=fmt):
            _fermer(); exporter_factures(table, fn)

        for w in (f, lbl):
            w.bind("<Enter>",    _enter)
            w.bind("<Leave>",    _leave)
            w.bind("<Button-1>", _click)

    menu.bind("<FocusOut>", lambda e: _fermer())
    menu.focus_set()


# ══════════════════════════════════════════════════════════
#   EXPORTER FACTURES
# ══════════════════════════════════════════════════════════
EN_TETES_FA = ["ID", "N Facture", "Client", "Prix HT",
               "TVA (%)", "Prix TTC", "Reduction (%)", "Statut", "Paiement"]
CHAMPS_FA   = ["id", "numero_facture", "client", "prix_ht",
               "tva", "prix_ttc", "reduction", "statut", "mode_paiement"]


def _get_factures_data():
    rows = []
    factures = session.query(Facture).order_by(Facture.id.desc()).all()
    for f in factures:
        client = session.query(Client).filter_by(id=f.client_id).first()
        rows.append({
            "id":             f.id,
            "numero_facture": f.numero_facture or "-",
            "client":         client.nom if client else "N/A",
            "prix_ht":        float(f.prix_ht or 0),
            "tva":            float(f.tva or 0),
            "prix_ttc":       float(f.prix_ttc or 0),
            "reduction":      float(f.reduction or 0),
            "statut":         f.statut or "-",
            "mode_paiement":  f.mode_paiement or "-",
        })
    return rows


def exporter_factures(table, format_):
    rows = _get_factures_data()
    if not rows:
        messagebox.showwarning("Vide", "Aucune facture a exporter.")
        return

    ext_map = {
        "excel": ("Fichier Excel", "*.xlsx", ".xlsx"),
        "word":  ("Document Word",  "*.docx", ".docx"),
        "pdf":   ("Fichier PDF",    "*.pdf",  ".pdf"),
        "csv":   ("Fichier CSV",    "*.csv",  ".csv"),
        "json":  ("Fichier JSON",   "*.json", ".json"),
    }
    label, pattern, ext = ext_map[format_]
    chemin = filedialog.asksaveasfilename(
        title=f"Exporter en {label}",
        defaultextension=ext,
        filetypes=[(label, pattern), ("Tous", "*.*")],
        initialfile=f"factures_{date.today()}{ext}",
    )
    if not chemin:
        return

    try:
        if format_ == "excel":   _fac_export_excel(chemin, rows)
        elif format_ == "word":  _fac_export_word(chemin, rows)
        elif format_ == "pdf":   _fac_export_pdf_liste(chemin, rows)
        elif format_ == "csv":   _fac_export_csv(chemin, rows)
        elif format_ == "json":  _fac_export_json(chemin, rows)

        if messagebox.askyesno("Export reussi",
                               f"Fichier cree :\n{chemin}\n\nOuvrir maintenant ?"):
            os.startfile(chemin)
    except Exception as ex:
        messagebox.showerror("Erreur export", f"Erreur :\n{ex}")


def _fac_export_excel(chemin, rows):
    if not EXCEL_OK:
        messagebox.showerror("Manquant", "pip install openpyxl"); return
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Factures"
    ws.merge_cells("A1:I1")
    ws["A1"] = f"Factures — Exporte le {date.today()}"
    ws["A1"].font      = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill      = PatternFill("solid", fgColor="E65100")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36
    for ci, h in enumerate(EN_TETES_FA, 1):
        cell = ws.cell(row=2, column=ci, value=h)
        cell.font      = Font(bold=True, color="FFFFFF", size=10)
        cell.fill      = PatternFill("solid", fgColor="BF360C")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 24
    bd   = Side(style="thin", color="E8D8C8")
    bord = Border(left=bd, right=bd, top=bd, bottom=bd)
    fp   = PatternFill("solid", fgColor="FFF8F0")
    fi   = PatternFill("solid", fgColor="FFFFFF")
    for ri, row in enumerate(rows, 3):
        fill = fp if ri % 2 == 0 else fi
        for ci, champ in enumerate(CHAMPS_FA, 1):
            cell = ws.cell(row=ri, column=ci, value=row[champ])
            cell.fill      = fill
            cell.border    = bord
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.font      = Font(size=9)
            if champ == "statut":
                col_map = {"Payee": "2E7D32", "Non payee": "C62828",
                           "En attente": "E65100", "Annulee": "757575"}
                cell.font = Font(size=9, bold=True,
                                 color=col_map.get(str(row[champ]), "212121"))
    larg = [6, 14, 20, 10, 8, 12, 10, 12, 14]
    for i, w in enumerate(larg, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w
    total = sum(r["prix_ttc"] for r in rows)
    rt = len(rows) + 3
    ws.cell(row=rt, column=6, value=f"TOTAL : {total:,.2f} MAD")
    ws.cell(row=rt, column=6).font = Font(bold=True, color="E65100", size=10)
    ws.cell(row=rt, column=6).fill = PatternFill("solid", fgColor="FFF0E5")
    wb.save(chemin)


def _fac_export_word(chemin, rows):
    if not WORD_OK:
        messagebox.showerror("Manquant", "pip install python-docx"); return
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2);  section.right_margin  = Cm(2)
    titre = doc.add_heading("Factures VentePro", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.runs[0].font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    titre.runs[0].font.size      = Pt(18)
    st = doc.add_paragraph(f"Exporte le {date.today()}  |  {len(rows)} facture(s)")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x33)
    st.runs[0].font.size      = Pt(10)
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=len(EN_TETES_FA))
    tbl.style = "Table Grid"
    for i, h in enumerate(EN_TETES_FA):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold = True; run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E65100"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear"); tc_pr.append(shd)
    for row in rows:
        tr = tbl.add_row()
        for i, champ in enumerate(CHAMPS_FA):
            cell = tr.cells[i]
            cell.text = str(row[champ])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            if len(cell.paragraphs[0].runs) > 0:
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)
    doc.save(chemin)


def _fac_export_pdf_liste(chemin, rows):
    if not PDF_OK:
        messagebox.showerror("Manquant", "pip install reportlab"); return
    doc    = SimpleDocTemplate(chemin, pagesize=A4,
                               leftMargin=1.5*cm, rightMargin=1.5*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []
    story.append(Paragraph("Liste des Factures — VentePro",
        ParagraphStyle("T", parent=styles["Title"],
                       fontSize=18, textColor=colors.HexColor("#E65100"),
                       spaceAfter=4)))
    story.append(Paragraph(
        f"Exporte le {date.today()}  —  {len(rows)} facture(s)",
        ParagraphStyle("S", parent=styles["Normal"],
                       fontSize=9, textColor=colors.HexColor("#996633"),
                       spaceAfter=14)))
    data = [EN_TETES_FA]
    for row in rows:
        data.append([str(row[c]) for c in CHAMPS_FA])
    col_w = [1*cm, 2.5*cm, 3.5*cm, 2*cm, 1.5*cm, 2.2*cm, 2*cm, 2.2*cm, 2.5*cm]
    tbl   = Table(data, colWidths=col_w, repeatRows=1)
    tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#E65100")),
        ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 8),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 7),
        ("TOPPADDING",    (0, 0), (-1, 0), 7),
        ("FONTSIZE",      (0, 1), (-1, -1), 7.5),
        ("ROWBACKGROUNDS",(0, 1), (-1, -1),
         [colors.HexColor("#FFF8F0"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E8D8C8")),
        ("TOPPADDING",    (0, 1), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 4),
    ]))
    story.append(tbl)
    doc.build(story)


def _fac_export_csv(chemin, rows):
    with open(chemin, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=CHAMPS_FA)
        writer.writeheader()
        writer.writerows(rows)


def _fac_export_json(chemin, rows):
    with open(chemin, "w", encoding="utf-8") as f:
        json.dump({
            "export_date":   str(date.today()),
            "total_records": len(rows),
            "factures":      rows,
        }, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════
#   AFFICHER FACTURES
# ══════════════════════════════════════════════════════════
def afficher_factures(parent):
    for widget in parent.winfo_children():
        widget.destroy()

    t       = get_theme()
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")

    # En-tete
    header = ctk.CTkFrame(parent, fg_color=t["card"], corner_radius=0)
    header.pack(fill="x")
    ctk.CTkLabel(header, text="🧾  Gestion des Factures",
                 font=("Arial", 22, "bold"),
                 text_color="#E65100").pack(side="left", padx=24, pady=14)

    # Barre outils
    toolbar = ctk.CTkFrame(parent, fg_color="transparent")
    toolbar.pack(fill="x", padx=20, pady=(10, 4))

    btns = [
        ("➕  Ajouter",   "#E65100", "#BF360C", lambda: ajouter_facture(table)),
        ("✏️  Modifier",  "#FF8F00", "#E65100", lambda: modifier_facture(table)),
        ("👁️  Consulter", "#2E7D32", "#1B5E20", lambda: consulter_facture(table)),
        ("📧  Envoyer",   "#1565C0", "#0D47A1", lambda: envoyer_facture(table)),
        ("🗑️  Supprimer", "#C62828", "#8B0000", lambda: supprimer_facture(table)),
    ]
    for txt, fg, hover, cmd in btns:
        ctk.CTkButton(toolbar, text=txt, width=118, height=36,
                      fg_color=fg, hover_color=hover,
                      font=("Arial", 11, "bold"),
                      corner_radius=8,
                      command=cmd).pack(side="left", padx=3)

    # ── Bouton EXPORTER avec menu deroulant ───────────────────
    btn_export_widget = ctk.CTkButton(
        toolbar, text="📤  Exporter ▾", width=138, height=36,
        fg_color="#6A1B9A", hover_color="#4A148C",
        font=("Arial", 11, "bold"), corner_radius=8,
        command=lambda: _show_dropdown_export(btn_export_widget, table, t)
    )
    btn_export_widget.pack(side="left", padx=3)

    # ── Barre de recherche générale ──────────────────────────
    search_frame = ctk.CTkFrame(parent, fg_color="transparent")
    search_frame.pack(fill="x", padx=20, pady=(4, 2))

    count_var = tk.StringVar()
    ctk.CTkLabel(search_frame, textvariable=count_var,
                 font=("Arial", 11),
                 text_color="#FF8F00").pack(side="right", padx=10)

    # ── Variables d etat ──────────────────────────────────
    filtre_actif = {"valeur": None, "type": None}
    dropdown_ref = [None]

    FILTRES = [
        ("Tous",           "#555555", "filtre"),
        ("Payee",          "#43A047", "filtre"),
        ("Non payee",      "#E53935", "filtre"),
        ("En attente",     "#FFA000", "filtre"),
        ("Annulee",        "#9E9E9E", "filtre"),
        ("Especes",        "#1976D2", "paiement"),
        ("Virement",       "#1976D2", "paiement"),
        ("Cheque",         "#1976D2", "paiement"),
        ("Carte bancaire", "#1976D2", "paiement"),
    ]

    ICONES_F = {
        "Tous": "📋", "Payee": "✅", "Non payee": "❌",
        "En attente": "⏳", "Annulee": "🚫",
        "Especes": "💵", "Virement": "🏦",
        "Cheque": "📄", "Carte bancaire": "💳",
    }

    # ── Conteneur barre ───────────────────────────────────
    barre = tk.Frame(search_frame, bg=t["bg"])
    barre.pack(side="right", padx=5)

    # Bouton filtre
    var_btn_txt = tk.StringVar(value="⚙  Filtre ▾")
    btn_filtre  = tk.Button(
        barre, textvariable=var_btn_txt,
        bg="#FF8F00", fg="white",
        font=("Arial", 10, "bold"),
        relief="flat", cursor="hand2",
        padx=10, pady=6,
        activebackground="#E65100",
        activeforeground="white",
    )
    btn_filtre.pack(side="left", padx=(0, 4))

    # Cadre du champ texte
    frame_input = tk.Frame(barre,
                           bg=t["card"],
                           highlightbackground="#E65100",
                           highlightthickness=1)
    frame_input.pack(side="left")

    recherche_var = tk.StringVar()
    PLACEHOLDER   = "🔍  Rechercher dans toutes les colonnes..."

    e_search = tk.Entry(
        frame_input,
        textvariable=recherche_var,
        font=("Arial", 11),
        bg=t["card"], fg="#888888",
        insertbackground=t["text"],
        relief="flat", bd=6,
        width=36,
    )
    e_search.insert(0, PLACEHOLDER)
    e_search.pack(side="left", ipady=5)

    btn_clear = tk.Button(
        frame_input, text=" ✕ ",
        bg=t["card"], fg="#C62828",
        font=("Arial", 10, "bold"),
        relief="flat", cursor="hand2",
        activebackground=t["card"],
        activeforeground="#8B0000", bd=0,
    )
    # btn_clear masque par defaut

    # ── Placeholder ───────────────────────────────────────
    def _show_ph():
        if not recherche_var.get():
            e_search.configure(fg="#888888")
            e_search.delete(0, "end")
            e_search.insert(0, PLACEHOLDER)

    def _hide_ph(e=None):
        if e_search.get() == PLACEHOLDER:
            e_search.delete(0, "end")
            e_search.configure(fg=t["text"])

    e_search.bind("<FocusIn>",  _hide_ph)
    e_search.bind("<FocusOut>", lambda e: _show_ph())

    # ── Effacer tout ──────────────────────────────────────
    def _effacer_tout():
        filtre_actif["valeur"] = None
        filtre_actif["type"]   = None
        var_btn_txt.set("⚙  Filtre ▾")
        btn_filtre.configure(bg="#FF8F00")
        recherche_var.set("")
        _show_ph()
        btn_clear.pack_forget()
        # table pas encore défini ici → on passe par une lambda
        _refresh_table()

    btn_clear.configure(command=_effacer_tout)

    # ── Recherche combinée (texte + filtre) ───────────────
    # NOTE : _refresh_table() et _appliquer_recherche() utilisent
    # la variable "table" qui est définie APRES ce bloc.
    # On stocke une référence mutable pour éviter le problème.
    table_ref = [None]

    def _refresh_table():
        if table_ref[0] is None:
            return
        charger_factures(table_ref[0], count_var)

    def _appliquer_recherche():
        if table_ref[0] is None:
            return
        tbl = table_ref[0]

        texte     = recherche_var.get().strip()
        if texte == PLACEHOLDER:
            texte = ""
        texte_low = texte.lower()

        statut_v = filtre_actif["valeur"]
        type_v   = filtre_actif["type"]

        # Charger tout
        factures_all = session.query(Facture).order_by(Facture.id.desc()).all()

        # Filtre statut/paiement
        if statut_v and statut_v != "Tous":
            if type_v == "filtre":
                factures_all = [f for f in factures_all
                                if (f.statut or "") == statut_v]
            else:
                factures_all = [f for f in factures_all
                                if (f.mode_paiement or "") == statut_v]

        # Recherche texte dans TOUTES les colonnes
        if texte_low:
            def _match(f):
                client = session.query(Client).filter_by(id=f.client_id).first()
                nom_c  = (client.nom or "") if client else ""
                return any(texte_low in str(v).lower() for v in [
                    f.numero_facture, nom_c,
                    f.prix_ht, f.tva, f.prix_ttc, f.reduction,
                    f.statut, f.mode_paiement,
                ])
            factures_all = [f for f in factures_all if _match(f)]

        # Afficher
        for row in tbl.get_children():
            tbl.delete(row)

        total_ttc = total_paye = 0.0
        for i, f in enumerate(factures_all):
            client     = session.query(Client).filter_by(id=f.client_id).first()
            statut_raw = f.statut or ""
            icone_s    = STATUT_ICONE.get(statut_raw, "")
            icone_p    = PAIEMENT_ICONE.get(f.mode_paiement or "", "")
            tags       = ["pair" if i % 2 == 0 else "impair"]
            if statut_raw in STATUT_COULEUR:
                tags.append(statut_raw)

            ttc = float(f.prix_ttc or 0)
            total_ttc += ttc
            if statut_raw == "Payee":
                total_paye += ttc

            tbl.insert("", "end", tags=tuple(tags), values=(
                f.id,
                f.numero_facture or "-",
                client.nom if client else "N/A",
                f"{float(f.prix_ht or 0):,.2f}",
                f"{float(f.tva or 0):.0f}%",
                f"{ttc:,.2f}",
                f"{float(f.reduction or 0):.0f}%",
                f"{icone_s} {statut_raw}".strip() if icone_s else statut_raw,
                f"{icone_p} {f.mode_paiement or '-'}".strip() if icone_p else (f.mode_paiement or "-"),
            ))

        nb = len(factures_all)
        suf = f"  |  Filtre: {statut_v}" if statut_v and statut_v != "Tous" else ""
        if texte:
            suf += f'  |  "{texte}"'
        count_var.set(
            f"{nb} resultat(s){suf}  |  "
            f"TTC: {total_ttc:,.2f}  |  Paye: {total_paye:,.2f} MAD"
        )

        # Afficher/masquer ✕
        if texte or (statut_v and statut_v != "Tous"):
            btn_clear.pack(side="right")
        else:
            btn_clear.pack_forget()

    # Frappe en temps réel
    def _on_frappe(*_):
        texte = recherche_var.get()
        if texte and texte != PLACEHOLDER:
            _appliquer_recherche()
        elif not texte:
            btn_clear.pack_forget()
            _refresh_table()

    recherche_var.trace("w", _on_frappe)

    # ── Dropdown filtre ───────────────────────────────────
    def _fermer_dropdown():
        try:
            if dropdown_ref[0] and dropdown_ref[0].winfo_exists():
                dropdown_ref[0].destroy()
        except Exception:
            pass
        dropdown_ref[0] = None

    def _activer_filtre(valeur, type_filtre):
        _fermer_dropdown()
        filtre_actif["valeur"] = valeur
        filtre_actif["type"]   = type_filtre
        if valeur == "Tous":
            var_btn_txt.set("⚙  Filtre ▾")
            btn_filtre.configure(bg="#FF8F00")
        else:
            icone = ICONES_F.get(valeur, "🔍")
            var_btn_txt.set(f"{icone} {valeur} ▾")
            col = {"Payee": "#43A047", "Non payee": "#E53935",
                   "En attente": "#FFA000", "Annulee": "#9E9E9E"
                   }.get(valeur, "#1976D2")
            btn_filtre.configure(bg=col)
        _appliquer_recherche()

    def _ouvrir_dropdown(e=None):
        _fermer_dropdown()

        is_d  = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
        bg_m  = "#1A1A35" if is_d else "#FFFFFF"
        bg_h  = "#252545" if is_d else "#FFF0E5"
        brd   = "#2A2A4A" if is_d else "#E8D8C8"
        muted = "#A0A0C0" if is_d else "#999999"

        dd = tk.Toplevel()
        dd.overrideredirect(True)
        dd.configure(bg=brd)
        dd.attributes("-topmost", True)
        dropdown_ref[0] = dd

        btn_filtre.update_idletasks()
        x = btn_filtre.winfo_rootx()
        y = btn_filtre.winfo_rooty() + btn_filtre.winfo_height() + 2
        dd.geometry(f"240x{len(FILTRES)*38+28}+{x}+{y}")

        inner = tk.Frame(dd, bg=bg_m)
        inner.pack(fill="both", expand=True, padx=1, pady=1)

        tk.Label(inner, text="  Filtrer par statut / paiement",
                 bg=bg_m, fg=muted,
                 font=("Arial", 8, "bold"),
                 anchor="w").pack(fill="x", padx=8, pady=(5, 2))

        for label, couleur, type_item in FILTRES:
            est_actif = filtre_actif["valeur"] == label
            bg_i = bg_h if est_actif else bg_m

            rf = tk.Frame(inner, bg=bg_i, cursor="hand2")
            rf.pack(fill="x", padx=4, pady=1)

            cv = tk.Canvas(rf, width=10, height=10,
                           bg=bg_i, highlightthickness=0)
            cv.create_oval(1, 1, 9, 9, fill=couleur, outline="")
            cv.pack(side="left", padx=(10, 4), pady=10)

            tk.Label(rf, text="✓" if est_actif else "  ",
                     bg=bg_i, fg=couleur,
                     font=("Arial", 10, "bold"), width=2).pack(side="left")

            lb = tk.Label(rf, text=label,
                          bg=bg_i, fg=couleur,
                          font=("Arial", 10, "bold"),
                          anchor="w", pady=8)
            lb.pack(side="left", fill="x", expand=True)

            def _mk(fr=rf, lb_=lb, cv_=cv, act=est_actif):
                def _in(e):
                    fr.configure(bg=bg_h); lb_.configure(bg=bg_h); cv_.configure(bg=bg_h)
                def _out(e):
                    c = bg_h if act else bg_m
                    fr.configure(bg=c); lb_.configure(bg=c); cv_.configure(bg=c)
                return _in, _out

            _in, _out = _mk()
            for w in (rf, lb, cv):
                w.bind("<Enter>", _in)
                w.bind("<Leave>", _out)
                w.bind("<Button-1>",
                       lambda e, v=label, tp=type_item: _activer_filtre(v, tp))

        dd.bind("<FocusOut>", lambda e: _fermer_dropdown())
        dd.focus_set()

    btn_filtre.configure(command=_ouvrir_dropdown)

    # Tableau
    frame_table = ctk.CTkFrame(parent, corner_radius=12)
    frame_table.pack(fill="both", expand=True, padx=20, pady=(4, 16))

    bg_row1, bg_row2 = _style_treeview(is_dark)

    colonnes = ("ID", "N Facture", "Client", "Prix HT",
                "TVA", "Prix TTC", "Reduction", "Statut", "Paiement")

    table = ttk.Treeview(frame_table, columns=colonnes,
                         show="headings", height=22,
                         style="FA.Treeview",
                         selectmode="browse")

    largeurs = {"ID": 45, "N Facture": 120, "Client": 160,
                "Prix HT": 90, "TVA": 60, "Prix TTC": 100,
                "Reduction": 80, "Statut": 110, "Paiement": 130}

    for col in colonnes:
        table.heading(col, text=col,
                      command=lambda c=col: _trier(table, c))
        table.column(col, width=largeurs.get(col, 100), anchor="center")

    # Tags par statut
    table.tag_configure("pair",      background=bg_row1)
    table.tag_configure("impair",    background=bg_row2)
    table.tag_configure("Payee",     foreground="#43A047")
    table.tag_configure("Non payee", foreground="#E53935")
    table.tag_configure("En attente",foreground="#FFA000")
    table.tag_configure("Annulee",   foreground="#9E9E9E")

    sb_v = ttk.Scrollbar(frame_table, orient="vertical",   command=table.yview)
    sb_h = ttk.Scrollbar(frame_table, orient="horizontal", command=table.xview)
    table.configure(yscroll=sb_v.set, xscroll=sb_h.set)
    sb_v.pack(side="right",  fill="y")
    sb_h.pack(side="bottom", fill="x")
    table.pack(fill="both", expand=True, padx=2, pady=2)

    table.bind("<Double-1>", lambda e: detail_facture(table))

    # Connecter la reference du tableau a la barre de recherche
    table_ref[0] = table
    charger_factures(table, count_var)
    return table


# ══════════════════════════════════════════════════════════
#   CHARGER FACTURES
# ══════════════════════════════════════════════════════════
def charger_factures(table, count_var=None):
    for row in table.get_children():
        table.delete(row)

    factures = session.query(Facture).order_by(Facture.id.desc()).all()
    total_ttc = 0.0
    total_paye = 0.0

    for i, f in enumerate(factures):
        client = session.query(Client).filter_by(id=f.client_id).first()
        statut_raw = f.statut or ""
        icone_s    = STATUT_ICONE.get(statut_raw, "")
        icone_p    = PAIEMENT_ICONE.get(f.mode_paiement or "", "")

        tags = ["pair" if i % 2 == 0 else "impair"]
        if statut_raw in STATUT_COULEUR:
            tags.append(statut_raw)

        ttc = float(f.prix_ttc or 0)
        total_ttc += ttc
        if statut_raw == "Payee":
            total_paye += ttc

        table.insert("", "end", tags=tuple(tags), values=(
            f.id,
            f.numero_facture or "-",
            client.nom if client else "N/A",
            f"{float(f.prix_ht or 0):,.2f}",
            f"{float(f.tva or 0):.0f}%",
            f"{ttc:,.2f}",
            f"{float(f.reduction or 0):.0f}%",
            f"{icone_s} {statut_raw}".strip() if icone_s else statut_raw,
            f"{icone_p} {f.mode_paiement or '-'}".strip() if icone_p else (f.mode_paiement or "-"),
        ))

    if count_var is not None:
        count_var.set(
            f"{len(factures)} facture(s)  |  "
            f"TTC : {total_ttc:,.2f}  |  "
            f"Paye : {total_paye:,.2f} MAD"
        )


# ══════════════════════════════════════════════════════════
#   FILTRER PAR STATUT
# ══════════════════════════════════════════════════════════
def _filtrer(table, statut, count_var=None):
    for row in table.get_children():
        table.delete(row)

    if statut == "Tous":
        charger_factures(table, count_var)
        return

    factures = session.query(Facture).filter_by(statut=statut).all()
    for i, f in enumerate(factures):
        client = session.query(Client).filter_by(id=f.client_id).first()
        statut_raw = f.statut or ""
        icone_s    = STATUT_ICONE.get(statut_raw, "")
        icone_p    = PAIEMENT_ICONE.get(f.mode_paiement or "", "")
        tags = ["pair" if i % 2 == 0 else "impair", statut_raw]

        table.insert("", "end", tags=tuple(tags), values=(
            f.id,
            f.numero_facture or "-",
            client.nom if client else "N/A",
            f"{float(f.prix_ht or 0):,.2f}",
            f"{float(f.tva or 0):.0f}%",
            f"{float(f.prix_ttc or 0):,.2f}",
            f"{float(f.reduction or 0):.0f}%",
            f"{icone_s} {statut_raw}".strip() if icone_s else statut_raw,
            f"{icone_p} {f.mode_paiement or '-'}".strip() if icone_p else (f.mode_paiement or "-"),
        ))

    if count_var is not None:
        count_var.set(f"{len(factures)} facture(s) — {statut}")


# ══════════════════════════════════════════════════════════
#   RECHERCHER
# ══════════════════════════════════════════════════════════
def rechercher(texte, table):
    for row in table.get_children():
        table.delete(row)

    texte_low = texte.lower()
    factures  = session.query(Facture).all()

    for i, f in enumerate(factures):
        client = session.query(Client).filter_by(id=f.client_id).first()
        nom_c  = client.nom if client else ""

        if (texte_low in nom_c.lower()
                or texte_low in (f.numero_facture or "").lower()
                or texte_low in (f.statut or "").lower()
                or texte_low in (f.mode_paiement or "").lower()):

            statut_raw = f.statut or ""
            icone_s    = STATUT_ICONE.get(statut_raw, "")
            icone_p    = PAIEMENT_ICONE.get(f.mode_paiement or "", "")
            tags = ["pair" if i % 2 == 0 else "impair"]
            if statut_raw in STATUT_COULEUR:
                tags.append(statut_raw)

            table.insert("", "end", tags=tuple(tags), values=(
                f.id,
                f.numero_facture or "-",
                nom_c or "N/A",
                f"{float(f.prix_ht or 0):,.2f}",
                f"{float(f.tva or 0):.0f}%",
                f"{float(f.prix_ttc or 0):,.2f}",
                f"{float(f.reduction or 0):.0f}%",
                f"{icone_s} {statut_raw}".strip() if icone_s else statut_raw,
                f"{icone_p} {f.mode_paiement or '-'}".strip() if icone_p else (f.mode_paiement or "-"),
            ))


# ══════════════════════════════════════════════════════════
#   AJOUTER FACTURE
# ══════════════════════════════════════════════════════════
def ajouter_facture(table):
    clients = session.query(Client).all()
    win, body, t = _fenetre_modale("➕  Nouvelle Facture", 480, 660)

    _lbl(body, "N Facture", t)
    e_num = _entry(body, t)

    _lbl(body, "Client", t)
    cb_client = _combo(body, t, [c.nom for c in clients])

    # Prix HT + TVA cote a cote
    row1 = tk.Frame(body, bg=t["bg"])
    row1.pack(fill="x", pady=(10, 0))

    col1 = tk.Frame(row1, bg=t["bg"])
    col1.pack(side="left", expand=True, fill="x", padx=(0, 8))
    tk.Label(col1, text="Prix HT (MAD)", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    e_ht = tk.Entry(col1, font=("Arial", 11), bg=t["card"], fg=t["text"],
                    insertbackground=t["text"], relief="flat", bd=6)
    e_ht.pack(fill="x", ipady=4)

    col2 = tk.Frame(row1, bg=t["bg"])
    col2.pack(side="left", expand=True, fill="x")
    tk.Label(col2, text="TVA (%)", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    e_tva = tk.Entry(col2, font=("Arial", 11), bg=t["card"], fg=t["text"],
                     insertbackground=t["text"], relief="flat", bd=6)
    e_tva.insert(0, "20")
    e_tva.pack(fill="x", ipady=4)

    # Reduction
    _lbl(body, "Reduction (%)", t)
    e_red = _entry(body, t, default="0")

    # Apercu TTC
    lbl_ttc = tk.Label(body, text="",
                       bg=t["bg"], fg="#FF8F00",
                       font=("Arial", 11, "bold"), anchor="w")
    lbl_ttc.pack(fill="x", pady=(6, 0))

    def maj_ttc(*_):
        try:
            ht  = float(e_ht.get())
            tva = float(e_tva.get())
            red = float(e_red.get() or 0)
            ttc = round(ht * (1 + tva / 100) * (1 - red / 100), 2)
            red_mnt = round(ht * (1 + tva / 100) * red / 100, 2)
            lbl_ttc.config(
                text=f"HT: {ht:,.2f}  →  TTC avant remise: {ht*(1+tva/100):,.2f}  "
                     f"−  Remise: {red_mnt:,.2f}  =  TTC final: {ttc:,.2f} MAD")
        except ValueError:
            lbl_ttc.config(text="")

    e_ht.bind("<KeyRelease>",  maj_ttc)
    e_tva.bind("<KeyRelease>", maj_ttc)
    e_red.bind("<KeyRelease>", maj_ttc)

    # Statut + Mode paiement cote a cote
    row2 = tk.Frame(body, bg=t["bg"])
    row2.pack(fill="x", pady=(10, 0))

    col3 = tk.Frame(row2, bg=t["bg"])
    col3.pack(side="left", expand=True, fill="x", padx=(0, 8))
    tk.Label(col3, text="Statut", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    cb_statut = ttk.Combobox(col3, values=STATUTS,
                             font=("Arial", 11), state="readonly")
    cb_statut.set("En attente")
    cb_statut.pack(fill="x", ipady=4)

    col4 = tk.Frame(row2, bg=t["bg"])
    col4.pack(side="left", expand=True, fill="x")
    tk.Label(col4, text="Mode de paiement", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    cb_paiement = ttk.Combobox(col4, values=PAIEMENTS,
                               font=("Arial", 11), state="readonly")
    cb_paiement.set("Especes")
    cb_paiement.pack(fill="x", ipady=4)

    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", pady=(12, 4))

    def sauvegarder():
        try:
            client = session.query(Client).filter_by(
                nom=cb_client.var.get()).first()
            if not client:
                messagebox.showerror("Erreur", "Selectionnez un client.", parent=win)
                return
            if not e_num.get().strip():
                messagebox.showerror("Erreur", "Le numero de facture est obligatoire.", parent=win)
                return

            ht  = float(e_ht.get())
            tva = float(e_tva.get())
            red = float(e_red.get() or 0)
            ttc = round(ht * (1 + tva / 100) * (1 - red / 100), 2)

            f = Facture(
                numero_facture=e_num.get().strip(),
                client_id=client.id,
                prix_ht=ht, tva=tva,
                prix_ttc=ttc, reduction=red,
                statut=cb_statut.get(),
                mode_paiement=cb_paiement.get(),
            )
            session.add(f)
            session.commit()
            charger_factures(table)
            win.destroy()
            messagebox.showinfo("Succes", "Facture ajoutee avec succes !")
        except Exception as ex:
            messagebox.showerror("Erreur", f"Erreur : {ex}", parent=win)

    _btn_ok(btn_frame, "💾  Sauvegarder", sauvegarder).pack(
        side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   MODIFIER FACTURE
# ══════════════════════════════════════════════════════════
def modifier_facture(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une facture a modifier !")
        return

    valeurs = table.item(sel[0])["values"]
    facture = session.query(Facture).filter_by(id=valeurs[0]).first()
    if not facture:
        return

    clients       = session.query(Client).all()
    client_actuel = session.query(Client).filter_by(id=facture.client_id).first()

    win, body, t = _fenetre_modale("✏️  Modifier la Facture", 480, 660)

    _lbl(body, "N Facture", t)
    e_num = _entry(body, t, default=facture.numero_facture or "")

    _lbl(body, "Client", t)
    cb_client = _combo(body, t, [c.nom for c in clients],
                       default=client_actuel.nom if client_actuel else "")

    # Prix HT + TVA cote a cote
    row1 = tk.Frame(body, bg=t["bg"])
    row1.pack(fill="x", pady=(10, 0))

    col1 = tk.Frame(row1, bg=t["bg"])
    col1.pack(side="left", expand=True, fill="x", padx=(0, 8))
    tk.Label(col1, text="Prix HT (MAD)", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    e_ht = tk.Entry(col1, font=("Arial", 11), bg=t["card"], fg=t["text"],
                    insertbackground=t["text"], relief="flat", bd=6)
    e_ht.insert(0, str(facture.prix_ht or ""))
    e_ht.pack(fill="x", ipady=4)

    col2 = tk.Frame(row1, bg=t["bg"])
    col2.pack(side="left", expand=True, fill="x")
    tk.Label(col2, text="TVA (%)", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    e_tva = tk.Entry(col2, font=("Arial", 11), bg=t["card"], fg=t["text"],
                     insertbackground=t["text"], relief="flat", bd=6)
    e_tva.insert(0, str(facture.tva or "20"))
    e_tva.pack(fill="x", ipady=4)

    _lbl(body, "Reduction (%)", t)
    e_red = _entry(body, t, default=str(facture.reduction or "0"))

    # Apercu TTC
    lbl_ttc = tk.Label(body,
                       text=f"Prix TTC actuel : {float(facture.prix_ttc or 0):,.2f} MAD",
                       bg=t["bg"], fg="#FF8F00",
                       font=("Arial", 11, "bold"), anchor="w")
    lbl_ttc.pack(fill="x", pady=(6, 0))

    def maj_ttc(*_):
        try:
            ht  = float(e_ht.get())
            tva = float(e_tva.get())
            red = float(e_red.get() or 0)
            ttc = round(ht * (1 + tva / 100) * (1 - red / 100), 2)
            lbl_ttc.config(text=f"Prix TTC calcule : {ttc:,.2f} MAD")
        except ValueError:
            lbl_ttc.config(text="")

    e_ht.bind("<KeyRelease>",  maj_ttc)
    e_tva.bind("<KeyRelease>", maj_ttc)
    e_red.bind("<KeyRelease>", maj_ttc)

    row2 = tk.Frame(body, bg=t["bg"])
    row2.pack(fill="x", pady=(10, 0))

    col3 = tk.Frame(row2, bg=t["bg"])
    col3.pack(side="left", expand=True, fill="x", padx=(0, 8))
    tk.Label(col3, text="Statut", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    cb_statut = ttk.Combobox(col3, values=STATUTS,
                             font=("Arial", 11), state="readonly")
    cb_statut.set(facture.statut or "En attente")
    cb_statut.pack(fill="x", ipady=4)

    col4 = tk.Frame(row2, bg=t["bg"])
    col4.pack(side="left", expand=True, fill="x")
    tk.Label(col4, text="Mode de paiement", bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(0, 2))
    cb_paiement = ttk.Combobox(col4, values=PAIEMENTS,
                               font=("Arial", 11), state="readonly")
    cb_paiement.set(facture.mode_paiement or "Especes")
    cb_paiement.pack(fill="x", ipady=4)

    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", pady=(12, 4))

    def sauvegarder():
        try:
            client = session.query(Client).filter_by(
                nom=cb_client.var.get()).first()
            if not client:
                messagebox.showerror("Erreur", "Selectionnez un client.", parent=win)
                return

            ht  = float(e_ht.get())
            tva = float(e_tva.get())
            red = float(e_red.get() or 0)
            ttc = round(ht * (1 + tva / 100) * (1 - red / 100), 2)

            facture.numero_facture = e_num.get().strip()
            facture.client_id      = client.id
            facture.prix_ht        = ht
            facture.tva            = tva
            facture.reduction      = red
            facture.prix_ttc       = ttc
            facture.statut         = cb_statut.get()
            facture.mode_paiement  = cb_paiement.get()
            session.commit()
            charger_factures(table)
            win.destroy()
            messagebox.showinfo("Succes", "Facture modifiee avec succes !")
        except Exception as ex:
            messagebox.showerror("Erreur", f"Erreur : {ex}", parent=win)

    _btn_ok(btn_frame, "💾  Sauvegarder", sauvegarder).pack(
        side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   SUPPRIMER FACTURE
# ══════════════════════════════════════════════════════════
def supprimer_facture(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une facture a supprimer !")
        return

    valeurs = table.item(sel[0])["values"]
    facture = session.query(Facture).filter_by(id=valeurs[0]).first()
    if not facture:
        return

    client = session.query(Client).filter_by(id=facture.client_id).first()

    if messagebox.askyesno(
        "Confirmation",
        f"Supprimer la facture {facture.numero_facture or facture.id} ?\n"
        f"Client  : {client.nom if client else 'N/A'}\n"
        f"Montant : {float(facture.prix_ttc or 0):,.2f} MAD\n\n"
        "Cette action est irreversible.",
    ):
        session.delete(facture)
        session.commit()
        charger_factures(table)
        messagebox.showinfo("Succes", "Facture supprimee avec succes !")


# ══════════════════════════════════════════════════════════
#   GENERER PDF FACTURE (reportlab)
# ══════════════════════════════════════════════════════════
def _generer_pdf(facture, client, chemin):
    if PDF_OK:
        _generer_pdf_reportlab(facture, client, chemin)
    elif FPDF_OK:
        _generer_pdf_fpdf(facture, client, chemin)
    else:
        raise RuntimeError(
            "Aucune librairie PDF disponible.\n"
            "Installez : pip install reportlab\n"
            "           ou pip install fpdf2"
        )


def _generer_pdf_reportlab(facture, client, chemin):
    doc    = SimpleDocTemplate(chemin, pagesize=A4,
                               leftMargin=2*cm, rightMargin=2*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []

    # ── En-tête ──────────────────────────────
    titre_style = ParagraphStyle(
        "Titre", parent=styles["Title"],
        fontSize=24, textColor=colors.HexColor("#E65100"),
        spaceAfter=2, alignment=TA_CENTER,
    )
    story.append(Paragraph("FACTURE", titre_style))
    story.append(Paragraph("VentePro",
        ParagraphStyle("Sub", parent=styles["Normal"],
                       fontSize=11, textColor=colors.HexColor("#FF8F00"),
                       spaceAfter=16, alignment=TA_CENTER)))

    story.append(HRFlowable(width="100%", thickness=2,
                             color=colors.HexColor("#E65100"), spaceAfter=12))

    # ── Infos facture + client ────────────────
    info_data = [
        ["N Facture :", facture.numero_facture or "-",
         "Client :", client.nom if client else "N/A"],
        ["Date :", str(date.today()),
         "Email :", client.email if client else "-"],
        ["Statut :", facture.statut or "-",
         "Telephone :", client.telephone if client else "-"],
        ["Paiement :", facture.mode_paiement or "-",
         "Ville :", client.ville if client else "-"],
    ]

    info_tbl = Table(info_data, colWidths=[3*cm, 5.5*cm, 3*cm, 5.5*cm])
    info_tbl.setStyle(TableStyle([
        ("FONTNAME",  (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE",  (0, 0), (-1, -1), 9.5),
        ("FONTNAME",  (0, 0), (0, -1), "Helvetica-Bold"),
        ("FONTNAME",  (2, 0), (2, -1), "Helvetica-Bold"),
        ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#E65100")),
        ("TEXTCOLOR", (2, 0), (2, -1), colors.HexColor("#E65100")),
        ("ROWBACKGROUNDS", (0, 0), (-1, -1),
         [colors.HexColor("#FFF8F0"), colors.white]),
        ("TOPPADDING",    (0, 0), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ("GRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#E8D8C8")),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 0.7*cm))

    # ── Tableau de calcul ─────────────────────
    ht  = float(facture.prix_ht or 0)
    tva = float(facture.tva or 0)
    red = float(facture.reduction or 0)
    ttc_avant = round(ht * (1 + tva / 100), 2)
    red_mnt   = round(ttc_avant * red / 100, 2)
    ttc_final = float(facture.prix_ttc or 0)

    detail_data = [
        ["Description", "Montant (MAD)"],
        ["Prix Hors Taxes (HT)", f"{ht:,.2f}"],
        [f"TVA ({tva:.0f}%)", f"{round(ht * tva / 100, 2):,.2f}"],
        ["Total TTC avant remise", f"{ttc_avant:,.2f}"],
        [f"Remise ({red:.0f}%)", f"- {red_mnt:,.2f}"],
        ["TOTAL TTC A PAYER", f"{ttc_final:,.2f}"],
    ]

    det_tbl = Table(detail_data, colWidths=[12*cm, 5*cm])
    det_style = TableStyle([
        # En-tete
        ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#E65100")),
        ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 10),
        ("ALIGN",         (1, 0), (1, -1), "RIGHT"),
        ("ALIGN",         (0, 0), (0, -1), "LEFT"),
        # Lignes
        ("FONTSIZE",      (0, 1), (-1, -2), 9.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -2),
         [colors.white, colors.HexColor("#FFF8F0")]),
        ("TOPPADDING",    (0, 0), (-1, -1), 7),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 7),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("RIGHTPADDING",  (0, 0), (-1, -1), 10),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E8D8C8")),
        # Ligne total
        ("BACKGROUND",  (0, -1), (-1, -1), colors.HexColor("#BF360C")),
        ("TEXTCOLOR",   (0, -1), (-1, -1), colors.white),
        ("FONTNAME",    (0, -1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE",    (0, -1), (-1, -1), 11),
    ])
    det_tbl.setStyle(det_style)
    story.append(det_tbl)
    story.append(Spacer(1, 0.6*cm))

    # ── Badge statut ──────────────────────────
    statut_colors_hex = {
        "Payee": "#2E7D32", "Non payee": "#C62828",
        "En attente": "#E65100", "Annulee": "#757575",
    }
    statut_col = statut_colors_hex.get(facture.statut or "", "#555555")

    statut_data = [[f"Statut : {facture.statut or '-'}"]]
    statut_tbl  = Table(statut_data, colWidths=[17*cm])
    statut_tbl.setStyle(TableStyle([
        ("BACKGROUND",    (0, 0), (-1, -1), colors.HexColor(statut_col)),
        ("TEXTCOLOR",     (0, 0), (-1, -1), colors.white),
        ("FONTNAME",      (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, -1), 11),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING",    (0, 0), (-1, -1), 8),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
    ]))
    story.append(statut_tbl)
    story.append(Spacer(1, 0.5*cm))

    # ── Pied de page ──────────────────────────
    story.append(HRFlowable(width="100%", thickness=1,
                             color=colors.HexColor("#E8D8C8"), spaceBefore=8))
    story.append(Paragraph(
        "VentePro — Merci pour votre confiance.",
        ParagraphStyle("Footer", parent=styles["Normal"],
                       fontSize=8.5, textColor=colors.HexColor("#999999"),
                       alignment=TA_CENTER, spaceAfter=0)))

    doc.build(story)


def _generer_pdf_fpdf(facture, client, chemin):
    """Fallback avec fpdf si reportlab absent."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_fill_color(230, 81, 0)
    pdf.rect(0, 0, 210, 18, "F")
    pdf.set_font("Arial", "B", 20)
    pdf.set_text_color(255, 255, 255)
    pdf.cell(0, 18, txt="FACTURE — VentePro", ln=True, align="C")
    pdf.set_text_color(0, 0, 0)
    pdf.ln(8)

    pdf.set_font("Arial", "B", 12)
    pdf.set_text_color(230, 81, 0)
    pdf.cell(0, 8, txt=f"N Facture : {facture.numero_facture or '-'}", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Arial", size=11)
    pdf.cell(0, 8, txt=f"Client    : {client.nom if client else 'N/A'}", ln=True)
    pdf.cell(0, 8, txt=f"Email     : {client.email if client else '-'}", ln=True)
    pdf.cell(0, 8, txt=f"Date      : {date.today()}", ln=True)
    pdf.ln(6)

    pdf.set_font("Arial", "B", 11)
    pdf.set_text_color(230, 81, 0)
    pdf.cell(0, 8, txt="Detail de la facture :", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.set_font("Arial", size=11)
    pdf.cell(0, 8, txt=f"Prix HT   : {float(facture.prix_ht or 0):,.2f} MAD", ln=True)
    pdf.cell(0, 8, txt=f"TVA       : {float(facture.tva or 0):.0f} %", ln=True)
    pdf.cell(0, 8, txt=f"Reduction : {float(facture.reduction or 0):.0f} %", ln=True)
    pdf.ln(4)
    pdf.set_font("Arial", "B", 13)
    pdf.set_text_color(191, 54, 12)
    pdf.cell(0, 10, txt=f"TOTAL TTC : {float(facture.prix_ttc or 0):,.2f} MAD", ln=True)
    pdf.set_text_color(0, 0, 0)
    pdf.ln(6)
    pdf.set_font("Arial", size=11)
    pdf.cell(0, 8, txt=f"Statut    : {facture.statut or '-'}", ln=True)
    pdf.cell(0, 8, txt=f"Paiement  : {facture.mode_paiement or '-'}", ln=True)

    pdf.output(chemin)


# ══════════════════════════════════════════════════════════
#   IMPRIMER FACTURE
# ══════════════════════════════════════════════════════════
def imprimer_facture(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une facture a imprimer !")
        return

    valeurs = table.item(sel[0])["values"]
    facture = session.query(Facture).filter_by(id=valeurs[0]).first()
    client  = session.query(Client).filter_by(id=facture.client_id).first()

    # Choisir ou sauvegarder
    chemin = filedialog.asksaveasfilename(
        title="Enregistrer la facture PDF",
        defaultextension=".pdf",
        filetypes=[("PDF", "*.pdf"), ("Tous", "*.*")],
        initialfile=f"facture_{facture.numero_facture or facture.id}.pdf",
    )
    if not chemin:
        return

    try:
        _generer_pdf(facture, client, chemin)
        if messagebox.askyesno("PDF cree",
                               f"Facture enregistree :\n{chemin}\n\nOuvrir maintenant ?"):
            os.startfile(chemin)
    except Exception as ex:
        messagebox.showerror("Erreur PDF", f"Erreur lors de la generation :\n{ex}")


# ══════════════════════════════════════════════════════════
#   ENVOYER FACTURE
# ══════════════════════════════════════════════════════════
def envoyer_facture(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une facture a envoyer !")
        return

    valeurs = table.item(sel[0])["values"]
    facture = session.query(Facture).filter_by(id=valeurs[0]).first()
    client  = session.query(Client).filter_by(id=facture.client_id).first()

    t = get_theme()
    win, body, t = _fenetre_modale("📧  Envoyer la Facture", 460, 500)
    win.resizable(True, True)
    win.minsize(400, 460)

    # Resume
    info_bg = "#1A1A35" if t["bg"] in ("#1A1A2E",) else "#FFF0E5"
    resume  = tk.Frame(body, bg=info_bg)
    resume.pack(fill="x", pady=(0, 10))

    for label, val in [
        ("N Facture",  facture.numero_facture or "-"),
        ("Client",     client.nom if client else "N/A"),
        ("Montant TTC",f"{float(facture.prix_ttc or 0):,.2f} MAD"),
        ("Statut",     facture.statut or "-"),
    ]:
        row = tk.Frame(resume, bg=info_bg)
        row.pack(fill="x", padx=14, pady=3)
        tk.Label(row, text=f"{label} :", bg=info_bg,
                 fg="#FF8F00", font=("Arial", 10, "bold"),
                 width=12, anchor="w").pack(side="left")
        tk.Label(row, text=val, bg=info_bg,
                 fg=t["text"], font=("Arial", 10)).pack(side="left")

    _lbl(body, "Email destinataire", t)
    email_def = getattr(client, "email", "") if client else ""
    e_email = _entry(body, t, default=email_def or "")

    _lbl(body, "Message (optionnel)", t)
    msg_def = (
        f"Bonjour {client.nom if client else ''},\n\n"
        f"Veuillez trouver ci-joint votre facture N {facture.numero_facture or facture.id} "
        f"d'un montant de {float(facture.prix_ttc or 0):,.2f} MAD.\n\n"
        "Cordialement."
    )
    e_msg = tk.Text(body, height=5, font=("Arial", 10),
                    bg=t["card"], fg=t["text"],
                    insertbackground=t["text"],
                    relief="flat", bd=6, wrap="word")
    e_msg.insert("1.0", msg_def)
    e_msg.pack(fill="x")

    _lbl(body, "Methode d'envoi", t)
    var_methode   = tk.StringVar(value="Email")
    frame_methode = tk.Frame(body, bg=t["bg"])
    frame_methode.pack(fill="x")

    for m in ["Email", "WhatsApp", "Impression"]:
        tk.Radiobutton(frame_methode, text=m,
                       variable=var_methode, value=m,
                       bg=t["bg"], fg=t["text"],
                       selectcolor=t["card"],
                       font=("Arial", 11),
                       activebackground=t["bg"]).pack(side="left", padx=8, pady=4)

    def confirmer_envoi():
        methode = var_methode.get()
        try:
            if methode == "Impression":
                tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                tmp.close()
                _generer_pdf(facture, client, tmp.name)
                os.startfile(tmp.name, "print")
                win.destroy()
                messagebox.showinfo("Impression",
                    f"Facture {facture.numero_facture or facture.id} envoyee a l'imprimante.")
            else:
                if not e_email.get().strip():
                    messagebox.showerror("Erreur",
                        "Saisissez une adresse email.", parent=win)
                    return
                try:
                    from utils.pdf_export import envoyer_facture_email, _charger_config, _config_valide, ouvrir_config_email

                    # Vérifier si la config est remplie
                    cfg = _charger_config()
                    if not _config_valide(cfg):
                        rep = messagebox.askyesno(
                            "Email non configure",
                            "L'email d'envoi n'est pas configure.\n\n"
                            "Voulez-vous ouvrir la configuration maintenant ?",
                            parent=win)
                        if rep:
                            ouvrir_config_email(parent=win)
                        return

                    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
                    tmp.close()
                    _generer_pdf(facture, client, tmp.name)

                    # Récupérer le message saisi par l'utilisateur
                    message_utilisateur = e_msg.get("1.0", "end").strip()
                    sujet = f"Votre facture N° {facture.numero_facture or facture.id} — VentePro"

                    try:
                        succes = envoyer_facture_email(
                            email_destinataire=e_email.get().strip(),
                            nom_client=client.nom if client else "",
                            fichier_pdf=tmp.name,
                            sujet=sujet,
                            message=message_utilisateur,
                        )
                        win.destroy()
                        messagebox.showinfo("Succes",
                            f"✅ Facture N° {facture.numero_facture or facture.id}\n"
                            f"envoyee avec succes a :\n{e_email.get().strip()}")
                    except ValueError as ve:
                        # Erreur d'authentification ou de config
                        msg_err = str(ve)
                        rep = messagebox.askyesno(
                            "Echec d'envoi",
                            f"{msg_err}\n\n"
                            "Voulez-vous ouvrir la configuration email ?",
                            parent=win)
                        if rep:
                            ouvrir_config_email(parent=win)

                except ImportError:
                    win.destroy()
                    messagebox.showinfo("Info",
                        f"Methode : {methode}\n"
                        f"Destinataire : {e_email.get()}\n\n"
                        "Module utils.pdf_export non disponible.\n"
                        "Verifiez que le fichier utils/pdf_export.py existe.")
        except Exception as ex:
            messagebox.showerror("Erreur", f"Erreur : {ex}", parent=win)

    # Bouton configuration email rapide
    cfg_frame = tk.Frame(body, bg=t["bg"])
    cfg_frame.pack(fill="x", pady=(6, 0))
    tk.Button(cfg_frame,
              text="⚙️  Configurer l'email d'envoi",
              bg=t["card"], fg="#FF8F00",
              font=("Arial", 9, "bold"),
              relief="flat", cursor="hand2",
              padx=10, pady=4,
              activebackground=t["bg"],
              activeforeground="#E65100",
              command=lambda: _ouvrir_cfg(win)
              ).pack(side="left")

    def _ouvrir_cfg(parent_win):
        try:
            from utils.pdf_export import ouvrir_config_email
            ouvrir_config_email(parent=parent_win)
        except ImportError:
            messagebox.showerror("Erreur",
                "utils/pdf_export.py introuvable.", parent=parent_win)

    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", pady=(12, 4))

    _btn_ok(btn_frame, "✅  Valider l'envoi",
            confirmer_envoi, "#1565C0", "#0D47A1").pack(
            side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   CONSULTER FACTURE — fiche professionnelle complète
# ══════════════════════════════════════════════════════════
def consulter_facture(table):
    """Affiche la fiche complète professionnelle d'une facture."""
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une facture a consulter !")
        return

    valeurs = table.item(sel[0])["values"]
    facture = session.query(Facture).filter_by(id=valeurs[0]).first()
    if not facture:
        return

    client  = session.query(Client).filter_by(id=facture.client_id).first()
    t       = get_theme()
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")

    win = tk.Toplevel()
    win.title(f"Facture N° {facture.numero_facture or facture.id}")
    win.geometry("640x700")
    win.configure(bg=t["bg"])
    win.resizable(True, True)
    win.grab_set()

    # ── En-tête coloré ───────────────────────────────────────
    hdr = tk.Frame(win, bg="#E65100", height=56)
    hdr.pack(fill="x")
    hdr.pack_propagate(False)

    tk.Label(hdr, text=f"🧾  FACTURE N° {facture.numero_facture or facture.id}",
             bg="#E65100", fg="white",
             font=("Arial", 16, "bold")).pack(side="left", padx=20, pady=14)

    # Badge statut dans l'en-tête
    statut_col_hdr = {
        "Payee": "#2E7D32", "Non payee": "#C62828",
        "En attente": "#FF8F00", "Annulee": "#757575",
    }
    scol = statut_col_hdr.get(facture.statut or "", "#555555")
    icone_s = STATUT_ICONE.get(facture.statut or "", "")
    tk.Label(hdr,
             text=f" {icone_s} {facture.statut or '-'} ",
             bg=scol, fg="white",
             font=("Arial", 10, "bold"),
             padx=8, pady=4).pack(side="right", padx=16, pady=14)

    # ── Corps scrollable ─────────────────────────────────────
    canvas_frame = tk.Frame(win, bg=t["bg"])
    canvas_frame.pack(fill="both", expand=True)

    canvas = tk.Canvas(canvas_frame, bg=t["bg"], highlightthickness=0)
    scrollbar = ttk.Scrollbar(canvas_frame, orient="vertical",
                              command=canvas.yview)
    canvas.configure(yscrollcommand=scrollbar.set)
    scrollbar.pack(side="right", fill="y")
    canvas.pack(side="left", fill="both", expand=True)

    body = tk.Frame(canvas, bg=t["bg"])
    body_id = canvas.create_window((0, 0), window=body, anchor="nw")

    def _on_configure(e):
        canvas.configure(scrollregion=canvas.bbox("all"))
        canvas.itemconfig(body_id, width=canvas.winfo_width())

    body.bind("<Configure>", _on_configure)
    canvas.bind("<Configure>", lambda e: canvas.itemconfig(
        body_id, width=canvas.winfo_width()))

    # Scroll souris
    def _on_mousewheel(e):
        canvas.yview_scroll(int(-1 * (e.delta / 120)), "units")
    canvas.bind_all("<MouseWheel>", _on_mousewheel)
    win.bind("<Destroy>", lambda e: canvas.unbind_all("<MouseWheel>"))

    info_bg  = "#1A1A35" if is_dark else "#FFF8F0"
    sep_col  = "#2A2A4A" if is_dark else "#E8D8C8"

    def _section(titre, couleur="#E65100"):
        f = tk.Frame(body, bg=t["bg"])
        f.pack(fill="x", padx=16, pady=(14, 4))
        tk.Label(f, text=titre, bg=t["bg"], fg=couleur,
                 font=("Arial", 12, "bold"), anchor="w").pack(side="left")
        tk.Frame(f, bg=sep_col, height=1).pack(
            side="left", fill="x", expand=True, padx=(10, 0), pady=6)

    def _ligne(parent, label, val, val_color=None):
        bg = parent.cget("bg") if hasattr(parent, "cget") else info_bg
        row = tk.Frame(parent, bg=info_bg)
        row.pack(fill="x", padx=14, pady=3)
        tk.Label(row, text=f"{label} :", bg=info_bg, fg="#FF8F00",
                 font=("Arial", 10, "bold"), width=16, anchor="w").pack(side="left")
        tk.Label(row, text=str(val), bg=info_bg,
                 fg=val_color or t["text"],
                 font=("Arial", 10, "bold" if val_color else "normal")).pack(side="left")

    # ── Section 1 : Informations facture ─────────────────────
    _section("📋  Informations de la Facture")
    bloc1 = tk.Frame(body, bg=info_bg,
                     highlightbackground=sep_col,
                     highlightthickness=1)
    bloc1.pack(fill="x", padx=16, pady=(0, 4))
    tk.Frame(bloc1, bg="#E65100", height=2).pack(fill="x")

    _ligne(bloc1, "N° Facture",      facture.numero_facture or "-")
    _ligne(bloc1, "Date d'emission", str(date.today()))
    _ligne(bloc1, "Mode de paiement",
           f"{PAIEMENT_ICONE.get(facture.mode_paiement or '', '')} {facture.mode_paiement or '-'}")
    _ligne(bloc1, "Statut",
           f"{STATUT_ICONE.get(facture.statut or '', '')} {facture.statut or '-'}",
           val_color=statut_col_hdr.get(facture.statut or "", None))
    tk.Frame(bloc1, bg=t["bg"], height=8).pack()

    # ── Section 2 : Informations client ──────────────────────
    _section("👤  Destinataire", "#1565C0")
    bloc2 = tk.Frame(body, bg=info_bg,
                     highlightbackground=sep_col,
                     highlightthickness=1)
    bloc2.pack(fill="x", padx=16, pady=(0, 4))
    tk.Frame(bloc2, bg="#1565C0", height=2).pack(fill="x")

    if client:
        _ligne(bloc2, "Nom / Raison sociale", client.nom or "-")
        _ligne(bloc2, "Email",                client.email or "-")
        _ligne(bloc2, "Telephone",            getattr(client, "telephone", "-") or "-")
        _ligne(bloc2, "Adresse",              getattr(client, "adresse",   "-") or "-")
        _ligne(bloc2, "Ville",                getattr(client, "ville",     "-") or "-")
    else:
        _ligne(bloc2, "Client", "N/A")
    tk.Frame(bloc2, bg=t["bg"], height=8).pack()

    # ── Section 3 : Détail financier ─────────────────────────
    _section("💰  Détail Financier", "#2E7D32")
    bloc3 = tk.Frame(body, bg=info_bg,
                     highlightbackground=sep_col,
                     highlightthickness=1)
    bloc3.pack(fill="x", padx=16, pady=(0, 4))
    tk.Frame(bloc3, bg="#2E7D32", height=2).pack(fill="x")

    ht       = float(facture.prix_ht  or 0)
    tva      = float(facture.tva      or 0)
    red      = float(facture.reduction or 0)
    ttc_av   = round(ht * (1 + tva / 100), 2)
    tva_mnt  = round(ht * tva / 100, 2)
    red_mnt  = round(ttc_av * red / 100, 2)
    ttc_fin  = float(facture.prix_ttc or 0)

    _ligne(bloc3, "Prix Hors Taxe (HT)", f"{ht:,.2f} MAD")
    _ligne(bloc3, f"TVA ({tva:.0f}%)",   f"+ {tva_mnt:,.2f} MAD")
    _ligne(bloc3, "TTC avant remise",    f"= {ttc_av:,.2f} MAD")
    if red > 0:
        _ligne(bloc3, f"Remise ({red:.0f}%)", f"- {red_mnt:,.2f} MAD")

    # Ligne total mise en évidence
    total_frame = tk.Frame(bloc3, bg="#E65100")
    total_frame.pack(fill="x", padx=14, pady=(6, 6))
    tk.Label(total_frame, text="💳  TOTAL TTC À PAYER :",
             bg="#E65100", fg="white",
             font=("Arial", 12, "bold")).pack(side="left", padx=12, pady=8)
    tk.Label(total_frame, text=f"{ttc_fin:,.2f} MAD",
             bg="#E65100", fg="white",
             font=("Arial", 14, "bold")).pack(side="right", padx=12, pady=8)
    tk.Frame(bloc3, bg=t["bg"], height=8).pack()

    # ── Boutons d'action ─────────────────────────────────────
    tk.Frame(body, bg=sep_col, height=1).pack(fill="x", padx=16, pady=(10, 0))
    btn_row = tk.Frame(body, bg=t["bg"])
    btn_row.pack(fill="x", padx=16, pady=10)

    _btn_ok(btn_row, "📤  Exporter PDF",
            lambda: [win.destroy(), imprimer_facture(table)],
            "#1565C0", "#0D47A1").pack(side="left", padx=(0, 6))
    _btn_ok(btn_row, "📧  Envoyer",
            lambda: [win.destroy(), envoyer_facture(table)],
            "#2E7D32", "#1B5E20").pack(side="left", padx=(0, 6))
    _btn_ok(btn_row, "✏️  Modifier",
            lambda: [win.destroy(), modifier_facture(table)],
            "#FF8F00", "#E65100").pack(side="left", padx=(0, 6))
    tk.Button(btn_row, text="Fermer", command=win.destroy,
              bg=t["card"], fg=t["text"],
              font=("Arial", 11, "bold"),
              relief="flat", cursor="hand2", pady=9,
              padx=16).pack(side="right")


# ══════════════════════════════════════════════════════════
#   DETAIL FACTURE (double-clic) — alias vers consulter
# ══════════════════════════════════════════════════════════
def detail_facture(table):
    """Double-clic → ouvre la fiche détail professionnelle."""
    consulter_facture(table)