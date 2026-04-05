import customtkinter as ctk
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from database.db import session
from database.models import Livraison, Client, Produit, Devis
from datetime import date
from utils.theme import get_theme
import os


# ══════════════════════════════════════════════════════════
#   IMPORTS OPTIONNELS
# ══════════════════════════════════════════════════════════
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_OK = True
except ImportError:
    EXCEL_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_OK = True
except ImportError:
    WORD_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.units import cm
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import json
    JSON_OK = True
except ImportError:
    JSON_OK = False

import csv


# ══════════════════════════════════════════════════════════
#   CONSTANTES
# ══════════════════════════════════════════════════════════
STATUTS = ["En attente", "En cours", "Livre", "Annule"]

STATUT_ICONE = {
    "En attente": "⏳",
    "En cours":   "🔄",
    "Livre":      "✅",
    "Annule":     "❌",
}

EN_TETES = ["ID", "N BL", "Client", "Adresse", "Produit",
            "Prix TTC", "Quantite", "Statut"]
CHAMPS   = ["id", "numero_bl", "client", "adresse", "produit",
            "prix", "quantite", "statut"]


# ══════════════════════════════════════════════════════════
#   STYLE TREEVIEW
# ══════════════════════════════════════════════════════════
def _style_treeview(is_dark):
    style = ttk.Style()
    style.theme_use("clam")

    bg_tree = "#12122A" if is_dark else "#FFFFFF"
    bg_head = "#E65100"
    bg_sel  = "#FF6D00" if is_dark else "#FFB74D"
    fg_text = "#E8E0FF" if is_dark else "#1A1A1A"
    bg_row1 = "#1A1A35" if is_dark else "#FFFFFF"
    bg_row2 = "#12122A" if is_dark else "#FFF8F0"

    style.configure("LV.Treeview",
                    background=bg_tree, foreground=fg_text,
                    rowheight=32, fieldbackground=bg_tree,
                    borderwidth=0, font=("Arial", 10))
    style.configure("LV.Treeview.Heading",
                    background=bg_head, foreground="white",
                    font=("Arial", 10, "bold"),
                    relief="flat", borderwidth=0)
    style.map("LV.Treeview",
              background=[("selected", bg_sel)],
              foreground=[("selected", "#FFFFFF")])
    style.map("LV.Treeview.Heading",
              background=[("active", "#BF360C")])
    return bg_row1, bg_row2


# ══════════════════════════════════════════════════════════
#   HELPERS UI
# ══════════════════════════════════════════════════════════
def _fenetre_modale(titre, largeur=480, hauteur=600):
    t   = get_theme()
    win = tk.Toplevel()
    win.title(titre)
    win.geometry(f"{largeur}x{hauteur}")
    win.configure(bg=t["bg"])
    win.resizable(False, False)
    win.grab_set()

    hdr = tk.Frame(win, bg="#E65100", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=titre, bg="#E65100", fg="white",
             font=("Arial", 14, "bold")).pack(pady=12)

    # Corps scrollable
    outer = tk.Frame(win, bg=t["bg"])
    outer.pack(fill="both", expand=True)

    canvas_s = tk.Canvas(outer, bg=t["bg"], highlightthickness=0)
    sb = ttk.Scrollbar(outer, orient="vertical", command=canvas_s.yview)
    canvas_s.configure(yscrollcommand=sb.set)
    sb.pack(side="right", fill="y")
    canvas_s.pack(side="left", fill="both", expand=True)

    body = tk.Frame(canvas_s, bg=t["bg"])
    win_id = canvas_s.create_window((0, 0), window=body, anchor="nw")

    def _on_resize(e):
        canvas_s.itemconfig(win_id, width=e.width)
    canvas_s.bind("<Configure>", _on_resize)

    def _on_scroll(e):
        canvas_s.configure(scrollregion=canvas_s.bbox("all"))
    body.bind("<Configure>", _on_scroll)

    # Scroll molette — bind local uniquement sur ce canvas
    def _wheel(e):
        try:
            canvas_s.yview_scroll(int(-1 * (e.delta / 120)), "units")
        except tk.TclError:
            pass

    canvas_s.bind("<MouseWheel>", _wheel)
    body.bind("<MouseWheel>", _wheel)

    # Desactiver le scroll quand la fenetre se ferme
    def _on_close():
        try:
            canvas_s.unbind("<MouseWheel>")
            body.unbind("<MouseWheel>")
        except Exception:
            pass
        win.destroy()

    win.protocol("WM_DELETE_WINDOW", _on_close)

    return win, body, t


def _lbl(body, text, t):
    tk.Label(body, text=text, bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(
             fill="x", padx=24, pady=(10, 2))


def _entry(body, t, default=""):
    e = tk.Entry(body, font=("Arial", 11),
                 bg=t["card"], fg=t["text"],
                 insertbackground=t["text"],
                 relief="flat", bd=6)
    e.pack(fill="x", padx=24, ipady=4)
    if default:
        e.insert(0, str(default))
    return e


def _combo(body, t, values, default=""):
    var = tk.StringVar(value=default)
    cb  = ttk.Combobox(body, textvariable=var, values=values,
                       font=("Arial", 11), state="readonly")
    cb.pack(fill="x", padx=24, ipady=4)
    cb.var = var
    return cb


def _btn_ok(parent, text, cmd, color="#E65100", hover="#BF360C"):
    return tk.Button(
        parent, text=text, command=cmd,
        bg=color, fg="white",
        font=("Arial", 12, "bold"),
        relief="flat", cursor="hand2",
        padx=16, pady=9,
        activebackground=hover,
        activeforeground="white",
    )


def _sep(body, t):
    tk.Frame(body,
             bg="#2A2A4A" if t["bg"] in ("#1A1A2E", "#0F0F1A") else "#E8D8C8",
             height=1).pack(fill="x", padx=24, pady=(14, 0))


# ══════════════════════════════════════════════════════════
#   TRI COLONNES
# ══════════════════════════════════════════════════════════
_tri_etat = {}

def _trier(table, col):
    rows = [(table.set(r, col), r) for r in table.get_children("")]
    rev  = _tri_etat.get(col, False)
    try:
        rows.sort(key=lambda x: float(
            x[0].replace(",", "").replace(" ", "")), reverse=rev)
    except ValueError:
        rows.sort(key=lambda x: x[0].lower(), reverse=rev)
    for i, (_, r) in enumerate(rows):
        table.move(r, "", i)
    _tri_etat[col] = not rev


# ══════════════════════════════════════════════════════════
#   MENUS DEROULANTS IMPORT / EXPORT
# ══════════════════════════════════════════════════════════
def _show_dropdown(btn_widget, items, t):
    """Affiche un menu deroulant generique sous btn_widget."""
    is_dark  = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
    bg_menu  = "#1A1A35" if is_dark else "#FFFFFF"
    bg_hover = "#252545" if is_dark else "#FFF0E5"
    border_c = "#2A2A4A" if is_dark else "#E8D8C8"

    menu = tk.Toplevel()
    menu.overrideredirect(True)
    menu.configure(bg=border_c)
    menu.attributes("-topmost", True)

    btn_widget.update_idletasks()
    x = btn_widget.winfo_rootx()
    y = btn_widget.winfo_rooty() + btn_widget.winfo_height() + 2
    menu.geometry(f"210x{len(items) * 40 + 16}+{x}+{y}")

    inner = tk.Frame(menu, bg=bg_menu, bd=0)
    inner.pack(fill="both", expand=True, padx=1, pady=1)

    tk.Label(inner, text="Choisir le format :",
             bg=bg_menu, fg="#FF8F00",
             font=("Arial", 8, "bold"),
             anchor="w").pack(fill="x", padx=10, pady=(6, 2))

    def _fermer():
        try:
            menu.destroy()
        except Exception:
            pass

    for label, couleur, cmd_fn in items:
        f   = tk.Frame(inner, bg=bg_menu, cursor="hand2")
        f.pack(fill="x", padx=4, pady=1)
        lbl = tk.Label(f, text=label, bg=bg_menu, fg=couleur,
                       font=("Arial", 10, "bold"),
                       anchor="w", padx=12, pady=8)
        lbl.pack(fill="x")

        def _enter(e, fr=f, lb=lbl):
            fr.configure(bg=bg_hover); lb.configure(bg=bg_hover)
        def _leave(e, fr=f, lb=lbl):
            fr.configure(bg=bg_menu);  lb.configure(bg=bg_menu)
        def _click(e, fn=cmd_fn):
            _fermer(); fn()

        for w in (f, lbl):
            w.bind("<Enter>",   _enter)
            w.bind("<Leave>",   _leave)
            w.bind("<Button-1>", _click)

    menu.bind("<FocusOut>", lambda e: _fermer())
    menu.focus_set()


def _show_dropdown_import(btn_widget, table, t):
    items = [
        ("📊  Excel (.xlsx)", "#1565C0", lambda: importer_fichier(table, "excel")),
        ("📄  CSV (.csv)",    "#00838F", lambda: importer_fichier(table, "csv")),
        ("📋  JSON (.json)",  "#6A1B9A", lambda: importer_fichier(table, "json")),
    ]
    _show_dropdown(btn_widget, items, t)


def _show_dropdown_export(btn_widget, table, t):
    items = [
        ("📊  Excel (.xlsx)",  "#1565C0", lambda: exporter_livraisons(table, "excel")),
        ("📝  Word (.docx)",   "#1976D2", lambda: exporter_livraisons(table, "word")),
        ("📕  PDF (.pdf)",     "#C62828", lambda: exporter_livraisons(table, "pdf")),
        ("📄  CSV (.csv)",     "#00838F", lambda: exporter_livraisons(table, "csv")),
        ("📋  JSON (.json)",   "#6A1B9A", lambda: exporter_livraisons(table, "json")),
    ]
    _show_dropdown(btn_widget, items, t)


# ══════════════════════════════════════════════════════════
#   AFFICHER LIVRAISONS
# ══════════════════════════════════════════════════════════
def afficher_livraisons(parent):
    for widget in parent.winfo_children():
        widget.destroy()

    t       = get_theme()
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")

    # En-tete
    header = ctk.CTkFrame(parent, fg_color=t["card"], corner_radius=0)
    header.pack(fill="x")
    ctk.CTkLabel(header, text="🚚  Gestion des Livraisons",
                 font=("Arial", 22, "bold"),
                 text_color="#E65100").pack(side="left", padx=24, pady=14)

    # Barre outils
    toolbar = ctk.CTkFrame(parent, fg_color="transparent")
    toolbar.pack(fill="x", padx=20, pady=(10, 4))

    btns_actions = [
        ("➕  Ajouter",      "#E65100", "#BF360C", lambda: ajouter_livraison(table)),
        ("✏️  Modifier",     "#FF8F00", "#E65100", lambda: modifier_livraison(table)),
        ("🗑️  Supprimer",   "#C62828", "#8B0000", lambda: supprimer_livraison(table)),
        ("🖨️  Imprimer",    "#2E7D32", "#1B5E20", lambda: imprimer_livraison(table)),
    ]
    for txt, fg, hover, cmd in btns_actions:
        ctk.CTkButton(toolbar, text=txt, width=128, height=36,
                      fg_color=fg, hover_color=hover,
                      font=("Arial", 11, "bold"),
                      corner_radius=8,
                      command=cmd).pack(side="left", padx=4)

    # Bouton Importer dropdown
    btn_imp = ctk.CTkButton(
        toolbar, text="📥  Importer ▾", width=138, height=36,
        fg_color="#1565C0", hover_color="#0D47A1",
        font=("Arial", 11, "bold"), corner_radius=8,
        command=lambda: _show_dropdown_import(btn_imp, table, t)
    )
    btn_imp.pack(side="left", padx=4)

    # Bouton Exporter dropdown
    btn_exp = ctk.CTkButton(
        toolbar, text="📤  Exporter ▾", width=138, height=36,
        fg_color="#6A1B9A", hover_color="#4A148C",
        font=("Arial", 11, "bold"), corner_radius=8,
        command=lambda: _show_dropdown_export(btn_exp, table, t)
    )
    btn_exp.pack(side="left", padx=4)

    # Recherche + compteur
    count_var = tk.StringVar()
    ctk.CTkLabel(toolbar, textvariable=count_var,
                 font=("Arial", 11),
                 text_color="#FF8F00").pack(side="right", padx=10)

    recherche_var = tk.StringVar()
    ctk.CTkEntry(toolbar,
                 placeholder_text="🔍  Rechercher...",
                 width=240, height=36,
                 textvariable=recherche_var,
                 font=("Arial", 11)).pack(side="right", padx=5)
    recherche_var.trace("w",
        lambda *a: rechercher(recherche_var.get(), table))

    # Tableau
    frame_table = ctk.CTkFrame(parent, corner_radius=12)
    frame_table.pack(fill="both", expand=True, padx=20, pady=(4, 16))

    bg_row1, bg_row2 = _style_treeview(is_dark)

    colonnes = ("ID", "N BL", "Client", "Adresse",
                "Produit", "Prix TTC", "Qte", "Statut")

    table = ttk.Treeview(frame_table, columns=colonnes,
                         show="headings", height=22,
                         style="LV.Treeview",
                         selectmode="browse")

    largeurs = {"ID": 45, "N BL": 100, "Client": 140,
                "Adresse": 180, "Produit": 130,
                "Prix TTC": 90, "Qte": 55, "Statut": 110}

    for col in colonnes:
        table.heading(col, text=col,
                      command=lambda c=col: _trier(table, c))
        table.column(col, width=largeurs.get(col, 100), anchor="center")

    table.tag_configure("pair",        background=bg_row1)
    table.tag_configure("impair",      background=bg_row2)
    table.tag_configure("Livre",       foreground="#43A047")
    table.tag_configure("Annule",      foreground="#E53935")
    table.tag_configure("En cours",    foreground="#42A5F5")
    table.tag_configure("En attente",  foreground="#FFA000")

    sb_v = ttk.Scrollbar(frame_table, orient="vertical",   command=table.yview)
    sb_h = ttk.Scrollbar(frame_table, orient="horizontal", command=table.xview)
    table.configure(yscroll=sb_v.set, xscroll=sb_h.set)
    sb_v.pack(side="right",  fill="y")
    sb_h.pack(side="bottom", fill="x")
    table.pack(fill="both", expand=True, padx=2, pady=2)

    table.bind("<Double-1>", lambda e: detail_livraison(table))

    charger_livraisons(table, count_var)
    return table


# ══════════════════════════════════════════════════════════
#   CHARGER LIVRAISONS
# ══════════════════════════════════════════════════════════
def charger_livraisons(table, count_var=None):
    for row in table.get_children():
        table.delete(row)

    livraisons = session.query(Livraison).order_by(Livraison.id.desc()).all()

    for i, l in enumerate(livraisons):
        client  = session.query(Client).filter_by(id=l.client_id).first()
        produit = session.query(Produit).filter_by(id=l.produit_id).first()

        statut_raw = l.statut or ""
        icone      = STATUT_ICONE.get(statut_raw, "")
        tags       = ("pair" if i % 2 == 0 else "impair",)
        if statut_raw in STATUT_ICONE:
            tags = tags + (statut_raw,)

        table.insert("", "end", tags=tags, values=(
            l.id,
            l.numero_bl or "-",
            client.nom  if client  else "N/A",
            l.adresse   or "-",
            produit.nom if produit else "N/A",
            f"{float(l.prix):,.2f}" if l.prix else "0.00",
            l.quantite,
            f"{icone} {statut_raw}".strip() if icone else statut_raw,
        ))

    if count_var is not None:
        count_var.set(f"{len(livraisons)} livraison(s)")


# ══════════════════════════════════════════════════════════
#   RECHERCHER
# ══════════════════════════════════════════════════════════
def rechercher(texte, table):
    for row in table.get_children():
        table.delete(row)

    texte_low  = texte.lower()
    livraisons = session.query(Livraison).all()

    for i, l in enumerate(livraisons):
        client  = session.query(Client).filter_by(id=l.client_id).first()
        produit = session.query(Produit).filter_by(id=l.produit_id).first()
        nom_c   = client.nom  if client  else ""
        nom_p   = produit.nom if produit else ""
        adr     = l.adresse   or ""

        if (texte_low in nom_c.lower()
                or texte_low in nom_p.lower()
                or texte_low in adr.lower()
                or texte_low in (l.numero_bl or "").lower()
                or texte_low in (l.statut    or "").lower()):

            statut_raw = l.statut or ""
            icone      = STATUT_ICONE.get(statut_raw, "")
            tags       = ("pair" if i % 2 == 0 else "impair",)

            table.insert("", "end", tags=tags, values=(
                l.id,
                l.numero_bl or "-",
                nom_c or "N/A",
                adr   or "-",
                nom_p or "N/A",
                f"{float(l.prix):,.2f}" if l.prix else "0.00",
                l.quantite,
                f"{icone} {statut_raw}".strip() if icone else statut_raw,
            ))


# ══════════════════════════════════════════════════════════
#   AJOUTER LIVRAISON
# ══════════════════════════════════════════════════════════
def ajouter_livraison(table):
    clients  = session.query(Client).all()
    produits = session.query(Produit).all()

    win, body, t = _fenetre_modale("➕  Nouvelle Livraison", 480, 620)

    _lbl(body, "N Bon de Livraison", t)
    e_num = _entry(body, t)

    _lbl(body, "Client", t)
    cb_client = _combo(body, t, [c.nom for c in clients])

    _lbl(body, "Adresse de livraison", t)
    e_adr = _entry(body, t)

    _lbl(body, "Produit", t)
    cb_produit = _combo(body, t, [p.nom for p in produits])

    _lbl(body, "Quantite", t)
    e_qte = _entry(body, t)

    _lbl(body, "Statut", t)
    cb_statut = _combo(body, t, STATUTS, default="En attente")

    # Apercu prix
    lbl_prev = tk.Label(body, text="", bg=t["bg"],
                        fg="#FF8F00", font=("Arial", 10, "italic"))
    lbl_prev.pack(fill="x", padx=24, pady=(6, 0))

    def maj_prev(*_):
        try:
            p = session.query(Produit).filter_by(
                nom=cb_produit.var.get()).first()
            if p and e_qte.get():
                qte   = int(e_qte.get())
                total = round(float(p.prix_ttc) * qte, 2)
                lbl_prev.config(
                    text=f"Prix TTC : {p.prix_ttc:,.2f}  |  Total : {total:,.2f} MAD")
        except Exception:
            pass

    cb_produit.var.trace("w", maj_prev)
    e_qte.bind("<KeyRelease>", maj_prev)

    def sauvegarder():
        try:
            client  = session.query(Client).filter_by(
                nom=cb_client.var.get()).first()
            produit = session.query(Produit).filter_by(
                nom=cb_produit.var.get()).first()
            if not client:
                messagebox.showerror("Erreur", "Selectionnez un client.", parent=win)
                return
            if not produit:
                messagebox.showerror("Erreur", "Selectionnez un produit.", parent=win)
                return

            qte = int(e_qte.get())

            liv = Livraison(
                numero_bl=e_num.get(),
                client_id=client.id,
                devis_id=None,
                adresse=e_adr.get(),
                produit_id=produit.id,
                prix=produit.prix_ttc,
                quantite=qte,
                statut=cb_statut.var.get(),
            )
            session.add(liv)
            session.commit()
            charger_livraisons(table)
            win.destroy()
            messagebox.showinfo("Succes", "Livraison ajoutee avec succes !")
        except Exception as ex:
            messagebox.showerror("Erreur", f"Erreur : {ex}", parent=win)

    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", padx=24, pady=(12, 4))
    _btn_ok(btn_frame, "💾  Sauvegarder", sauvegarder).pack(
        side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   MODIFIER LIVRAISON  (toutes les infos)
# ══════════════════════════════════════════════════════════
def modifier_livraison(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une livraison a modifier !")
        return

    valeurs   = table.item(sel[0])["values"]
    livraison = session.query(Livraison).filter_by(id=valeurs[0]).first()
    if not livraison:
        return

    clients  = session.query(Client).all()
    produits = session.query(Produit).all()

    client_actuel  = session.query(Client).filter_by(id=livraison.client_id).first()
    produit_actuel = session.query(Produit).filter_by(id=livraison.produit_id).first()

    win, body, t = _fenetre_modale("✏️  Modifier la Livraison", 480, 680)

    # ── Section : Informations generales ─────
    tk.Label(body, text="INFORMATIONS GENERALES",
             bg=t["bg"], fg=PALETTE_ORG,
             font=("Arial", 9, "bold")).pack(
             fill="x", padx=24, pady=(6, 0))
    _sep(body, t)

    _lbl(body, "N Bon de Livraison", t)
    e_num = _entry(body, t, default=livraison.numero_bl or "")

    _lbl(body, "Client", t)
    cb_client = _combo(body, t,
                       [c.nom for c in clients],
                       default=client_actuel.nom if client_actuel else "")

    _lbl(body, "Adresse de livraison", t)
    e_adr = _entry(body, t, default=livraison.adresse or "")

    # ── Section : Produit & Quantite ─────────
    tk.Label(body, text="PRODUIT & QUANTITE",
             bg=t["bg"], fg=PALETTE_ORG,
             font=("Arial", 9, "bold")).pack(
             fill="x", padx=24, pady=(14, 0))
    _sep(body, t)

    _lbl(body, "Produit", t)
    cb_produit = _combo(body, t,
                        [p.nom for p in produits],
                        default=produit_actuel.nom if produit_actuel else "")

    _lbl(body, "Quantite", t)
    e_qte = _entry(body, t, default=str(livraison.quantite or 1))

    # Apercu prix
    lbl_prev = tk.Label(body, text="", bg=t["bg"],
                        fg="#FF8F00", font=("Arial", 10, "italic"))
    lbl_prev.pack(fill="x", padx=24, pady=(4, 0))

    def maj_prev(*_):
        try:
            p = session.query(Produit).filter_by(
                nom=cb_produit.var.get()).first()
            if p and e_qte.get():
                qte   = int(e_qte.get())
                total = round(float(p.prix_ttc) * qte, 2)
                lbl_prev.config(
                    text=f"Prix TTC : {p.prix_ttc:,.2f}  |  Total : {total:,.2f} MAD")
        except Exception:
            pass

    cb_produit.var.trace("w", maj_prev)
    e_qte.bind("<KeyRelease>", maj_prev)
    maj_prev()

    # ── Section : Statut ─────────────────────
    tk.Label(body, text="STATUT",
             bg=t["bg"], fg=PALETTE_ORG,
             font=("Arial", 9, "bold")).pack(
             fill="x", padx=24, pady=(14, 0))
    _sep(body, t)

    _lbl(body, "Statut de livraison", t)
    cb_statut = _combo(body, t, STATUTS,
                       default=livraison.statut or "En attente")

    # Indicateur visuel statut
    statut_colors = {
        "En attente": "#FFA000",
        "En cours":   "#42A5F5",
        "Livre":      "#43A047",
        "Annule":     "#E53935",
    }

    frame_statut = tk.Frame(body, bg=t["bg"])
    frame_statut.pack(fill="x", padx=24, pady=(6, 0))

    lbl_statut_color = tk.Label(frame_statut, text="", bg=t["bg"],
                                font=("Arial", 10, "bold"))
    lbl_statut_color.pack(side="left")

    def maj_statut_color(*_):
        s   = cb_statut.var.get()
        col = statut_colors.get(s, "#888888")
        icn = STATUT_ICONE.get(s, "")
        lbl_statut_color.config(text=f"{icn}  {s}", fg=col)

    cb_statut.var.trace("w", maj_statut_color)
    maj_statut_color()

    # ── Boutons ───────────────────────────────
    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", padx=24, pady=(12, 20))

    def sauvegarder():
        try:
            client  = session.query(Client).filter_by(
                nom=cb_client.var.get()).first()
            produit = session.query(Produit).filter_by(
                nom=cb_produit.var.get()).first()
            if not client:
                messagebox.showerror("Erreur", "Selectionnez un client.", parent=win)
                return
            if not produit:
                messagebox.showerror("Erreur", "Selectionnez un produit.", parent=win)
                return

            livraison.numero_bl  = e_num.get()
            livraison.client_id  = client.id
            livraison.adresse    = e_adr.get()
            livraison.produit_id = produit.id
            livraison.prix       = produit.prix_ttc
            livraison.quantite   = int(e_qte.get())
            livraison.statut     = cb_statut.var.get()

            session.commit()
            charger_livraisons(table)
            win.destroy()
            messagebox.showinfo("Succes", "Livraison modifiee avec succes !")
        except Exception as ex:
            messagebox.showerror("Erreur", f"Erreur : {ex}", parent=win)

    _btn_ok(btn_frame, "💾  Sauvegarder", sauvegarder).pack(
        side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


PALETTE_ORG = "#FF8F00"


# ══════════════════════════════════════════════════════════
#   SUPPRIMER LIVRAISON
# ══════════════════════════════════════════════════════════
def supprimer_livraison(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez une livraison a supprimer !")
        return

    valeurs   = table.item(sel[0])["values"]
    livraison = session.query(Livraison).filter_by(id=valeurs[0]).first()
    if not livraison:
        return

    client = session.query(Client).filter_by(id=livraison.client_id).first()
    nom_c  = client.nom if client else "N/A"

    if messagebox.askyesno(
        "Confirmation",
        f"Supprimer la livraison {livraison.numero_bl or livraison.id} ?\n"
        f"Client  : {nom_c}\n"
        f"Adresse : {livraison.adresse or '-'}\n\n"
        "Cette action est irreversible.",
    ):
        session.delete(livraison)
        session.commit()
        charger_livraisons(table)
        messagebox.showinfo("Succes", "Livraison supprimee avec succes !")


# ══════════════════════════════════════════════════════════
#   DETAIL LIVRAISON  (double-clic)
# ══════════════════════════════════════════════════════════
def detail_livraison(table):
    sel = table.selection()
    if not sel:
        return

    valeurs   = table.item(sel[0])["values"]
    livraison = session.query(Livraison).filter_by(id=valeurs[0]).first()
    if not livraison:
        return

    client  = session.query(Client).filter_by(id=livraison.client_id).first()
    produit = session.query(Produit).filter_by(id=livraison.produit_id).first()

    t = get_theme()
    win, body, t = _fenetre_modale(
        f"Detail — {livraison.numero_bl or livraison.id}", 440, 460)

    info_bg = "#1A1A35" if t["bg"] in ("#1A1A2E", "#0F0F1A") else "#FFF0E5"
    frame   = tk.Frame(body, bg=info_bg)
    frame.pack(fill="both", expand=True, padx=4)

    statut_colors = {
        "En attente": "#FFA000", "En cours": "#42A5F5",
        "Livre": "#43A047", "Annule": "#E53935",
    }

    champs = [
        ("N BL",      livraison.numero_bl or "-"),
        ("Client",    client.nom  if client  else "N/A"),
        ("Adresse",   livraison.adresse or "-"),
        ("Produit",   produit.nom if produit else "N/A"),
        ("Prix TTC",  f"{float(livraison.prix):,.2f} MAD" if livraison.prix else "-"),
        ("Quantite",  str(livraison.quantite)),
        ("Statut",    livraison.statut or "-"),
    ]

    for label, val in champs:
        row = tk.Frame(frame, bg=info_bg)
        row.pack(fill="x", padx=16, pady=5)
        tk.Label(row, text=f"{label} :", bg=info_bg,
                 fg="#FF8F00", font=("Arial", 11, "bold"),
                 width=10, anchor="w").pack(side="left")

        fg_val = statut_colors.get(val, t["text"]) if label == "Statut" else t["text"]
        fw_val = "bold" if label == "Statut" else "normal"
        tk.Label(row, text=val, bg=info_bg,
                 fg=fg_val, font=("Arial", 11, fw_val)).pack(side="left")

        tk.Frame(frame, bg="#333355" if t["bg"] in ("#1A1A2E",)
                 else "#E8D8C8", height=1).pack(fill="x", padx=16)

    btn_row = tk.Frame(body, bg=t["bg"])
    btn_row.pack(fill="x", padx=24, pady=(12, 4))

    _btn_ok(btn_row, "✏️  Modifier",
            lambda: [win.destroy(), modifier_livraison(table)],
            "#FF8F00", "#E65100").pack(side="left", padx=(0, 6))
    _btn_ok(btn_row, "🖨️  Imprimer",
            lambda: _imprimer_une(livraison, client, produit),
            "#2E7D32", "#1B5E20").pack(side="left", padx=(0, 6))
    tk.Button(btn_row, text="Fermer", command=win.destroy,
              bg="#E65100", fg="white",
              font=("Arial", 12, "bold"),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   UTILITAIRE : collecter les données
# ══════════════════════════════════════════════════════════
def _get_livraisons_data(selection_only=False, table=None):
    rows = []

    if selection_only and table:
        sel = table.selection()
        ids = [table.item(s)["values"][0] for s in sel]
        livs = [session.query(Livraison).filter_by(id=i).first() for i in ids]
        livs = [l for l in livs if l]
    else:
        livs = session.query(Livraison).order_by(Livraison.id.desc()).all()

    for l in livs:
        client  = session.query(Client).filter_by(id=l.client_id).first()
        produit = session.query(Produit).filter_by(id=l.produit_id).first()
        rows.append({
            "id":        l.id,
            "numero_bl": l.numero_bl or "-",
            "client":    client.nom  if client  else "N/A",
            "adresse":   l.adresse   or "-",
            "produit":   produit.nom if produit else "N/A",
            "prix":      float(l.prix) if l.prix else 0.0,
            "quantite":  l.quantite,
            "statut":    l.statut or "-",
        })
    return rows


# ══════════════════════════════════════════════════════════
#   EXPORTER LIVRAISONS
# ══════════════════════════════════════════════════════════
def exporter_livraisons(table, format_):
    rows = _get_livraisons_data()
    if not rows:
        messagebox.showwarning("Vide", "Aucune donnee a exporter.")
        return

    ext_map = {
        "excel": ("Fichier Excel", "*.xlsx", ".xlsx"),
        "word":  ("Document Word",  "*.docx", ".docx"),
        "pdf":   ("Fichier PDF",    "*.pdf",  ".pdf"),
        "csv":   ("Fichier CSV",    "*.csv",  ".csv"),
        "json":  ("Fichier JSON",   "*.json", ".json"),
    }
    label, pattern, ext = ext_map[format_]

    chemin = filedialog.asksaveasfilename(
        title=f"Exporter en {label}",
        defaultextension=ext,
        filetypes=[(label, pattern), ("Tous", "*.*")],
        initialfile=f"livraisons_{date.today()}{ext}",
    )
    if not chemin:
        return

    try:
        if format_ == "excel":  _export_excel(chemin, rows)
        elif format_ == "word": _export_word(chemin, rows)
        elif format_ == "pdf":  _export_pdf(chemin, rows)
        elif format_ == "csv":  _export_csv(chemin, rows)
        elif format_ == "json": _export_json(chemin, rows)

        if messagebox.askyesno("Export reussi",
                               f"Fichier cree :\n{chemin}\n\nOuvrir maintenant ?"):
            os.startfile(chemin)
    except Exception as ex:
        messagebox.showerror("Erreur export", f"Erreur :\n{ex}")


# ── Export Excel ──────────────────────────────────────────
def _export_excel(chemin, rows):
    if not EXCEL_OK:
        messagebox.showerror("Manquant", "pip install openpyxl"); return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Livraisons"

    ws.merge_cells("A1:H1")
    ws["A1"] = f"Bons de Livraison — Exporte le {date.today()}"
    ws["A1"].font      = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill      = PatternFill("solid", fgColor="E65100")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    for ci, h in enumerate(EN_TETES, 1):
        cell = ws.cell(row=2, column=ci, value=h)
        cell.font      = Font(bold=True, color="FFFFFF", size=10)
        cell.fill      = PatternFill("solid", fgColor="BF360C")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 24

    bd   = Side(style="thin", color="E8D8C8")
    bord = Border(left=bd, right=bd, top=bd, bottom=bd)
    fp   = PatternFill("solid", fgColor="FFF8F0")
    fi   = PatternFill("solid", fgColor="FFFFFF")

    statut_col = {
        "En attente": "E65100", "En cours": "1565C0",
        "Livre": "2E7D32", "Annule": "C62828",
    }

    for ri, row in enumerate(rows, 3):
        fill = fp if ri % 2 == 0 else fi
        for ci, champ in enumerate(CHAMPS, 1):
            val  = row[champ]
            cell = ws.cell(row=ri, column=ci, value=val)
            cell.fill      = fill
            cell.border    = bord
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.font      = Font(size=9)
            if champ == "statut":
                fg = statut_col.get(str(val), "212121")
                cell.font = Font(size=9, bold=True, color=fg)

    larg = [6, 12, 18, 22, 18, 10, 6, 14]
    for i, w in enumerate(larg, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    wb.save(chemin)


# ── Export Word ───────────────────────────────────────────
def _export_word(chemin, rows):
    if not WORD_OK:
        messagebox.showerror("Manquant", "pip install python-docx"); return

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    titre = doc.add_heading("Bons de Livraison", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.runs[0].font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    titre.runs[0].font.size      = Pt(18)

    st = doc.add_paragraph(f"Exporte le {date.today()}  |  {len(rows)} livraison(s)")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x33)
    st.runs[0].font.size      = Pt(10)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=len(EN_TETES))
    tbl.style = "Table Grid"

    for i, h in enumerate(EN_TETES):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E65100")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    statut_col = {
        "En attente": ("FFF3E0", "E65100"),
        "En cours":   ("E3F2FD", "1565C0"),
        "Livre":      ("E8F5E9", "2E7D32"),
        "Annule":     ("FFEBEE", "C62828"),
    }

    for row in rows:
        tr = tbl.add_row()
        for i, champ in enumerate(CHAMPS):
            val  = row[champ]
            cell = tr.cells[i]
            cell.text = f"{float(val):,.2f}" if champ == "prix" else str(val)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)
            if champ == "statut":
                bg, fg = statut_col.get(str(val), ("FFFFFF", "212121"))
                run.font.bold      = True
                run.font.color.rgb = RGBColor(
                    int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc_pr = cell._tc.get_or_add_tcPr()
                shd   = OxmlElement("w:shd")
                shd.set(qn("w:fill"), bg)
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tc_pr.append(shd)

    doc.save(chemin)


# ── Export PDF ────────────────────────────────────────────
def _export_pdf(chemin, rows):
    if not PDF_OK:
        messagebox.showerror("Manquant", "pip install reportlab"); return

    doc    = SimpleDocTemplate(chemin, pagesize=A4,
                               leftMargin=1.5*cm, rightMargin=1.5*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []

    story.append(Paragraph("Bons de Livraison",
        ParagraphStyle("T", parent=styles["Title"],
                       fontSize=18, textColor=colors.HexColor("#E65100"),
                       spaceAfter=4)))
    story.append(Paragraph(
        f"Exporte le {date.today()}  —  {len(rows)} livraison(s)",
        ParagraphStyle("S", parent=styles["Normal"],
                       fontSize=9, textColor=colors.HexColor("#996633"),
                       spaceAfter=14)))

    data = [EN_TETES]
    for row in rows:
        data.append([
            str(row["id"]), str(row["numero_bl"]),
            str(row["client"]), str(row["adresse"]),
            str(row["produit"]), f"{row['prix']:,.2f}",
            str(row["quantite"]), str(row["statut"]),
        ])

    col_w = [1*cm, 2*cm, 3*cm, 3.5*cm, 3*cm, 2*cm, 1.2*cm, 2.3*cm]
    tbl   = Table(data, colWidths=col_w, repeatRows=1)

    ts = TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#E65100")),
        ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 8),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING",    (0, 0), (-1, 0), 8),
        ("FONTSIZE",      (0, 1), (-1, -1), 7.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.HexColor("#FFF8F0"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E8D8C8")),
        ("TOPPADDING",    (0, 1), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 5),
    ])
    statut_col = {
        "En attente": colors.HexColor("#E65100"),
        "En cours":   colors.HexColor("#1565C0"),
        "Livre":      colors.HexColor("#2E7D32"),
        "Annule":     colors.HexColor("#C62828"),
    }
    for ri, row in enumerate(rows, 1):
        col = statut_col.get(row["statut"])
        if col:
            ts.add("TEXTCOLOR", (7, ri), (7, ri), col)
            ts.add("FONTNAME",  (7, ri), (7, ri), "Helvetica-Bold")
    tbl.setStyle(ts)
    story.append(tbl)
    doc.build(story)


# ── Export CSV ────────────────────────────────────────────
def _export_csv(chemin, rows):
    with open(chemin, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=CHAMPS)
        writer.writeheader()
        writer.writerows(rows)


# ── Export JSON ───────────────────────────────────────────
def _export_json(chemin, rows):
    with open(chemin, "w", encoding="utf-8") as f:
        json.dump({
            "export_date": str(date.today()),
            "total_records": len(rows),
            "livraisons": rows,
        }, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════
#   IMPORTER LIVRAISONS
# ══════════════════════════════════════════════════════════
def importer_fichier(table, format_):
    ext_map = {
        "excel": [("Fichier Excel", "*.xlsx"), ("Tous", "*.*")],
        "csv":   [("Fichier CSV",   "*.csv"),  ("Tous", "*.*")],
        "json":  [("Fichier JSON",  "*.json"), ("Tous", "*.*")],
    }
    chemin = filedialog.askopenfilename(
        title=f"Importer depuis {format_.upper()}",
        filetypes=ext_map.get(format_, [("Tous", "*.*")]),
    )
    if not chemin:
        return

    try:
        if format_ == "excel":  lignes = _import_excel(chemin)
        elif format_ == "csv":  lignes = _import_csv(chemin)
        elif format_ == "json": lignes = _import_json(chemin)
        else:                   lignes = []

        if not lignes:
            messagebox.showwarning("Vide", "Aucune ligne trouvee dans le fichier.")
            return

        _fenetre_import_preview(table, lignes, chemin)
    except Exception as ex:
        messagebox.showerror("Erreur import", f"Erreur :\n{ex}")


def _import_excel(chemin):
    if not EXCEL_OK:
        messagebox.showerror("Manquant", "pip install openpyxl"); return []
    wb   = openpyxl.load_workbook(chemin, data_only=True)
    ws   = wb.active
    rows = list(ws.iter_rows(values_only=True))
    start = 2 if (rows and rows[0] and
                  str(rows[0][0] or "").lower() in ("bons de livraison", "titre")) else 1
    return [[str(c) if c is not None else "" for c in row]
            for row in rows[start:] if any(c is not None for c in row)]


def _import_csv(chemin):
    lignes = []
    with open(chemin, encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        next(reader, None)
        for row in reader:
            if row:
                lignes.append(row)
    return lignes


def _import_json(chemin):
    with open(chemin, encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict):
        data = data.get("livraisons", [])
    return [[str(item.get(k, "")) for k in CHAMPS]
            for item in data if isinstance(item, dict)]


def _fenetre_import_preview(table, lignes, chemin):
    t   = get_theme()
    win = tk.Toplevel()
    win.title(f"Apercu import — {len(lignes)} ligne(s)")
    win.geometry("860x500")
    win.configure(bg=t["bg"])
    win.grab_set()

    hdr = tk.Frame(win, bg="#1565C0", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=f"📥  Apercu Import — {os.path.basename(chemin)}",
             bg="#1565C0", fg="white",
             font=("Arial", 13, "bold")).pack(pady=12)

    frame_prev = tk.Frame(win, bg=t["bg"])
    frame_prev.pack(fill="both", expand=True, padx=16, pady=10)

    prev_tree = ttk.Treeview(frame_prev, columns=EN_TETES,
                             show="headings", height=14,
                             style="LV.Treeview")
    for col in EN_TETES:
        prev_tree.heading(col, text=col)
        prev_tree.column(col, width=90, anchor="center")

    for i, row in enumerate(lignes[:50]):
        padded = list(row) + [""] * max(0, len(EN_TETES) - len(row))
        prev_tree.insert("", "end",
                         tags=("pair" if i % 2 == 0 else "impair",),
                         values=padded[:len(EN_TETES)])

    sb = ttk.Scrollbar(frame_prev, orient="vertical", command=prev_tree.yview)
    prev_tree.configure(yscroll=sb.set)
    sb.pack(side="right", fill="y")
    prev_tree.pack(fill="both", expand=True)

    if len(lignes) > 50:
        tk.Label(win, text=f"... et {len(lignes)-50} autres lignes.",
                 bg=t["bg"], fg="#FF8F00",
                 font=("Arial", 9, "italic")).pack()

    tk.Label(win,
             text=f"{len(lignes)} lignes detectees. Seules les lignes avec client et produit existants seront importees.",
             bg=t["bg"], fg=t["text"],
             font=("Arial", 9)).pack(pady=(0, 6))

    btn_frame = tk.Frame(win, bg=t["bg"])
    btn_frame.pack(pady=8)

    def confirmer():
        nb_ok, nb_err = _inserer_lignes(lignes)
        charger_livraisons(table)
        win.destroy()
        messagebox.showinfo("Import termine",
            f"Inseres : {nb_ok}\nIgnores : {nb_err}")

    _btn_ok(btn_frame, f"✅  Importer {len(lignes)} ligne(s)",
            confirmer, "#1565C0", "#0D47A1").pack(side="left", padx=8)
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2",
              padx=16, pady=9).pack(side="left", padx=8)


def _inserer_lignes(lignes):
    nb_ok = nb_err = 0
    for row in lignes:
        try:
            padded     = list(row) + [""] * max(0, 8 - len(row))
            nom_client = str(padded[2]).strip()
            nom_produit= str(padded[4]).strip()

            client  = session.query(Client).filter_by(nom=nom_client).first()
            produit = session.query(Produit).filter_by(nom=nom_produit).first()
            if not client or not produit:
                nb_err += 1; continue

            num_bl = str(padded[1]).strip() or None
            if num_bl and session.query(Livraison).filter_by(
                    numero_bl=num_bl).first():
                nb_err += 1; continue

            liv = Livraison(
                numero_bl=num_bl,
                client_id=client.id,
                devis_id=None,
                adresse=str(padded[3]).strip() or "-",
                produit_id=produit.id,
                prix=float(str(padded[5]).replace(",", "") or 0),
                quantite=int(float(str(padded[6]).replace(",", "") or 1)),
                statut=str(padded[7]).strip() or "En attente",
            )
            session.add(liv)
            nb_ok += 1
        except Exception:
            nb_err += 1
    session.commit()
    return nb_ok, nb_err


# ══════════════════════════════════════════════════════════
#   IMPRIMER
# ══════════════════════════════════════════════════════════
def imprimer_livraison(table):
    sel = table.selection()

    if sel:
        choix = messagebox.askyesnocancel(
            "Imprimer",
            "Oui = selection  |  Non = toutes  |  Annuler = annuler",
        )
        if choix is None:
            return
        rows = _get_livraisons_data(selection_only=choix, table=table)
    else:
        rows = _get_livraisons_data()

    if not rows:
        messagebox.showwarning("Vide", "Aucune donnee a imprimer.")
        return

    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()

    try:
        if PDF_OK:
            _export_pdf(tmp.name, rows)
            os.startfile(tmp.name, "print")
            messagebox.showinfo("Impression",
                f"Envoi imprimante : {len(rows)} livraison(s).")
        elif EXCEL_OK:
            tmp2 = tmp.name.replace(".pdf", ".xlsx")
            _export_excel(tmp2, rows)
            os.startfile(tmp2)
            messagebox.showinfo("Impression",
                "reportlab absent — fichier Excel ouvert pour impression manuelle.\n"
                "pip install reportlab pour impression directe.")
        else:
            tmp2 = tmp.name.replace(".pdf", ".csv")
            _export_csv(tmp2, rows)
            os.startfile(tmp2)
    except Exception as ex:
        messagebox.showerror("Erreur impression", f"Erreur : {ex}")


def _imprimer_une(livraison, client, produit):
    rows = [{
        "id":        livraison.id,
        "numero_bl": livraison.numero_bl or "-",
        "client":    client.nom  if client  else "N/A",
        "adresse":   livraison.adresse or "-",
        "produit":   produit.nom if produit else "N/A",
        "prix":      float(livraison.prix) if livraison.prix else 0.0,
        "quantite":  livraison.quantite,
        "statut":    livraison.statut or "-",
    }]
    import tempfile
    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    tmp.close()
    try:
        if PDF_OK:
            _export_pdf(tmp.name, rows)
            os.startfile(tmp.name, "print")
        elif EXCEL_OK:
            tmp2 = tmp.name.replace(".pdf", ".xlsx")
            _export_excel(tmp2, rows)
            os.startfile(tmp2)
    except Exception as ex:
        messagebox.showerror("Erreur", f"Erreur impression : {ex}")