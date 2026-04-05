import customtkinter as ctk
import tkinter as tk
from tkinter import ttk, filedialog, messagebox
from database.db import session
from database.models import Client
from datetime import date
from utils.theme import get_theme
import os

# ══════════════════════════════════════════════════════════
#   IMPORTS OPTIONNELS
# ══════════════════════════════════════════════════════════
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_OK = True
except ImportError:
    EXCEL_OK = False

try:
    import pandas as pd
    PANDAS_OK = True
except ImportError:
    PANDAS_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
    from reportlab.lib.units import cm
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    WORD_OK = True
except ImportError:
    WORD_OK = False

import csv
import json


# ══════════════════════════════════════════════════════════
#   CONSTANTES
# ══════════════════════════════════════════════════════════
EN_TETES = ["ID", "Nom", "Email", "Telephone", "Adresse", "Ville", "Date creation"]
CHAMPS   = ["id", "nom", "email", "telephone", "adresse", "ville", "date_creation"]


# ══════════════════════════════════════════════════════════
#   STYLE TREEVIEW
# ══════════════════════════════════════════════════════════
def _style_treeview(is_dark):
    style = ttk.Style()
    style.theme_use("clam")

    bg_tree = "#12122A" if is_dark else "#FFFFFF"
    bg_head = "#E65100"
    bg_sel  = "#FF6D00" if is_dark else "#FFB74D"
    fg_text = "#E8E0FF" if is_dark else "#1A1A1A"
    bg_row1 = "#1A1A35" if is_dark else "#FFFFFF"
    bg_row2 = "#12122A" if is_dark else "#FFF8F0"

    style.configure("CL.Treeview",
                    background=bg_tree, foreground=fg_text,
                    rowheight=32, fieldbackground=bg_tree,
                    borderwidth=0, font=("Arial", 10))
    style.configure("CL.Treeview.Heading",
                    background=bg_head, foreground="white",
                    font=("Arial", 10, "bold"),
                    relief="flat", borderwidth=0)
    style.map("CL.Treeview",
              background=[("selected", bg_sel)],
              foreground=[("selected", "#FFFFFF")])
    style.map("CL.Treeview.Heading",
              background=[("active", "#BF360C")])
    return bg_row1, bg_row2


# ══════════════════════════════════════════════════════════
#   HELPERS UI
# ══════════════════════════════════════════════════════════
def _fenetre_modale(titre, largeur=460, hauteur=520):
    t   = get_theme()
    win = tk.Toplevel()
    win.title(titre)
    win.geometry(f"{largeur}x{hauteur}")
    win.configure(bg=t["bg"])
    win.resizable(False, False)
    win.grab_set()

    hdr = tk.Frame(win, bg="#E65100", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=titre, bg="#E65100", fg="white",
             font=("Arial", 14, "bold")).pack(pady=12)

    body = tk.Frame(win, bg=t["bg"])
    body.pack(fill="both", expand=True, padx=24, pady=16)
    return win, body, t


def _lbl(body, text, t):
    tk.Label(body, text=text, bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(
             fill="x", pady=(10, 2))


def _entry(body, t, default=""):
    e = tk.Entry(body, font=("Arial", 11),
                 bg=t["card"], fg=t["text"],
                 insertbackground=t["text"],
                 relief="flat", bd=6)
    e.pack(fill="x", ipady=4)
    if default:
        e.insert(0, str(default))
    return e


def _btn_ok(parent, text, cmd, color="#E65100", hover="#BF360C"):
    return tk.Button(
        parent, text=text, command=cmd,
        bg=color, fg="white",
        font=("Arial", 12, "bold"),
        relief="flat", cursor="hand2",
        padx=16, pady=9,
        activebackground=hover,
        activeforeground="white",
    )


def _sep(body, t):
    tk.Frame(body,
             bg="#2A2A4A" if t["bg"] in ("#1A1A2E", "#0F0F1A") else "#E8D8C8",
             height=1).pack(fill="x", pady=(10, 0))


# ══════════════════════════════════════════════════════════
#   TRI COLONNES
# ══════════════════════════════════════════════════════════
_tri_etat = {}

def _trier(table, col):
    rows = [(table.set(r, col), r) for r in table.get_children("")]
    rev  = _tri_etat.get(col, False)
    try:
        rows.sort(key=lambda x: float(x[0].replace(",", "")), reverse=rev)
    except ValueError:
        rows.sort(key=lambda x: x[0].lower(), reverse=rev)
    for i, (_, r) in enumerate(rows):
        table.move(r, "", i)
    _tri_etat[col] = not rev


# ══════════════════════════════════════════════════════════
#   MENU DEROULANT GENERIQUE
# ══════════════════════════════════════════════════════════
def _show_dropdown(btn_widget, items, t):
    is_dark  = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
    bg_menu  = "#1A1A35" if is_dark else "#FFFFFF"
    bg_hover = "#252545" if is_dark else "#FFF0E5"
    border_c = "#2A2A4A" if is_dark else "#E8D8C8"

    menu = tk.Toplevel()
    menu.overrideredirect(True)
    menu.configure(bg=border_c)
    menu.attributes("-topmost", True)

    btn_widget.update_idletasks()
    x = btn_widget.winfo_rootx()
    y = btn_widget.winfo_rooty() + btn_widget.winfo_height() + 2
    menu.geometry(f"210x{len(items) * 40 + 16}+{x}+{y}")

    inner = tk.Frame(menu, bg=bg_menu, bd=0)
    inner.pack(fill="both", expand=True, padx=1, pady=1)

    tk.Label(inner, text="Choisir le format :",
             bg=bg_menu, fg="#FF8F00",
             font=("Arial", 8, "bold"),
             anchor="w").pack(fill="x", padx=10, pady=(6, 2))

    def _fermer():
        try:
            menu.destroy()
        except Exception:
            pass

    for label, couleur, cmd_fn in items:
        f   = tk.Frame(inner, bg=bg_menu, cursor="hand2")
        f.pack(fill="x", padx=4, pady=1)
        lbl = tk.Label(f, text=label, bg=bg_menu, fg=couleur,
                       font=("Arial", 10, "bold"),
                       anchor="w", padx=12, pady=8)
        lbl.pack(fill="x")

        def _enter(e, fr=f, lb=lbl):
            fr.configure(bg=bg_hover); lb.configure(bg=bg_hover)
        def _leave(e, fr=f, lb=lbl):
            fr.configure(bg=bg_menu);  lb.configure(bg=bg_menu)
        def _click(e, fn=cmd_fn):
            _fermer(); fn()

        for w in (f, lbl):
            w.bind("<Enter>",    _enter)
            w.bind("<Leave>",    _leave)
            w.bind("<Button-1>", _click)

    menu.bind("<FocusOut>", lambda e: _fermer())
    menu.focus_set()


def _show_dropdown_import(btn_widget, table, t):
    items = [
        ("📊  Excel (.xlsx)", "#1565C0", lambda: _importer(table, "excel")),
        ("📄  CSV (.csv)",    "#00838F", lambda: _importer(table, "csv")),
        ("📋  JSON (.json)",  "#6A1B9A", lambda: _importer(table, "json")),
    ]
    _show_dropdown(btn_widget, items, t)


def _show_dropdown_export(btn_widget, table, t):
    items = [
        ("📊  Excel (.xlsx)",  "#1565C0", lambda: _exporter(table, "excel")),
        ("📝  Word (.docx)",   "#1976D2", lambda: _exporter(table, "word")),
        ("📕  PDF (.pdf)",     "#C62828", lambda: _exporter(table, "pdf")),
        ("📄  CSV (.csv)",     "#00838F", lambda: _exporter(table, "csv")),
        ("📋  JSON (.json)",   "#6A1B9A", lambda: _exporter(table, "json")),
    ]
    _show_dropdown(btn_widget, items, t)


# ══════════════════════════════════════════════════════════
#   AFFICHER CLIENTS
# ══════════════════════════════════════════════════════════
def afficher_clients(parent):
    for widget in parent.winfo_children():
        widget.destroy()

    t       = get_theme()
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")

    # En-tete
    header = ctk.CTkFrame(parent, fg_color=t["card"], corner_radius=0)
    header.pack(fill="x")
    ctk.CTkLabel(header, text="👥  Gestion des Clients",
                 font=("Arial", 22, "bold"),
                 text_color="#E65100").pack(side="left", padx=24, pady=14)

    # Barre outils
    toolbar = ctk.CTkFrame(parent, fg_color="transparent")
    toolbar.pack(fill="x", padx=20, pady=(10, 4))

    btns = [
        ("➕  Ajouter",    "#E65100", "#BF360C", lambda: ajouter_client(table)),
        ("✏️  Modifier",   "#FF8F00", "#E65100", lambda: modifier_client(table)),
        ("🗑️  Supprimer",  "#C62828", "#8B0000", lambda: supprimer_client(table)),
    ]
    for txt, fg, hover, cmd in btns:
        ctk.CTkButton(toolbar, text=txt, width=128, height=36,
                      fg_color=fg, hover_color=hover,
                      font=("Arial", 11, "bold"),
                      corner_radius=8,
                      command=cmd).pack(side="left", padx=4)

    # Bouton Importer dropdown
    btn_imp = ctk.CTkButton(
        toolbar, text="📥  Importer ▾", width=138, height=36,
        fg_color="#1565C0", hover_color="#0D47A1",
        font=("Arial", 11, "bold"), corner_radius=8,
        command=lambda: _show_dropdown_import(btn_imp, table, t)
    )
    btn_imp.pack(side="left", padx=4)

    # Compteur + Recherche
    count_var = tk.StringVar()
    ctk.CTkLabel(toolbar, textvariable=count_var,
                 font=("Arial", 11),
                 text_color="#FF8F00").pack(side="right", padx=10)

    recherche_var = tk.StringVar()
    ctk.CTkEntry(toolbar,
                 placeholder_text="🔍  Rechercher...",
                 width=240, height=36,
                 textvariable=recherche_var,
                 font=("Arial", 11)).pack(side="right", padx=5)
    recherche_var.trace("w",
        lambda *a: rechercher(recherche_var.get(), table))

    # Tableau
    frame_table = ctk.CTkFrame(parent, corner_radius=12)
    frame_table.pack(fill="both", expand=True, padx=20, pady=(4, 16))

    bg_row1, bg_row2 = _style_treeview(is_dark)

    colonnes = ("ID", "Nom", "Email", "Telephone", "Adresse", "Ville", "Date")

    table = ttk.Treeview(frame_table, columns=colonnes,
                         show="headings", height=22,
                         style="CL.Treeview",
                         selectmode="browse")

    largeurs = {"ID": 45, "Nom": 170, "Email": 190,
                "Telephone": 120, "Adresse": 180,
                "Ville": 120, "Date": 110}

    for col in colonnes:
        table.heading(col, text=col,
                      command=lambda c=col: _trier(table, c))
        table.column(col, width=largeurs.get(col, 120), anchor="center")

    table.tag_configure("pair",   background=bg_row1)
    table.tag_configure("impair", background=bg_row2)

    sb_v = ttk.Scrollbar(frame_table, orient="vertical",   command=table.yview)
    sb_h = ttk.Scrollbar(frame_table, orient="horizontal", command=table.xview)
    table.configure(yscroll=sb_v.set, xscroll=sb_h.set)
    sb_v.pack(side="right",  fill="y")
    sb_h.pack(side="bottom", fill="x")
    table.pack(fill="both", expand=True, padx=2, pady=2)

    table.bind("<Double-1>", lambda e: modifier_client(table))

    charger_clients(table, count_var)
    return table


# ══════════════════════════════════════════════════════════
#   CHARGER CLIENTS
# ══════════════════════════════════════════════════════════
def charger_clients(table, count_var=None):
    for row in table.get_children():
        table.delete(row)

    clients = session.query(Client).order_by(Client.id.desc()).all()

    for i, c in enumerate(clients):
        table.insert("", "end",
                     tags=("pair" if i % 2 == 0 else "impair",),
                     values=(
                         c.id,
                         c.nom        or "-",
                         c.email      or "-",
                         c.telephone  or "-",
                         c.adresse    or "-",
                         c.ville      or "-",
                         str(c.date_creation) if c.date_creation else "-",
                     ))

    if count_var is not None:
        count_var.set(f"{len(clients)} client(s)")


# ══════════════════════════════════════════════════════════
#   RECHERCHER
# ══════════════════════════════════════════════════════════
def rechercher(texte, table):
    for row in table.get_children():
        table.delete(row)

    clients = session.query(Client).filter(
        Client.nom.ilike(f"%{texte}%")
    ).all()

    for i, c in enumerate(clients):
        table.insert("", "end",
                     tags=("pair" if i % 2 == 0 else "impair",),
                     values=(
                         c.id,
                         c.nom       or "-",
                         c.email     or "-",
                         c.telephone or "-",
                         c.adresse   or "-",
                         c.ville     or "-",
                         str(c.date_creation) if c.date_creation else "-",
                     ))


# ══════════════════════════════════════════════════════════
#   AJOUTER CLIENT
# ══════════════════════════════════════════════════════════
def ajouter_client(table):
    win, body, t = _fenetre_modale("➕  Nouveau Client", 460, 530)

    _lbl(body, "Nom *", t)
    e_nom = _entry(body, t)

    _lbl(body, "Email", t)
    e_email = _entry(body, t)

    _lbl(body, "Telephone", t)
    e_tel = _entry(body, t)

    _lbl(body, "Adresse", t)
    e_adr = _entry(body, t)

    _lbl(body, "Ville", t)
    e_ville = _entry(body, t)

    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", pady=(12, 4))

    def sauvegarder():
        if not e_nom.get().strip():
            messagebox.showerror("Erreur", "Le nom est obligatoire.", parent=win)
            return
        c = Client(
            nom=e_nom.get().strip(),
            email=e_email.get().strip(),
            telephone=e_tel.get().strip(),
            adresse=e_adr.get().strip(),
            ville=e_ville.get().strip(),
            date_creation=date.today(),
        )
        session.add(c)
        session.commit()
        charger_clients(table)
        win.destroy()
        messagebox.showinfo("Succes", "Client ajoute avec succes !")

    _btn_ok(btn_frame, "💾  Sauvegarder", sauvegarder).pack(
        side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   MODIFIER CLIENT
# ══════════════════════════════════════════════════════════
def modifier_client(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez un client a modifier !")
        return

    valeurs = table.item(sel[0])["values"]
    client  = session.query(Client).filter_by(id=valeurs[0]).first()
    if not client:
        return

    win, body, t = _fenetre_modale("✏️  Modifier le Client", 460, 530)

    _lbl(body, "Nom *", t)
    e_nom = _entry(body, t, default=client.nom or "")

    _lbl(body, "Email", t)
    e_email = _entry(body, t, default=client.email or "")

    _lbl(body, "Telephone", t)
    e_tel = _entry(body, t, default=client.telephone or "")

    _lbl(body, "Adresse", t)
    e_adr = _entry(body, t, default=client.adresse or "")

    _lbl(body, "Ville", t)
    e_ville = _entry(body, t, default=client.ville or "")

    _sep(body, t)
    btn_frame = tk.Frame(body, bg=t["bg"])
    btn_frame.pack(fill="x", pady=(12, 4))

    def sauvegarder():
        if not e_nom.get().strip():
            messagebox.showerror("Erreur", "Le nom est obligatoire.", parent=win)
            return
        client.nom       = e_nom.get().strip()
        client.email     = e_email.get().strip()
        client.telephone = e_tel.get().strip()
        client.adresse   = e_adr.get().strip()
        client.ville     = e_ville.get().strip()
        session.commit()
        charger_clients(table)
        win.destroy()
        messagebox.showinfo("Succes", "Client modifie avec succes !")

    _btn_ok(btn_frame, "💾  Sauvegarder", sauvegarder).pack(
        side="left", expand=True, fill="x", padx=(0, 6))
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=9).pack(
              side="left", expand=True, fill="x")


# ══════════════════════════════════════════════════════════
#   SUPPRIMER CLIENT
# ══════════════════════════════════════════════════════════
def supprimer_client(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez un client a supprimer !")
        return

    valeurs = table.item(sel[0])["values"]
    client  = session.query(Client).filter_by(id=valeurs[0]).first()
    if not client:
        return

    if messagebox.askyesno(
        "Confirmation",
        f"Supprimer le client : {client.nom} ?\n"
        f"Email : {client.email or '-'}\n"
        f"Ville : {client.ville or '-'}\n\n"
        "Attention : les ventes liees a ce client seront impactees.",
    ):
        session.delete(client)
        session.commit()
        charger_clients(table)
        messagebox.showinfo("Succes", "Client supprime avec succes !")


# ══════════════════════════════════════════════════════════
#   UTILITAIRE : collecte des données
# ══════════════════════════════════════════════════════════
def _get_clients_data():
    clients = session.query(Client).order_by(Client.id.desc()).all()
    rows = []
    for c in clients:
        rows.append({
            "id":            c.id,
            "nom":           c.nom           or "",
            "email":         c.email         or "",
            "telephone":     c.telephone     or "",
            "adresse":       c.adresse       or "",
            "ville":         c.ville         or "",
            "date_creation": str(c.date_creation) if c.date_creation else "",
        })
    return rows


# ══════════════════════════════════════════════════════════
#   EXPORTER CLIENTS
# ══════════════════════════════════════════════════════════
def _exporter(table, format_):
    rows = _get_clients_data()
    if not rows:
        messagebox.showwarning("Vide", "Aucun client a exporter.")
        return

    ext_map = {
        "excel": ("Fichier Excel", "*.xlsx", ".xlsx"),
        "word":  ("Document Word",  "*.docx", ".docx"),
        "pdf":   ("Fichier PDF",    "*.pdf",  ".pdf"),
        "csv":   ("Fichier CSV",    "*.csv",  ".csv"),
        "json":  ("Fichier JSON",   "*.json", ".json"),
    }
    label, pattern, ext = ext_map[format_]

    chemin = filedialog.asksaveasfilename(
        title=f"Exporter en {label}",
        defaultextension=ext,
        filetypes=[(label, pattern), ("Tous", "*.*")],
        initialfile=f"clients_{date.today()}{ext}",
    )
    if not chemin:
        return

    try:
        if format_ == "excel":  _export_excel(chemin, rows)
        elif format_ == "word": _export_word(chemin, rows)
        elif format_ == "pdf":  _export_pdf(chemin, rows)
        elif format_ == "csv":  _export_csv(chemin, rows)
        elif format_ == "json": _export_json(chemin, rows)

        if messagebox.askyesno("Export reussi",
                               f"Fichier cree :\n{chemin}\n\nOuvrir maintenant ?"):
            os.startfile(chemin)
    except Exception as ex:
        messagebox.showerror("Erreur export", f"Erreur :\n{ex}")


def _export_excel(chemin, rows):
    if not EXCEL_OK:
        messagebox.showerror("Manquant", "pip install openpyxl"); return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Clients"

    ws.merge_cells("A1:G1")
    ws["A1"] = f"Liste des Clients — Exporte le {date.today()}"
    ws["A1"].font      = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill      = PatternFill("solid", fgColor="E65100")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    for ci, h in enumerate(EN_TETES, 1):
        cell = ws.cell(row=2, column=ci, value=h)
        cell.font      = Font(bold=True, color="FFFFFF", size=10)
        cell.fill      = PatternFill("solid", fgColor="BF360C")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 24

    bd   = Side(style="thin", color="E8D8C8")
    bord = Border(left=bd, right=bd, top=bd, bottom=bd)
    fp   = PatternFill("solid", fgColor="FFF8F0")
    fi   = PatternFill("solid", fgColor="FFFFFF")

    for ri, row in enumerate(rows, 3):
        fill = fp if ri % 2 == 0 else fi
        for ci, champ in enumerate(CHAMPS, 1):
            cell = ws.cell(row=ri, column=ci, value=row[champ])
            cell.fill      = fill
            cell.border    = bord
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.font      = Font(size=9)

    larg = [6, 20, 24, 14, 22, 14, 14]
    for i, w in enumerate(larg, 1):
        ws.column_dimensions[openpyxl.utils.get_column_letter(i)].width = w

    wb.save(chemin)


def _export_word(chemin, rows):
    if not WORD_OK:
        messagebox.showerror("Manquant", "pip install python-docx"); return

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(1.5); section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2);  section.right_margin  = Cm(2)

    titre = doc.add_heading("Liste des Clients", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titre.runs[0].font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    titre.runs[0].font.size      = Pt(18)

    st = doc.add_paragraph(f"Exporte le {date.today()}  |  {len(rows)} client(s)")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x33)
    st.runs[0].font.size      = Pt(10)
    doc.add_paragraph()

    tbl = doc.add_table(rows=1, cols=len(EN_TETES))
    tbl.style = "Table Grid"

    for i, h in enumerate(EN_TETES):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.font.bold      = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size      = Pt(9)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        tc_pr = cell._tc.get_or_add_tcPr()
        shd   = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E65100"); shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear"); tc_pr.append(shd)

    for row in rows:
        tr = tbl.add_row()
        for i, champ in enumerate(CHAMPS):
            cell = tr.cells[i]
            cell.text = str(row[champ])
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            cell.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.save(chemin)


def _export_pdf(chemin, rows):
    if not PDF_OK:
        messagebox.showerror("Manquant", "pip install reportlab"); return

    doc    = SimpleDocTemplate(chemin, pagesize=A4,
                               leftMargin=1.5*cm, rightMargin=1.5*cm,
                               topMargin=2*cm, bottomMargin=2*cm)
    styles = getSampleStyleSheet()
    story  = []

    story.append(Paragraph("Liste des Clients",
        ParagraphStyle("T", parent=styles["Title"],
                       fontSize=18, textColor=colors.HexColor("#E65100"),
                       spaceAfter=4)))
    story.append(Paragraph(
        f"Exporte le {date.today()}  —  {len(rows)} client(s)",
        ParagraphStyle("S", parent=styles["Normal"],
                       fontSize=9, textColor=colors.HexColor("#996633"),
                       spaceAfter=14)))

    data = [EN_TETES]
    for row in rows:
        data.append([str(row[c]) for c in CHAMPS])

    col_w = [1*cm, 3*cm, 4*cm, 2.5*cm, 3.5*cm, 2.5*cm, 2.5*cm]
    tbl   = Table(data, colWidths=col_w, repeatRows=1)

    ts = TableStyle([
        ("BACKGROUND",    (0, 0), (-1, 0), colors.HexColor("#E65100")),
        ("TEXTCOLOR",     (0, 0), (-1, 0), colors.white),
        ("FONTNAME",      (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",      (0, 0), (-1, 0), 8),
        ("ALIGN",         (0, 0), (-1, -1), "CENTER"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING",    (0, 0), (-1, 0), 8),
        ("FONTSIZE",      (0, 1), (-1, -1), 7.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.HexColor("#FFF8F0"), colors.white]),
        ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#E8D8C8")),
        ("TOPPADDING",    (0, 1), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 5),
    ])
    tbl.setStyle(ts)
    story.append(tbl)
    doc.build(story)


def _export_csv(chemin, rows):
    with open(chemin, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv.DictWriter(f, fieldnames=CHAMPS)
        writer.writeheader()
        writer.writerows(rows)


def _export_json(chemin, rows):
    with open(chemin, "w", encoding="utf-8") as f:
        json.dump({
            "export_date":   str(date.today()),
            "total_records": len(rows),
            "clients":       rows,
        }, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════
#   IMPORTER CLIENTS
# ══════════════════════════════════════════════════════════
def _importer(table, format_):
    ext_map = {
        "excel": [("Fichier Excel", "*.xlsx *.xls"), ("Tous", "*.*")],
        "csv":   [("Fichier CSV",   "*.csv"),         ("Tous", "*.*")],
        "json":  [("Fichier JSON",  "*.json"),         ("Tous", "*.*")],
    }
    chemin = filedialog.askopenfilename(
        title=f"Importer depuis {format_.upper()}",
        filetypes=ext_map.get(format_, [("Tous", "*.*")]),
    )
    if not chemin:
        return

    try:
        if format_ == "excel":  lignes = _import_excel(chemin)
        elif format_ == "csv":  lignes = _import_csv(chemin)
        elif format_ == "json": lignes = _import_json(chemin)
        else:                   lignes = []

        if not lignes:
            messagebox.showwarning("Vide", "Aucune ligne trouvee dans le fichier.")
            return

        _fenetre_preview(table, lignes, chemin)
    except Exception as ex:
        messagebox.showerror("Erreur import", f"Erreur :\n{ex}")


def _import_excel(chemin):
    # Essayer pandas d'abord (compatible .xls), sinon openpyxl
    if PANDAS_OK:
        df = pd.read_excel(chemin)
        col_map = {
            "nom": "nom", "name": "nom",
            "email": "email", "mail": "email",
            "telephone": "telephone", "tel": "telephone", "phone": "telephone",
            "adresse": "adresse", "address": "adresse",
            "ville": "ville", "city": "ville",
        }
        rows = []
        for _, row in df.iterrows():
            ligne = {}
            for col in df.columns:
                key = col_map.get(col.lower().strip(), col.lower().strip())
                ligne[key] = str(row[col]) if pd.notna(row[col]) else ""
            rows.append([
                ligne.get("nom", ""),
                ligne.get("email", ""),
                ligne.get("telephone", ""),
                ligne.get("adresse", ""),
                ligne.get("ville", ""),
            ])
        return rows

    if not EXCEL_OK:
        messagebox.showerror("Manquant",
            "Installez pandas ou openpyxl :\npip install pandas openpyxl")
        return []

    wb   = openpyxl.load_workbook(chemin, data_only=True)
    ws   = wb.active
    all_rows = list(ws.iter_rows(values_only=True))
    # Detecter en-tete
    start = 1 if (all_rows and all_rows[0] and
                  any(str(c or "").lower() in ("nom", "name", "email")
                      for c in all_rows[0])) else 0
    return [[str(c) if c is not None else "" for c in row]
            for row in all_rows[start:] if any(c is not None for c in row)]


def _import_csv(chemin):
    lignes = []
    with open(chemin, encoding="utf-8-sig", newline="") as f:
        reader = csv.reader(f)
        next(reader, None)
        for row in reader:
            if row:
                lignes.append(row)
    return lignes


def _import_json(chemin):
    with open(chemin, encoding="utf-8") as f:
        data = json.load(f)
    if isinstance(data, dict):
        data = data.get("clients", data.get("customers", []))
    champs_j = ["nom", "email", "telephone", "adresse", "ville"]
    return [[str(item.get(k, "")) for k in champs_j]
            for item in data if isinstance(item, dict)]


def _fenetre_preview(table, lignes, chemin):
    t   = get_theme()
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
    win = tk.Toplevel()
    win.title(f"Apercu import — {len(lignes)} ligne(s)")
    win.geometry("780x460")
    win.configure(bg=t["bg"])
    win.grab_set()

    hdr = tk.Frame(win, bg="#1565C0", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=f"📥  Apercu Import — {os.path.basename(chemin)}",
             bg="#1565C0", fg="white",
             font=("Arial", 13, "bold")).pack(pady=12)

    frame_prev = tk.Frame(win, bg=t["bg"])
    frame_prev.pack(fill="both", expand=True, padx=16, pady=10)

    cols_prev = ["Nom", "Email", "Telephone", "Adresse", "Ville"]

    prev_tree = ttk.Treeview(frame_prev, columns=cols_prev,
                             show="headings", height=12,
                             style="CL.Treeview")
    for col in cols_prev:
        prev_tree.heading(col, text=col)
        prev_tree.column(col, width=140, anchor="center")

    bg_row1 = "#1A1A35" if is_dark else "#FFFFFF"
    bg_row2 = "#12122A" if is_dark else "#FFF8F0"
    prev_tree.tag_configure("pair",   background=bg_row1)
    prev_tree.tag_configure("impair", background=bg_row2)

    for i, row in enumerate(lignes[:50]):
        # Adapter : ignorer ID si present (8 colonnes) ou prendre 5 premiers
        if len(row) >= 7:
            vals = row[1:6]   # sauter ID
        elif len(row) >= 5:
            vals = row[:5]
        else:
            vals = list(row) + [""] * (5 - len(row))

        prev_tree.insert("", "end",
                         tags=("pair" if i % 2 == 0 else "impair",),
                         values=vals[:5])

    sb = ttk.Scrollbar(frame_prev, orient="vertical", command=prev_tree.yview)
    prev_tree.configure(yscroll=sb.set)
    sb.pack(side="right", fill="y")
    prev_tree.pack(fill="both", expand=True)

    if len(lignes) > 50:
        tk.Label(win, text=f"... et {len(lignes)-50} autres lignes.",
                 bg=t["bg"], fg="#FF8F00",
                 font=("Arial", 9, "italic")).pack()

    tk.Label(win,
             text=f"{len(lignes)} ligne(s) detectee(s). Les doublons sur le nom sont ignores.",
             bg=t["bg"], fg=t["text"],
             font=("Arial", 9)).pack(pady=(0, 4))

    btn_frame = tk.Frame(win, bg=t["bg"])
    btn_frame.pack(pady=8)

    def confirmer():
        nb_ok, nb_err = _inserer_lignes(lignes)
        charger_clients(table)
        win.destroy()
        messagebox.showinfo(
            "Import termine",
            f"Import termine.\n\n"
            f"Inseres avec succes : {nb_ok}\n"
            f"Ignores (erreurs)   : {nb_err}",
        )

    _btn_ok(btn_frame, f"✅  Importer {len(lignes)} client(s)",
            confirmer, "#1565C0", "#0D47A1").pack(side="left", padx=8)
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2",
              padx=16, pady=9).pack(side="left", padx=8)


def _inserer_lignes(lignes):
    nb_ok = nb_err = 0
    for row in lignes:
        try:
            # Accepter 5 colonnes (sans ID) ou 7 (avec ID)
            if len(row) >= 7:
                _, nom, email, tel, adr, ville, _ = (
                    row[0], row[1], row[2], row[3], row[4], row[5], row[6])
            elif len(row) >= 5:
                nom, email, tel, adr, ville = row[0], row[1], row[2], row[3], row[4]
            else:
                nb_err += 1; continue

            nom = str(nom).strip()
            if not nom:
                nb_err += 1; continue

            # Eviter les doublons sur le nom
            if session.query(Client).filter_by(nom=nom).first():
                nb_err += 1; continue

            c = Client(
                nom=nom,
                email=str(email).strip() if email else "",
                telephone=str(tel).strip() if tel else "",
                adresse=str(adr).strip() if adr else "",
                ville=str(ville).strip() if ville else "",
                date_creation=date.today(),
            )
            session.add(c)
            nb_ok += 1
        except Exception:
            nb_err += 1
    session.commit()
    return nb_ok, nb_err