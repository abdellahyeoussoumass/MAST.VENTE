import customtkinter as ctk
import tkinter as tk
from tkinter import ttk, messagebox, filedialog
from database.db import session
from database.models import BonCommande, Client, Produit
from datetime import date
from utils.theme import get_theme
import os


# ══════════════════════════════════════════════════════════
#   IMPORTS OPTIONNELS  (installés si disponibles)
# ══════════════════════════════════════════════════════════
try:
    import openpyxl
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    EXCEL_OK = True
except ImportError:
    EXCEL_OK = False

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    WORD_OK = True
except ImportError:
    WORD_OK = False

try:
    import csv
    CSV_OK = True
except ImportError:
    CSV_OK = False

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Table, TableStyle,
                                    Paragraph, Spacer)
    from reportlab.lib.units import cm
    PDF_OK = True
except ImportError:
    PDF_OK = False

try:
    import json
    JSON_OK = True
except ImportError:
    JSON_OK = False


# ══════════════════════════════════════════════════════════
#   STYLE TREEVIEW
# ══════════════════════════════════════════════════════════
def _style_treeview(is_dark):
    style = ttk.Style()
    style.theme_use("clam")

    bg_tree = "#12122A" if is_dark else "#FFFFFF"
    bg_head = "#E65100"
    bg_sel  = "#FF6D00" if is_dark else "#FFB74D"
    fg_text = "#E8E0FF" if is_dark else "#1A1A1A"
    bg_row1 = "#1A1A35" if is_dark else "#FFFFFF"
    bg_row2 = "#12122A" if is_dark else "#FFF8F0"

    style.configure("BC.Treeview",
                    background=bg_tree, foreground=fg_text,
                    rowheight=32, fieldbackground=bg_tree,
                    borderwidth=0, font=("Arial", 10))
    style.configure("BC.Treeview.Heading",
                    background=bg_head, foreground="white",
                    font=("Arial", 10, "bold"),
                    relief="flat", borderwidth=0)
    style.map("BC.Treeview",
              background=[("selected", bg_sel)],
              foreground=[("selected", "#FFFFFF")])
    style.map("BC.Treeview.Heading",
              background=[("active", "#BF360C")])
    return bg_row1, bg_row2


# ══════════════════════════════════════════════════════════
#   CONSTANTES STATUT
# ══════════════════════════════════════════════════════════
STATUT_ICONE = {
    "Paye":           "✅",
    "Pas encore paye":"⏳",
    "Annule":         "❌",
    "En cours":       "🔄",
}


# ══════════════════════════════════════════════════════════
#   HELPERS UI
# ══════════════════════════════════════════════════════════
def _fenetre_modale(titre, largeur=460, hauteur=580):
    t   = get_theme()
    win = tk.Toplevel()
    win.title(titre)
    win.geometry(f"{largeur}x{hauteur}")
    win.configure(bg=t["bg"])
    win.resizable(False, False)
    win.grab_set()

    hdr = tk.Frame(win, bg="#E65100", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=titre, bg="#E65100", fg="white",
             font=("Arial", 14, "bold")).pack(pady=12)

    body = tk.Frame(win, bg=t["bg"])
    body.pack(fill="both", expand=True, padx=24, pady=16)
    return win, body, t


def _lbl(body, text, t):
    tk.Label(body, text=text, bg=t["bg"], fg=t["text"],
             font=("Arial", 11, "bold"), anchor="w").pack(fill="x", pady=(10, 2))


def _entry(body, t, default=""):
    e = tk.Entry(body, font=("Arial", 11),
                 bg=t["card"], fg=t["text"],
                 insertbackground=t["text"],
                 relief="flat", bd=6)
    e.pack(fill="x", ipady=4)
    if default:
        e.insert(0, default)
    return e


def _combo(body, t, values, default=""):
    var = tk.StringVar(value=default)
    cb  = ttk.Combobox(body, textvariable=var, values=values,
                       font=("Arial", 11), state="readonly")
    cb.pack(fill="x", ipady=4)
    cb.var = var
    return cb


def _btn_ok(parent, text, cmd, color="#E65100", hover="#BF360C"):
    return tk.Button(
        parent, text=text, command=cmd,
        bg=color, fg="white",
        font=("Arial", 12, "bold"),
        relief="flat", cursor="hand2",
        padx=16, pady=9,
        activebackground=hover,
        activeforeground="white",
    )


def _resume_frame(body, t, champs):
    info_bg = "#1A1A35" if t["bg"] in ("#1A1A2E", "#0F0F1A") else "#FFF0E5"
    frame   = tk.Frame(body, bg=info_bg)
    frame.pack(fill="x", pady=(0, 10))
    for label, val in champs:
        row = tk.Frame(frame, bg=info_bg)
        row.pack(fill="x", padx=14, pady=3)
        tk.Label(row, text=f"{label} :", bg=info_bg,
                 fg="#FF8F00", font=("Arial", 10, "bold"),
                 width=12, anchor="w").pack(side="left")
        tk.Label(row, text=str(val), bg=info_bg,
                 fg=t["text"], font=("Arial", 10)).pack(side="left")


# ══════════════════════════════════════════════════════════
#   TRI COLONNES
# ══════════════════════════════════════════════════════════
_tri_etat = {}

def _trier(table, col):
    rows = [(table.set(r, col), r) for r in table.get_children("")]
    rev  = _tri_etat.get(col, False)
    try:
        rows.sort(key=lambda x: float(
            x[0].replace(",", "").replace("%", "").replace(" ", "")),
            reverse=rev)
    except ValueError:
        rows.sort(key=lambda x: x[0].lower(), reverse=rev)
    for i, (_, r) in enumerate(rows):
        table.move(r, "", i)
    _tri_etat[col] = not rev


# ══════════════════════════════════════════════════════════
#   MENUS DEROULANTS IMPORT / EXPORT
# ══════════════════════════════════════════════════════════
def _show_dropdown_import(btn_widget, table, t):
    """Affiche un menu deroulant sous le bouton Importer."""
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
    bg_menu   = "#1A1A35" if is_dark else "#FFFFFF"
    bg_hover  = "#252545" if is_dark else "#FFF0E5"
    fg_text   = "#E8E0FF" if is_dark else "#1A1A1A"
    border_c  = "#2A2A4A" if is_dark else "#E8D8C8"

    items = [
        ("📊  Excel (.xlsx)",  "#1565C0", "excel"),
        ("📄  CSV (.csv)",     "#00838F", "csv"),
        ("📋  JSON (.json)",   "#6A1B9A", "json"),
    ]

    menu = tk.Toplevel()
    menu.overrideredirect(True)
    menu.configure(bg=border_c)
    menu.attributes("-topmost", True)

    # Position sous le bouton
    btn_widget.update_idletasks()
    x = btn_widget.winfo_rootx()
    y = btn_widget.winfo_rooty() + btn_widget.winfo_height() + 2
    menu.geometry(f"200x{len(items) * 40 + 10}+{x}+{y}")

    inner = tk.Frame(menu, bg=bg_menu, bd=0)
    inner.pack(fill="both", expand=True, padx=1, pady=1)

    # Titre section
    tk.Label(inner, text="Choisir le format :",
             bg=bg_menu, fg="#FF8F00",
             font=("Arial", 8, "bold"),
             anchor="w").pack(fill="x", padx=10, pady=(6, 2))

    def _fermer():
        try:
            menu.destroy()
        except Exception:
            pass

    def _item_click(fmt):
        _fermer()
        importer_fichier(table, fmt)

    for label, couleur, fmt in items:
        f = tk.Frame(inner, bg=bg_menu, cursor="hand2")
        f.pack(fill="x", padx=4, pady=1)

        lbl = tk.Label(f, text=label,
                       bg=bg_menu, fg=couleur,
                       font=("Arial", 10, "bold"),
                       anchor="w", padx=12, pady=8)
        lbl.pack(fill="x")

        # Hover effect
        fmt_local = fmt
        def _enter(e, frame=f, lbl_=lbl):
            frame.configure(bg=bg_hover)
            lbl_.configure(bg=bg_hover)
        def _leave(e, frame=f, lbl_=lbl):
            frame.configure(bg=bg_menu)
            lbl_.configure(bg=bg_menu)

        f.bind("<Enter>", _enter)
        f.bind("<Leave>", _leave)
        lbl.bind("<Enter>", _enter)
        lbl.bind("<Leave>", _leave)
        f.bind("<Button-1>",   lambda e, fmt=fmt_local: _item_click(fmt))
        lbl.bind("<Button-1>", lambda e, fmt=fmt_local: _item_click(fmt))

    # Fermer si clic ailleurs
    menu.bind("<FocusOut>", lambda e: _fermer())
    menu.focus_set()


def _show_dropdown_export(btn_widget, table, t):
    """Affiche un menu deroulant sous le bouton Exporter."""
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
    bg_menu  = "#1A1A35" if is_dark else "#FFFFFF"
    bg_hover = "#252545" if is_dark else "#FFF0E5"
    border_c = "#2A2A4A" if is_dark else "#E8D8C8"

    items = [
        ("📊  Excel (.xlsx)",   "#1565C0", "excel"),
        ("📝  Word (.docx)",    "#1976D2", "word"),
        ("📕  PDF (.pdf)",      "#C62828", "pdf"),
        ("📄  CSV (.csv)",      "#00838F", "csv"),
        ("📋  JSON (.json)",    "#6A1B9A", "json"),
    ]

    menu = tk.Toplevel()
    menu.overrideredirect(True)
    menu.configure(bg=border_c)
    menu.attributes("-topmost", True)

    btn_widget.update_idletasks()
    x = btn_widget.winfo_rootx()
    y = btn_widget.winfo_rooty() + btn_widget.winfo_height() + 2
    menu.geometry(f"200x{len(items) * 40 + 10}+{x}+{y}")

    inner = tk.Frame(menu, bg=bg_menu, bd=0)
    inner.pack(fill="both", expand=True, padx=1, pady=1)

    tk.Label(inner, text="Choisir le format :",
             bg=bg_menu, fg="#FF8F00",
             font=("Arial", 8, "bold"),
             anchor="w").pack(fill="x", padx=10, pady=(6, 2))

    def _fermer():
        try:
            menu.destroy()
        except Exception:
            pass

    def _item_click(fmt):
        _fermer()
        exporter_bc(table, fmt)

    for label, couleur, fmt in items:
        f = tk.Frame(inner, bg=bg_menu, cursor="hand2")
        f.pack(fill="x", padx=4, pady=1)

        lbl = tk.Label(f, text=label,
                       bg=bg_menu, fg=couleur,
                       font=("Arial", 10, "bold"),
                       anchor="w", padx=12, pady=8)
        lbl.pack(fill="x")

        fmt_local = fmt
        def _enter(e, frame=f, lbl_=lbl):
            frame.configure(bg=bg_hover)
            lbl_.configure(bg=bg_hover)
        def _leave(e, frame=f, lbl_=lbl):
            frame.configure(bg=bg_menu)
            lbl_.configure(bg=bg_menu)

        f.bind("<Enter>", _enter)
        f.bind("<Leave>", _leave)
        lbl.bind("<Enter>", _enter)
        lbl.bind("<Leave>", _leave)
        f.bind("<Button-1>",   lambda e, fmt=fmt_local: _item_click(fmt))
        lbl.bind("<Button-1>", lambda e, fmt=fmt_local: _item_click(fmt))

    menu.bind("<FocusOut>", lambda e: _fermer())
    menu.focus_set()


# ══════════════════════════════════════════════════════════
#   AFFICHER BONS DE COMMANDE
# ══════════════════════════════════════════════════════════
def afficher_bons_commande(parent):
    for widget in parent.winfo_children():
        widget.destroy()

    t       = get_theme()
    is_dark = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")

    # En-tete
    header = ctk.CTkFrame(parent, fg_color=t["card"], corner_radius=0)
    header.pack(fill="x")
    ctk.CTkLabel(header, text="🛒  Bons de Commande",
                 font=("Arial", 22, "bold"),
                 text_color="#E65100").pack(side="left", padx=24, pady=14)

    # ── Barre outils ligne 1 : actions ───────
    toolbar = ctk.CTkFrame(parent, fg_color="transparent")
    toolbar.pack(fill="x", padx=20, pady=(10, 2))

    btns_actions = [
        ("➕  Ajouter",      "#E65100", "#BF360C", lambda: ajouter_bc(table)),
        ("✏️  Modifier",     "#FF8F00", "#E65100", lambda: modifier_statut(table)),
        ("🗑️  Supprimer",   "#C62828", "#8B0000", lambda: supprimer_bc(table)),
        ("👁️  Consulter",   "#2E7D32", "#1B5E20", lambda: detail_bc(table)),
    ]
    for txt, fg, hover, cmd in btns_actions:
        ctk.CTkButton(toolbar, text=txt, width=128, height=36,
                      fg_color=fg, hover_color=hover,
                      font=("Arial", 11, "bold"),
                      corner_radius=8,
                      command=cmd).pack(side="left", padx=4)

    # ── Bouton IMPORTER avec menu deroulant ───
    btn_import_widget = ctk.CTkButton(
        toolbar, text="📥  Importer ▾", width=138, height=36,
        fg_color="#1565C0", hover_color="#0D47A1",
        font=("Arial", 11, "bold"), corner_radius=8,
        command=lambda: _show_dropdown_import(btn_import_widget, table, t)
    )
    btn_import_widget.pack(side="left", padx=4)

    # ── Bouton EXPORTER avec menu deroulant ───
    btn_export_widget = ctk.CTkButton(
        toolbar, text="📤  Exporter ▾", width=138, height=36,
        fg_color="#6A1B9A", hover_color="#4A148C",
        font=("Arial", 11, "bold"), corner_radius=8,
        command=lambda: _show_dropdown_export(btn_export_widget, table, t)
    )
    btn_export_widget.pack(side="left", padx=4)

    # ── Recherche + filtre statut ─────────────
    search_frame = ctk.CTkFrame(parent, fg_color="transparent")
    search_frame.pack(fill="x", padx=20, pady=(2, 4))

    count_var    = tk.StringVar()
    filtre_actif = {"valeur": None}
    dropdown_ref = [None]
    table_ref    = [None]

    FILTRES_STATUT = [
        ("Tous",            "#555555"),
        ("Paye",            "#43A047"),
        ("Pas encore paye", "#FFA000"),
        ("En cours",        "#42A5F5"),
        ("Annule",          "#E53935"),
    ]
    ICONES_F = {
        "Tous": "📋", "Paye": "✅",
        "Pas encore paye": "⏳", "En cours": "🔄", "Annule": "❌",
    }

    # Compteur à droite
    ctk.CTkLabel(search_frame, textvariable=count_var,
                 font=("Arial", 11),
                 text_color="#FF8F00").pack(side="right", padx=10)

    # Bouton filtre statut
    var_btn_txt = tk.StringVar(value="⚙  Filtre ▾")
    btn_filtre  = tk.Button(
        search_frame, textvariable=var_btn_txt,
        bg="#FF8F00", fg="white",
        font=("Arial", 10, "bold"),
        relief="flat", cursor="hand2",
        padx=10, pady=6,
        activebackground="#E65100",
        activeforeground="white",
    )
    btn_filtre.pack(side="left", padx=(0, 4))

    # Cadre champ texte
    frame_input = tk.Frame(search_frame, bg=t["card"],
                           highlightbackground="#E65100",
                           highlightthickness=1)
    frame_input.pack(side="left")

    recherche_var = tk.StringVar()
    PLACEHOLDER   = "🔍  Rechercher dans tous les bons de commande..."

    e_search = tk.Entry(
        frame_input, textvariable=recherche_var,
        font=("Arial", 11), bg=t["card"], fg="#888888",
        insertbackground=t["text"],
        relief="flat", bd=6, width=38,
    )
    e_search.insert(0, PLACEHOLDER)
    e_search.pack(side="left", ipady=5)

    btn_clear = tk.Button(
        frame_input, text=" ✕ ",
        bg=t["card"], fg="#C62828",
        font=("Arial", 10, "bold"),
        relief="flat", cursor="hand2",
        activebackground=t["card"],
        activeforeground="#8B0000", bd=0,
    )

    def _show_ph():
        if not recherche_var.get():
            e_search.configure(fg="#888888")
            e_search.delete(0, "end")
            e_search.insert(0, PLACEHOLDER)

    def _hide_ph(e=None):
        if e_search.get() == PLACEHOLDER:
            e_search.delete(0, "end")
            e_search.configure(fg=t["text"])

    e_search.bind("<FocusIn>",  _hide_ph)
    e_search.bind("<FocusOut>", lambda e: _show_ph() if not recherche_var.get() else None)

    def _effacer_tout():
        filtre_actif["valeur"] = None
        var_btn_txt.set("⚙  Filtre ▾")
        btn_filtre.configure(bg="#FF8F00")
        recherche_var.set("")
        _show_ph()
        btn_clear.pack_forget()
        if table_ref[0]:
            charger_bcs(table_ref[0], count_var)

    btn_clear.configure(command=_effacer_tout)

    def _appliquer_recherche():
        if not table_ref[0]:
            return
        tbl = table_ref[0]
        texte = recherche_var.get().strip()
        if texte == PLACEHOLDER:
            texte = ""
        texte_low = texte.lower()
        statut_v  = filtre_actif["valeur"]

        bcs_all = session.query(BonCommande).order_by(BonCommande.id.desc()).all()

        if statut_v and statut_v != "Tous":
            bcs_all = [b for b in bcs_all if (b.statut or "") == statut_v]

        if texte_low:
            def _match(b):
                client  = session.query(Client).filter_by(id=b.client_id).first()
                produit = session.query(Produit).filter_by(id=b.produit_id).first()
                return any(texte_low in str(v).lower() for v in [
                    b.numero_bc,
                    client.nom  if client  else "",
                    produit.nom if produit else "",
                    b.categorie, b.statut,
                    b.prix_ht, b.prix_ttc, b.prix_total, b.date_bc,
                ])
            bcs_all = [b for b in bcs_all if _match(b)]

        for row in tbl.get_children():
            tbl.delete(row)

        for i, b in enumerate(bcs_all):
            client  = session.query(Client).filter_by(id=b.client_id).first()
            produit = session.query(Produit).filter_by(id=b.produit_id).first()
            statut_raw = b.statut or ""
            icone      = STATUT_ICONE.get(statut_raw, "")
            tags = ("pair" if i % 2 == 0 else "impair",)
            if statut_raw in STATUT_ICONE:
                tags = tags + (statut_raw,)
            tbl.insert("", "end", tags=tags, values=(
                b.id, b.numero_bc or "-",
                client.nom  if client  else "N/A",
                produit.nom if produit else "N/A",
                b.categorie or "-", b.quantite,
                f"{b.prix_ht:,.2f}", f"{b.prix_ttc:,.2f}",
                f"{b.prix_total:,.2f}",
                f"{icone} {statut_raw}".strip() if icone else statut_raw,
                b.date_bc,
            ))

        nb  = len(bcs_all)
        suf = f"  |  Filtre: {statut_v}" if statut_v and statut_v != "Tous" else ""
        if texte:
            suf += f"  |  \"{texte}\""
        count_var.set(f"{nb} bon(s) de commande{suf}")

        if texte or (statut_v and statut_v != "Tous"):
            btn_clear.pack(side="right")
        else:
            btn_clear.pack_forget()

    def _on_frappe(*_):
        texte = recherche_var.get()
        if texte and texte != PLACEHOLDER:
            _appliquer_recherche()
        elif not texte:
            btn_clear.pack_forget()
            if table_ref[0]:
                charger_bcs(table_ref[0], count_var)

    recherche_var.trace("w", _on_frappe)

    def _fermer_dd():
        try:
            if dropdown_ref[0] and dropdown_ref[0].winfo_exists():
                dropdown_ref[0].destroy()
        except Exception:
            pass
        dropdown_ref[0] = None

    def _activer_filtre(valeur):
        _fermer_dd()
        filtre_actif["valeur"] = valeur
        if valeur == "Tous":
            var_btn_txt.set("⚙  Filtre ▾")
            btn_filtre.configure(bg="#FF8F00")
        else:
            icone = ICONES_F.get(valeur, "🔍")
            var_btn_txt.set(f"{icone} {valeur} ▾")
            col_map = {
                "Paye":            "#2E7D32",
                "Pas encore paye": "#E65100",
                "En cours":        "#1565C0",
                "Annule":          "#C62828",
            }
            btn_filtre.configure(bg=col_map.get(valeur, "#FF8F00"))
        _appliquer_recherche()

    def _ouvrir_filtre(e=None):
        _fermer_dd()
        is_d  = t["bg"] in ("#1A1A2E", "#0F0F1A", "#0D0D1A")
        bg_m  = "#1A1A35" if is_d else "#FFFFFF"
        bg_h  = "#252545" if is_d else "#FFF0E5"
        brd   = "#2A2A4A" if is_d else "#E8D8C8"
        muted = "#A0A0C0" if is_d else "#999999"

        dd = tk.Toplevel()
        dd.overrideredirect(True)
        dd.configure(bg=brd)
        dd.attributes("-topmost", True)
        dropdown_ref[0] = dd

        btn_filtre.update_idletasks()
        x = btn_filtre.winfo_rootx()
        y = btn_filtre.winfo_rooty() + btn_filtre.winfo_height() + 2
        h = len(FILTRES_STATUT) * 38 + 28
        dd.geometry(f"230x{h}+{x}+{y}")

        inner = tk.Frame(dd, bg=bg_m)
        inner.pack(fill="both", expand=True, padx=1, pady=1)
        tk.Label(inner, text="  Filtrer par statut",
                 bg=bg_m, fg=muted, font=("Arial", 8, "bold"),
                 anchor="w").pack(fill="x", padx=8, pady=(5, 2))

        for label, couleur in FILTRES_STATUT:
            est_actif = filtre_actif["valeur"] == label
            bg_i = bg_h if est_actif else bg_m

            rf = tk.Frame(inner, bg=bg_i, cursor="hand2")
            rf.pack(fill="x", padx=4, pady=1)

            cv = tk.Canvas(rf, width=10, height=10,
                           bg=bg_i, highlightthickness=0)
            cv.create_oval(1, 1, 9, 9, fill=couleur, outline="")
            cv.pack(side="left", padx=(10, 4), pady=10)

            tk.Label(rf, text="✓" if est_actif else "  ",
                     bg=bg_i, fg=couleur,
                     font=("Arial", 10, "bold"), width=2).pack(side="left")

            lb = tk.Label(rf, text=label, bg=bg_i, fg=couleur,
                          font=("Arial", 10, "bold"),
                          anchor="w", pady=8)
            lb.pack(side="left", fill="x", expand=True)

            def _mk(fr=rf, lb_=lb, cv_=cv, act=est_actif):
                def _in(e):
                    fr.configure(bg=bg_h)
                    lb_.configure(bg=bg_h)
                    cv_.configure(bg=bg_h)
                def _out(e):
                    c = bg_h if act else bg_m
                    fr.configure(bg=c)
                    lb_.configure(bg=c)
                    cv_.configure(bg=c)
                return _in, _out

            _in, _out = _mk()
            for w in (rf, lb, cv):
                w.bind("<Enter>",    _in)
                w.bind("<Leave>",    _out)
                w.bind("<Button-1>", lambda e, v=label: _activer_filtre(v))

        dd.bind("<FocusOut>", lambda e: _fermer_dd())
        dd.focus_set()

    btn_filtre.configure(command=_ouvrir_filtre)

    # ── Tableau ──────────────────────────────
    frame_table = ctk.CTkFrame(parent, corner_radius=12)
    frame_table.pack(fill="both", expand=True, padx=20, pady=(4, 16))

    bg_row1, bg_row2 = _style_treeview(is_dark)

    colonnes = ("ID", "N BC", "Client", "Produit", "Categorie",
                "Qte", "Prix HT", "Prix TTC", "Total", "Statut", "Date")

    table = ttk.Treeview(frame_table, columns=colonnes,
                         show="headings", height=20,
                         style="BC.Treeview",
                         selectmode="browse")

    largeurs = {"ID": 45, "N BC": 100, "Client": 140,
                "Produit": 130, "Categorie": 100,
                "Qte": 55, "Prix HT": 85,
                "Prix TTC": 85, "Total": 100,
                "Statut": 110, "Date": 100}

    for col in colonnes:
        table.heading(col, text=col,
                      command=lambda c=col: _trier(table, c))
        table.column(col, width=largeurs.get(col, 90), anchor="center")

    table.tag_configure("pair",   background=bg_row1)
    table.tag_configure("impair", background=bg_row2)
    table.tag_configure("Paye",           foreground="#43A047")
    table.tag_configure("Pas encore paye",foreground="#FFA000")
    table.tag_configure("Annule",         foreground="#E53935")
    table.tag_configure("En cours",       foreground="#42A5F5")

    sb_v = ttk.Scrollbar(frame_table, orient="vertical",   command=table.yview)
    sb_h = ttk.Scrollbar(frame_table, orient="horizontal", command=table.xview)
    table.configure(yscroll=sb_v.set, xscroll=sb_h.set)
    sb_v.pack(side="right",  fill="y")
    sb_h.pack(side="bottom", fill="x")
    table.pack(fill="both", expand=True, padx=2, pady=2)

    table.bind("<Double-1>", lambda e: detail_bc(table))

    table_ref[0] = table
    charger_bcs(table, count_var)
    return table


# ══════════════════════════════════════════════════════════
#   CHARGER BCS
# ══════════════════════════════════════════════════════════
def charger_bcs(table, count_var=None):
    for row in table.get_children():
        table.delete(row)

    bcs = session.query(BonCommande).order_by(BonCommande.id.desc()).all()

    for i, b in enumerate(bcs):
        client  = session.query(Client).filter_by(id=b.client_id).first()
        produit = session.query(Produit).filter_by(id=b.produit_id).first()

        statut_raw = b.statut or ""
        icone      = STATUT_ICONE.get(statut_raw, "")
        tags       = ("pair" if i % 2 == 0 else "impair",)
        if statut_raw in STATUT_ICONE:
            tags = tags + (statut_raw,)

        table.insert("", "end", tags=tags, values=(
            b.id,
            b.numero_bc or "-",
            client.nom  if client  else "N/A",
            produit.nom if produit else "N/A",
            b.categorie or "-",
            b.quantite,
            f"{b.prix_ht:,.2f}",
            f"{b.prix_ttc:,.2f}",
            f"{b.prix_total:,.2f}",
            f"{icone} {statut_raw}".strip() if icone else statut_raw,
            b.date_bc,
        ))

    if count_var is not None:
        count_var.set(f"{len(bcs)} bons de commande")


# ══════════════════════════════════════════════════════════
#   RECHERCHER
# ══════════════════════════════════════════════════════════
def rechercher(texte, table):
    for row in table.get_children():
        table.delete(row)

    texte_low = texte.lower()
    bcs       = session.query(BonCommande).all()

    for i, b in enumerate(bcs):
        client  = session.query(Client).filter_by(id=b.client_id).first()
        produit = session.query(Produit).filter_by(id=b.produit_id).first()
        nom_c   = client.nom  if client  else ""
        nom_p   = produit.nom if produit else ""

        if (texte_low in nom_c.lower()
                or texte_low in nom_p.lower()
                or texte_low in (b.numero_bc   or "").lower()
                or texte_low in (b.statut      or "").lower()
                or texte_low in (b.categorie   or "").lower()):

            statut_raw = b.statut or ""
            icone      = STATUT_ICONE.get(statut_raw, "")
            tags       = ("pair" if i % 2 == 0 else "impair",)

            table.insert("", "end", tags=tags, values=(
                b.id, b.numero_bc or "-",
                nom_c or "N/A", nom_p or "N/A",
                b.categorie or "-", b.quantite,
                f"{b.prix_ht:,.2f}", f"{b.prix_ttc:,.2f}",
                f"{b.prix_total:,.2f}",
                f"{icone} {statut_raw}".strip() if icone else statut_raw,
                b.date_bc,
            ))


# ══════════════════════════════════════════════════════════
#   AJOUTER BC
# ══════════════════════════════════════════════════════════
def ajouter_bc(table):
    clients  = session.query(Client).all()
    produits = session.query(Produit).all()

    win, body, t = _fenetre_modale("➕  Nouveau Bon de Commande", 460, 580)

    _lbl(body, "N Bon de Commande", t)
    e_num = _entry(body, t)

    _lbl(body, "Client", t)
    cb_client = _combo(body, t, [c.nom for c in clients])

    _lbl(body, "Produit", t)
    cb_produit = _combo(body, t, [p.nom for p in produits])

    _lbl(body, "Quantite", t)
    e_qte = _entry(body, t)

    _lbl(body, "Statut", t)
    cb_statut = _combo(body, t,
        ["Paye", "Pas encore paye", "En cours", "Annule"],
        default="En cours")

    # Apercu prix
    lbl_prev = tk.Label(body, text="", bg=t["bg"],
                        fg="#FF8F00", font=("Arial", 10, "italic"))
    lbl_prev.pack(pady=(6, 0))

    def maj_prev(*_):
        try:
            p = session.query(Produit).filter_by(
                nom=cb_produit.var.get()).first()
            if p and e_qte.get():
                qte   = int(e_qte.get())
                total = round(p.prix_ttc * qte, 2)
                lbl_prev.config(
                    text=f"HT: {p.prix_ht:,.2f}  |  TTC: {p.prix_ttc:,.2f}  |  Total: {total:,.2f} MAD")
        except Exception:
            pass

    cb_produit.var.trace("w", maj_prev)
    e_qte.bind("<KeyRelease>", maj_prev)

    def sauvegarder():
        try:
            client  = session.query(Client).filter_by(
                nom=cb_client.var.get()).first()
            produit = session.query(Produit).filter_by(
                nom=cb_produit.var.get()).first()
            if not client:
                messagebox.showerror("Erreur", "Selectionnez un client.", parent=win)
                return
            if not produit:
                messagebox.showerror("Erreur", "Selectionnez un produit.", parent=win)
                return
            qte   = int(e_qte.get())
            total = round(produit.prix_ttc * qte, 2)

            bc = BonCommande(
                numero_bc=e_num.get(),
                client_id=client.id,
                produit_id=produit.id,
                categorie=produit.categorie,
                quantite=qte,
                prix_ht=produit.prix_ht,
                prix_ttc=produit.prix_ttc,
                prix_total=total,
                statut=cb_statut.var.get(),
                date_bc=date.today(),
            )
            session.add(bc)
            session.commit()
            charger_bcs(table)
            win.destroy()
            messagebox.showinfo("Succes", "Bon de commande ajoute avec succes !")
        except Exception as ex:
            messagebox.showerror("Erreur", f"Erreur : {ex}", parent=win)

    _btn_ok(body, "💾  Sauvegarder", sauvegarder).pack(fill="x", pady=(14, 4))
    tk.Button(body, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=8).pack(fill="x")


# ══════════════════════════════════════════════════════════
#   MODIFIER STATUT
# ══════════════════════════════════════════════════════════
def modifier_statut(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez un bon de commande !")
        return

    valeurs = table.item(sel[0])["values"]
    bc      = session.query(BonCommande).filter_by(id=valeurs[0]).first()
    if not bc:
        return

    client  = session.query(Client).filter_by(id=bc.client_id).first()
    produit = session.query(Produit).filter_by(id=bc.produit_id).first()

    win, body, t = _fenetre_modale("✏️  Modifier le Bon de Commande", 420, 380)

    _resume_frame(body, t, [
        ("N BC",     bc.numero_bc or "-"),
        ("Client",   client.nom  if client  else "N/A"),
        ("Produit",  produit.nom if produit else "N/A"),
        ("Total",    f"{bc.prix_total:,.2f} MAD"),
    ])

    _lbl(body, "Nouveau statut", t)
    cb = _combo(body, t,
        ["Paye", "Pas encore paye", "En cours", "Annule"],
        default=bc.statut or "En cours")

    def sauvegarder():
        bc.statut = cb.var.get()
        session.commit()
        charger_bcs(table)
        win.destroy()
        messagebox.showinfo("Succes", "Statut modifie avec succes !")

    _btn_ok(body, "💾  Sauvegarder", sauvegarder).pack(fill="x", pady=(14, 4))
    tk.Button(body, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2", pady=8).pack(fill="x")


# ══════════════════════════════════════════════════════════
#   SUPPRIMER BC
# ══════════════════════════════════════════════════════════
def supprimer_bc(table):
    sel = table.selection()
    if not sel:
        messagebox.showwarning("Attention", "Selectionnez un bon de commande !")
        return

    valeurs = table.item(sel[0])["values"]
    bc      = session.query(BonCommande).filter_by(id=valeurs[0]).first()
    if not bc:
        return

    client = session.query(Client).filter_by(id=bc.client_id).first()
    nom_c  = client.nom if client else "N/A"

    if messagebox.askyesno(
        "Confirmation",
        f"Supprimer le BC {bc.numero_bc or bc.id} ?\n"
        f"Client : {nom_c}\n"
        f"Total  : {bc.prix_total:,.2f} MAD\n\n"
        "Cette action est irreversible.",
    ):
        session.delete(bc)
        session.commit()
        charger_bcs(table)
        messagebox.showinfo("Succes", "Bon de commande supprime avec succes !")


# ══════════════════════════════════════════════════════════
#   UTILITAIRE : collecter les données
# ══════════════════════════════════════════════════════════
def _get_bcs_data(selection_only=False, table=None):
    """Retourne liste de dicts pour tous les BCs ou la sélection."""
    rows = []

    if selection_only and table:
        sel = table.selection()
        ids = [table.item(s)["values"][0] for s in sel]
        bcs = [session.query(BonCommande).filter_by(id=i).first() for i in ids]
        bcs = [b for b in bcs if b]
    else:
        bcs = session.query(BonCommande).order_by(BonCommande.id.desc()).all()

    for b in bcs:
        client  = session.query(Client).filter_by(id=b.client_id).first()
        produit = session.query(Produit).filter_by(id=b.produit_id).first()
        rows.append({
            "id":          b.id,
            "numero_bc":   b.numero_bc or "-",
            "client":      client.nom  if client  else "N/A",
            "produit":     produit.nom if produit else "N/A",
            "categorie":   b.categorie or "-",
            "quantite":    b.quantite,
            "prix_ht":     float(b.prix_ht),
            "prix_ttc":    float(b.prix_ttc),
            "prix_total":  float(b.prix_total),
            "statut":      b.statut or "-",
            "date_bc":     str(b.date_bc),
        })
    return rows


EN_TETES = ["ID", "N BC", "Client", "Produit", "Categorie",
            "Qte", "Prix HT", "Prix TTC", "Total", "Statut", "Date"]

CHAMPS = ["id", "numero_bc", "client", "produit", "categorie",
          "quantite", "prix_ht", "prix_ttc", "prix_total", "statut", "date_bc"]


# ══════════════════════════════════════════════════════════
#   EXPORTER BC
# ══════════════════════════════════════════════════════════
def exporter_bc(table, format_):
    rows = _get_bcs_data()
    if not rows:
        messagebox.showwarning("Vide", "Aucune donnee a exporter.")
        return

    ext_map = {
        "excel": ("Fichier Excel", "*.xlsx", ".xlsx"),
        "word":  ("Document Word",  "*.docx", ".docx"),
        "pdf":   ("Fichier PDF",    "*.pdf",  ".pdf"),
        "csv":   ("Fichier CSV",    "*.csv",  ".csv"),
        "json":  ("Fichier JSON",   "*.json", ".json"),
    }
    label, pattern, ext = ext_map[format_]

    chemin = filedialog.asksaveasfilename(
        title=f"Exporter en {label}",
        defaultextension=ext,
        filetypes=[(label, pattern), ("Tous", "*.*")],
        initialfile=f"bons_commande_{date.today()}{ext}",
    )
    if not chemin:
        return

    try:
        if format_ == "excel":
            _export_excel(chemin, rows)
        elif format_ == "word":
            _export_word(chemin, rows)
        elif format_ == "pdf":
            _export_pdf(chemin, rows)
        elif format_ == "csv":
            _export_csv(chemin, rows)
        elif format_ == "json":
            _export_json(chemin, rows)

        if messagebox.askyesno("Export reussi",
                               f"Fichier cree :\n{chemin}\n\nOuvrir maintenant ?"):
            os.startfile(chemin)

    except Exception as ex:
        messagebox.showerror("Erreur export", f"Erreur lors de l'export :\n{ex}")


# ── Export Excel ──────────────────────────────────────────
def _export_excel(chemin, rows):
    if not EXCEL_OK:
        messagebox.showerror("Manquant", "Installez openpyxl :\npip install openpyxl")
        return

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Bons de Commande"

    # Titre
    ws.merge_cells("A1:K1")
    ws["A1"] = f"Bons de Commande — Exporte le {date.today()}"
    ws["A1"].font      = Font(bold=True, size=14, color="FFFFFF")
    ws["A1"].fill      = PatternFill("solid", fgColor="E65100")
    ws["A1"].alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[1].height = 36

    # En-tetes
    for col_idx, h in enumerate(EN_TETES, start=1):
        cell = ws.cell(row=2, column=col_idx, value=h)
        cell.font      = Font(bold=True, color="FFFFFF", size=10)
        cell.fill      = PatternFill("solid", fgColor="BF360C")
        cell.alignment = Alignment(horizontal="center", vertical="center")
    ws.row_dimensions[2].height = 24

    # Donnees
    fill_pair   = PatternFill("solid", fgColor="FFF8F0")
    fill_impair = PatternFill("solid", fgColor="FFFFFF")
    border_side = Side(style="thin", color="E8D8C8")
    border      = Border(left=border_side, right=border_side,
                         top=border_side, bottom=border_side)

    for r_idx, row in enumerate(rows, start=3):
        fill = fill_pair if r_idx % 2 == 0 else fill_impair
        for c_idx, champ in enumerate(CHAMPS, start=1):
            val  = row[champ]
            cell = ws.cell(row=r_idx, column=c_idx, value=val)
            cell.fill      = fill
            cell.border    = border
            cell.alignment = Alignment(horizontal="center", vertical="center")
            cell.font      = Font(size=9)

            # Couleur statut
            if champ == "statut":
                col_map = {
                    "Paye":           "2E7D32",
                    "Pas encore paye":"E65100",
                    "Annule":         "C62828",
                    "En cours":       "1565C0",
                }
                fg = col_map.get(str(val), "212121")
                cell.font = Font(size=9, bold=True, color=fg)

    # Largeurs colonnes
    largeurs_xl = [6, 12, 18, 18, 14, 6, 10, 10, 12, 16, 12]
    for i, w in enumerate(largeurs_xl, start=1):
        ws.column_dimensions[
            openpyxl.utils.get_column_letter(i)].width = w

    # Ligne total
    total_all = sum(r["prix_total"] for r in rows)
    row_total = len(rows) + 3
    ws.cell(row=row_total, column=9, value=f"TOTAL : {total_all:,.2f} MAD")
    ws.cell(row=row_total, column=9).font = Font(bold=True, color="E65100", size=10)
    ws.cell(row=row_total, column=9).fill = PatternFill("solid", fgColor="FFF0E5")

    wb.save(chemin)


# ── Export Word ───────────────────────────────────────────
def _export_word(chemin, rows):
    if not WORD_OK:
        messagebox.showerror("Manquant", "Installez python-docx :\npip install python-docx")
        return

    doc = Document()

    # Marges
    for section in doc.sections:
        section.top_margin    = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin   = Cm(2)
        section.right_margin  = Cm(2)

    # Titre
    titre = doc.add_heading("Bons de Commande", level=1)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = titre.runs[0]
    run.font.color.rgb = RGBColor(0xE6, 0x51, 0x00)
    run.font.size      = Pt(18)

    # Sous-titre
    st = doc.add_paragraph(f"Exporte le {date.today()}  |  {len(rows)} enregistrement(s)")
    st.alignment = WD_ALIGN_PARAGRAPH.CENTER
    st.runs[0].font.color.rgb = RGBColor(0x99, 0x66, 0x33)
    st.runs[0].font.size      = Pt(10)

    doc.add_paragraph()

    # Tableau
    table_doc = doc.add_table(rows=1, cols=len(EN_TETES))
    table_doc.style = "Table Grid"

    # En-tete
    for i, h in enumerate(EN_TETES):
        cell = table_doc.rows[0].cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].font.bold      = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        cell.paragraphs[0].runs[0].font.size      = Pt(9)
        cell.paragraphs[0].alignment              = WD_ALIGN_PARAGRAPH.CENTER
        tc_pr = cell._tc.get_or_add_tcPr()
        from docx.oxml.ns import qn
        from docx.oxml import OxmlElement
        shd = OxmlElement("w:shd")
        shd.set(qn("w:fill"), "E65100")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:val"), "clear")
        tc_pr.append(shd)

    # Donnees
    color_map = {
        "Paye":           ("E8F5E9", "2E7D32"),
        "Pas encore paye":("FFF3E0", "E65100"),
        "Annule":         ("FFEBEE", "C62828"),
        "En cours":       ("E3F2FD", "1565C0"),
    }

    for row in rows:
        tr = table_doc.add_row()
        for i, champ in enumerate(CHAMPS):
            val  = row[champ]
            cell = tr.cells[i]

            if champ in ("prix_ht", "prix_ttc", "prix_total"):
                cell.text = f"{float(val):,.2f}"
            else:
                cell.text = str(val)

            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(8.5)

            if champ == "statut":
                bg, fg = color_map.get(str(val), ("FFFFFF", "212121"))
                run.font.bold      = True
                run.font.color.rgb = RGBColor(
                    int(fg[0:2], 16), int(fg[2:4], 16), int(fg[4:6], 16))
                from docx.oxml.ns import qn
                from docx.oxml import OxmlElement
                tc_pr = cell._tc.get_or_add_tcPr()
                shd   = OxmlElement("w:shd")
                shd.set(qn("w:fill"), bg)
                shd.set(qn("w:color"), "auto")
                shd.set(qn("w:val"), "clear")
                tc_pr.append(shd)

    # Total
    doc.add_paragraph()
    total_all = sum(r["prix_total"] for r in rows)
    p = doc.add_paragraph(f"Total general : {total_all:,.2f} MAD")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.runs[0].font.bold      = True
    p.runs[0].font.size      = Pt(11)
    p.runs[0].font.color.rgb = RGBColor(0xE6, 0x51, 0x00)

    doc.save(chemin)


# ── Export PDF ────────────────────────────────────────────
def _export_pdf(chemin, rows):
    if not PDF_OK:
        messagebox.showerror(
            "Manquant",
            "Installez reportlab :\npip install reportlab")
        return

    doc = SimpleDocTemplate(chemin, pagesize=A4,
                            leftMargin=1.5*cm, rightMargin=1.5*cm,
                            topMargin=2*cm, bottomMargin=2*cm)
    styles  = getSampleStyleSheet()
    story   = []

    # Titre
    titre_style = ParagraphStyle(
        "Titre", parent=styles["Title"],
        fontSize=18, textColor=colors.HexColor("#E65100"),
        spaceAfter=4,
    )
    story.append(Paragraph("Bons de Commande", titre_style))

    sous_style = ParagraphStyle(
        "Sous", parent=styles["Normal"],
        fontSize=9, textColor=colors.HexColor("#996633"),
        spaceAfter=14,
    )
    story.append(Paragraph(
        f"Exporte le {date.today()}  —  {len(rows)} enregistrement(s)",
        sous_style))

    # Donnees tableau
    data = [EN_TETES]
    for row in rows:
        data.append([
            str(row["id"]),
            str(row["numero_bc"]),
            str(row["client"]),
            str(row["produit"]),
            str(row["categorie"]),
            str(row["quantite"]),
            f"{row['prix_ht']:,.2f}",
            f"{row['prix_ttc']:,.2f}",
            f"{row['prix_total']:,.2f}",
            str(row["statut"]),
            str(row["date_bc"]),
        ])

    col_widths = [1*cm, 2*cm, 3*cm, 3*cm, 2.2*cm, 1*cm,
                  1.8*cm, 1.8*cm, 2*cm, 2.2*cm, 2*cm]

    tbl = Table(data, colWidths=col_widths, repeatRows=1)

    tbl_style = TableStyle([
        # En-tete
        ("BACKGROUND",  (0, 0), (-1, 0), colors.HexColor("#E65100")),
        ("TEXTCOLOR",   (0, 0), (-1, 0), colors.white),
        ("FONTNAME",    (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",    (0, 0), (-1, 0), 8),
        ("ALIGN",       (0, 0), (-1, 0), "CENTER"),
        ("BOTTOMPADDING", (0, 0), (-1, 0), 8),
        ("TOPPADDING",    (0, 0), (-1, 0), 8),
        # Donnees
        ("FONTSIZE",    (0, 1), (-1, -1), 7.5),
        ("ALIGN",       (0, 1), (-1, -1), "CENTER"),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1),
         [colors.HexColor("#FFF8F0"), colors.white]),
        ("GRID",        (0, 0), (-1, -1), 0.4, colors.HexColor("#E8D8C8")),
        ("TOPPADDING",  (0, 1), (-1, -1), 5),
        ("BOTTOMPADDING", (0, 1), (-1, -1), 5),
    ])

    # Couleur statut dans PDF
    statut_col_map = {
        "Paye":           colors.HexColor("#2E7D32"),
        "Pas encore paye":colors.HexColor("#E65100"),
        "Annule":         colors.HexColor("#C62828"),
        "En cours":       colors.HexColor("#1565C0"),
    }
    for r_idx, row in enumerate(rows, start=1):
        col = statut_col_map.get(row["statut"])
        if col:
            tbl_style.add("TEXTCOLOR", (9, r_idx), (9, r_idx), col)
            tbl_style.add("FONTNAME",  (9, r_idx), (9, r_idx), "Helvetica-Bold")

    tbl.setStyle(tbl_style)
    story.append(tbl)
    story.append(Spacer(1, 0.5*cm))

    # Total
    total_all   = sum(r["prix_total"] for r in rows)
    total_style = ParagraphStyle(
        "Total", parent=styles["Normal"],
        fontSize=11, textColor=colors.HexColor("#E65100"),
        fontName="Helvetica-Bold", alignment=2,
    )
    story.append(Paragraph(
        f"Total general : {total_all:,.2f} MAD", total_style))

    doc.build(story)


# ── Export CSV ────────────────────────────────────────────
def _export_csv(chemin, rows):
    import csv as csv_mod
    with open(chemin, "w", newline="", encoding="utf-8-sig") as f:
        writer = csv_mod.DictWriter(f, fieldnames=CHAMPS)
        writer.writeheader()
        writer.writerows(rows)


# ── Export JSON ───────────────────────────────────────────
def _export_json(chemin, rows):
    import json as json_mod
    with open(chemin, "w", encoding="utf-8") as f:
        json_mod.dump({
            "export_date": str(date.today()),
            "total_records": len(rows),
            "bons_commande": rows,
        }, f, ensure_ascii=False, indent=2)


# ══════════════════════════════════════════════════════════
#   IMPORTER BC
# ══════════════════════════════════════════════════════════
def importer_fichier(table, format_):
    ext_map = {
        "excel": [("Fichier Excel", "*.xlsx"), ("Tous", "*.*")],
        "csv":   [("Fichier CSV",   "*.csv"),  ("Tous", "*.*")],
        "json":  [("Fichier JSON",  "*.json"), ("Tous", "*.*")],
    }

    chemin = filedialog.askopenfilename(
        title=f"Importer depuis {format_.upper()}",
        filetypes=ext_map.get(format_, [("Tous", "*.*")]),
    )
    if not chemin:
        return

    try:
        if format_ == "excel":
            lignes = _import_excel(chemin)
        elif format_ == "csv":
            lignes = _import_csv(chemin)
        elif format_ == "json":
            lignes = _import_json(chemin)
        else:
            lignes = []

        if not lignes:
            messagebox.showwarning("Vide", "Aucune ligne trouvee dans le fichier.")
            return

        # Fenetre de confirmation
        _fenetre_import_preview(table, lignes, chemin)

    except Exception as ex:
        messagebox.showerror("Erreur import", f"Erreur lors de l'import :\n{ex}")


# ── Import Excel ──────────────────────────────────────────
def _import_excel(chemin):
    if not EXCEL_OK:
        messagebox.showerror("Manquant", "Installez openpyxl :\npip install openpyxl")
        return []

    wb   = openpyxl.load_workbook(chemin, data_only=True)
    ws   = wb.active
    rows = list(ws.iter_rows(values_only=True))

    # Detecter si ligne titre (row 1) et en-tetes (row 2)
    start = 1
    if rows and rows[0] and any(
            str(rows[0][0] or "").lower() in ("bons de commande", "titre")
            for _ in [1]):
        start = 2

    lignes = []
    for row in rows[start:]:
        if row and any(cell is not None for cell in row):
            lignes.append([str(c) if c is not None else "" for c in row])
    return lignes


# ── Import CSV ────────────────────────────────────────────
def _import_csv(chemin):
    import csv as csv_mod
    lignes = []
    with open(chemin, encoding="utf-8-sig", newline="") as f:
        reader = csv_mod.reader(f)
        header = next(reader, None)  # ignore en-tete
        for row in reader:
            if row:
                lignes.append(row)
    return lignes


# ── Import JSON ───────────────────────────────────────────
def _import_json(chemin):
    import json as json_mod
    with open(chemin, encoding="utf-8") as f:
        data = json_mod.load(f)

    # Accepte {"bons_commande": [...]} ou [...]
    if isinstance(data, dict):
        data = data.get("bons_commande", [])

    lignes = []
    for item in data:
        if isinstance(item, dict):
            lignes.append([str(item.get(k, "")) for k in CHAMPS])
    return lignes


# ══════════════════════════════════════════════════════════
#   FENETRE PREVIEW IMPORT
# ══════════════════════════════════════════════════════════
def _fenetre_import_preview(table, lignes, chemin):
    t   = get_theme()
    win = tk.Toplevel()
    win.title(f"Apercu import — {len(lignes)} ligne(s)")
    win.geometry("900x520")
    win.configure(bg=t["bg"])
    win.grab_set()

    hdr = tk.Frame(win, bg="#1565C0", height=50)
    hdr.pack(fill="x")
    tk.Label(hdr, text=f"📥  Apercu Import — {os.path.basename(chemin)}",
             bg="#1565C0", fg="white",
             font=("Arial", 13, "bold")).pack(pady=12)

    # Tableau preview
    frame_prev = tk.Frame(win, bg=t["bg"])
    frame_prev.pack(fill="both", expand=True, padx=16, pady=10)

    prev_cols = EN_TETES
    prev_tree = ttk.Treeview(frame_prev, columns=prev_cols,
                             show="headings", height=14,
                             style="BC.Treeview")
    for col in prev_cols:
        prev_tree.heading(col, text=col)
        prev_tree.column(col, width=80, anchor="center")

    for i, row in enumerate(lignes[:50]):  # max 50 lignes preview
        tag = "pair" if i % 2 == 0 else "impair"
        # Compléter si la ligne est trop courte
        padded = list(row) + [""] * max(0, len(prev_cols) - len(row))
        prev_tree.insert("", "end", tags=(tag,), values=padded[:len(prev_cols)])

    sb = ttk.Scrollbar(frame_prev, orient="vertical", command=prev_tree.yview)
    prev_tree.configure(yscroll=sb.set)
    sb.pack(side="right", fill="y")
    prev_tree.pack(fill="both", expand=True)

    if len(lignes) > 50:
        tk.Label(win, text=f"... et {len(lignes)-50} autres lignes.",
                 bg=t["bg"], fg="#FF8F00",
                 font=("Arial", 9, "italic")).pack()

    # Info
    info = tk.Label(win,
        text=f"{len(lignes)} lignes detectees.  Seules les lignes avec client et produit existants seront importees.",
        bg=t["bg"], fg=t["text"], font=("Arial", 9))
    info.pack(pady=(0, 6))

    # Boutons
    btn_frame = tk.Frame(win, bg=t["bg"])
    btn_frame.pack(pady=8)

    def confirmer():
        nb_ok, nb_err = _inserer_lignes(lignes)
        charger_bcs(table)
        win.destroy()
        messagebox.showinfo(
            "Import termine",
            f"Import termine.\n\n"
            f"Inseres avec succes : {nb_ok}\n"
            f"Ignores (erreur)    : {nb_err}",
        )

    _btn_ok(btn_frame, f"✅  Importer {len(lignes)} ligne(s)", confirmer,
            "#1565C0", "#0D47A1").pack(side="left", padx=8)
    tk.Button(btn_frame, text="Annuler", command=win.destroy,
              bg=t["card"], fg=t["text"], font=("Arial", 11),
              relief="flat", cursor="hand2",
              padx=16, pady=9).pack(side="left", padx=8)


# ══════════════════════════════════════════════════════════
#   INSERER LIGNES IMPORTEES EN BASE
# ══════════════════════════════════════════════════════════
def _inserer_lignes(lignes):
    nb_ok  = 0
    nb_err = 0

    for row in lignes:
        try:
            # row : [id, numero_bc, client, produit, categorie,
            #        quantite, prix_ht, prix_ttc, prix_total, statut, date_bc]
            padded = list(row) + [""] * max(0, 11 - len(row))

            nom_client  = str(padded[2]).strip()
            nom_produit = str(padded[3]).strip()

            client  = session.query(Client).filter_by(nom=nom_client).first()
            produit = session.query(Produit).filter_by(nom=nom_produit).first()

            if not client or not produit:
                nb_err += 1
                continue

            # Eviter doublons sur numero_bc
            num_bc = str(padded[1]).strip() or None
            if num_bc and session.query(BonCommande).filter_by(
                    numero_bc=num_bc).first():
                nb_err += 1
                continue

            qte   = int(float(str(padded[5]).replace(",", "") or 0))
            ht    = float(str(padded[6]).replace(",", "") or 0)
            ttc   = float(str(padded[7]).replace(",", "") or 0)
            total = float(str(padded[8]).replace(",", "") or 0)
            statut = str(padded[9]).strip() or "En cours"

            try:
                d = date.fromisoformat(str(padded[10]).strip())
            except Exception:
                d = date.today()

            bc = BonCommande(
                numero_bc=num_bc,
                client_id=client.id,
                produit_id=produit.id,
                categorie=produit.categorie,
                quantite=qte,
                prix_ht=ht,
                prix_ttc=ttc,
                prix_total=total,
                statut=statut,
                date_bc=d,
            )
            session.add(bc)
            nb_ok += 1

        except Exception:
            nb_err += 1
            continue

    session.commit()
    return nb_ok, nb_err


# ══════════════════════════════════════════════════════════
#   EXPORTER UN BC EN PDF (bouton dans fiche detail)
# ══════════════════════════════════════════════════════════
def _exporter_un_pdf(bc, client, produit):
    """Exporte un seul BC en PDF professionnel (style devis)."""
    chemin = filedialog.asksaveasfilename(
        title="Exporter le bon de commande en PDF",
        defaultextension=".pdf",
        filetypes=[("PDF", "*.pdf"), ("Tous", "*.*")],
        initialfile=f"bc_{bc.numero_bc or bc.id}.pdf",
    )
    if not chemin:
        return
    try:
        _generer_pdf_bc(bc, client, produit, chemin)
        if messagebox.askyesno("Export reussi",
                               f"BC exporte :\n{chemin}\n\nOuvrir maintenant ?"):
            os.startfile(chemin)
    except Exception as ex:
        messagebox.showerror("Erreur export", f"Erreur :\n{ex}")


def _generer_pdf_bc(bc, client, produit, chemin):
    """Genere un PDF professionnel au format Bon de Commande (style SAGE/devis)."""
    if not PDF_OK:
        raise RuntimeError(
            "reportlab non installe.\nExecutez : pip install reportlab"
        )
    from reportlab.lib.enums import TA_CENTER, TA_RIGHT, TA_LEFT
    from reportlab.platypus import HRFlowable

    doc    = SimpleDocTemplate(chemin, pagesize=A4,
                               leftMargin=1.8*cm, rightMargin=1.8*cm,
                               topMargin=1.8*cm, bottomMargin=1.8*cm)
    styles = getSampleStyleSheet()
    story  = []

    ORANGE = colors.HexColor("#E65100")
    GRAY   = colors.HexColor("#555555")
    LGRAY  = colors.HexColor("#F5F5F5")
    WHITE  = colors.white

    # ── Titre BON DE COMMANDE ────────────────────────────────────
    story.append(Paragraph(
        f"BON DE COMMANDE n° {bc.numero_bc or bc.id}",
        ParagraphStyle("TitrBC", parent=styles["Title"],
                       fontSize=24, textColor=ORANGE,
                       spaceAfter=2, alignment=TA_LEFT,
                       fontName="Helvetica-Bold")))
    story.append(HRFlowable(width="100%", thickness=3,
                             color=ORANGE, spaceAfter=10))

    # ── Bloc entreprise + destinataire ───────────────────────────
    client_nom   = client.nom                         if client else "N/A"
    client_adr   = getattr(client, "adresse",   "") or ""
    client_ville = getattr(client, "ville",     "") or ""
    client_email = getattr(client, "email",     "") or ""
    client_tel   = getattr(client, "telephone", "") or ""

    bloc_data = [[
        Paragraph("<b>VentePro</b><br/>— Votre adresse —",
                  ParagraphStyle("Ent", parent=styles["Normal"],
                                 fontSize=9, textColor=colors.black)),
        Paragraph(
            f"<b>Destinataire</b><br/><b>{client_nom}</b><br/>"
            f"{client_adr}<br/>{client_ville}<br/>"
            f"{client_email}<br/>{client_tel}",
            ParagraphStyle("Cli", parent=styles["Normal"],
                           fontSize=9, textColor=colors.black,
                           alignment=TA_RIGHT)),
    ]]
    bloc_tbl = Table(bloc_data, colWidths=[9*cm, 8*cm])
    bloc_tbl.setStyle(TableStyle([
        ("VALIGN",        (0, 0), (-1, -1), "TOP"),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(bloc_tbl)
    story.append(Spacer(1, 0.6*cm))

    # ── Informations BC ──────────────────────────────────────────
    info_data = [
        ["Date du BC :",      str(bc.date_bc),
         "Reference :",       bc.numero_bc or str(bc.id)],
        ["Validite :",        "30 jours",
         "Statut :",          bc.statut or "-"],
        ["Contact client :",  client_nom,
         "Telephone :",       client_tel or "-"],
    ]
    info_tbl = Table(info_data, colWidths=[4*cm, 5*cm, 4*cm, 4.2*cm])
    info_tbl.setStyle(TableStyle([
        ("FONTNAME",      (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE",      (0, 0), (-1, -1), 8.5),
        ("FONTNAME",      (0, 0), (0, -1),  "Helvetica-Bold"),
        ("FONTNAME",      (2, 0), (2, -1),  "Helvetica-Bold"),
        ("TEXTCOLOR",     (0, 0), (0, -1),  ORANGE),
        ("TEXTCOLOR",     (2, 0), (2, -1),  ORANGE),
        ("ROWBACKGROUNDS",(0, 0), (-1, -1),
         [colors.HexColor("#FFF8F0"), WHITE]),
        ("TOPPADDING",    (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ("GRID",          (0, 0), (-1, -1), 0.3, colors.HexColor("#E8D8C8")),
    ]))
    story.append(info_tbl)
    story.append(Spacer(1, 0.8*cm))

    # ── Tableau lignes BC ────────────────────────────────────────
    qte    = bc.quantite or 1
    ht_u   = float(bc.prix_ht or 0)
    ttc_u  = float(bc.prix_ttc or 0)
    tva_u  = round(ttc_u - ht_u, 2)
    tva_p  = round(tva_u / ht_u * 100, 1) if ht_u > 0 else 0.0
    tot_ht  = round(ht_u  * qte, 2)
    tot_tva = round(tva_u * qte, 2)
    tot_ttc = float(bc.prix_total or 0)
    prod_nom = produit.nom if produit else (bc.categorie or "-")
    cat_nom  = bc.categorie or "-"

    col_hdrs = ["Description", "Categorie", "Quantite", "Unite",
                "Prix unitaire HT", "% TVA", "Total TVA", "Total TTC"]
    lines = [
        col_hdrs,
        [prod_nom, cat_nom, str(qte), "u",
         f"{ht_u:,.2f}", f"{tva_p:.0f} %",
         f"{tot_tva:,.2f}", f"{tot_ttc:,.2f}"],
        ["", "", "", "Total HT",  "", "", f"{tot_ht:,.2f}",  ""],
        ["", "", "", "Total TVA", "", "", f"{tot_tva:,.2f}", ""],
        ["", "", "", "",          "", "", "Total TTC",       f"{tot_ttc:,.2f}"],
    ]
    col_w   = [4.5*cm, 2.2*cm, 1.6*cm, 1.2*cm, 2.8*cm, 1.4*cm, 2.2*cm, 2.2*cm]
    data_tbl = Table(lines, colWidths=col_w, repeatRows=1)
    data_tbl.setStyle(TableStyle([
        ("BACKGROUND",     (0, 0), (-1, 0), ORANGE),
        ("TEXTCOLOR",      (0, 0), (-1, 0), WHITE),
        ("FONTNAME",       (0, 0), (-1, 0), "Helvetica-Bold"),
        ("FONTSIZE",       (0, 0), (-1, 0), 8),
        ("ALIGN",          (2, 0), (-1, -1), "RIGHT"),
        ("ALIGN",          (0, 0), (1, -1),  "LEFT"),
        ("BOTTOMPADDING",  (0, 0), (-1, 0), 8),
        ("TOPPADDING",     (0, 0), (-1, 0), 8),
        ("FONTSIZE",       (0, 1), (-1, -1), 8.5),
        ("ROWBACKGROUNDS", (0, 1), (-1, 2),
         [WHITE, colors.HexColor("#FFF8F0")]),
        ("TOPPADDING",     (0, 1), (-1, -1), 5),
        ("BOTTOMPADDING",  (0, 1), (-1, -1), 5),
        ("LEFTPADDING",    (0, 0), (-1, -1), 5),
        ("RIGHTPADDING",   (0, 0), (-1, -1), 5),
        ("GRID",           (0, 0), (-1, 2),  0.4, colors.HexColor("#E8D8C8")),
        ("FONTNAME",       (3, 2), (-1, -1), "Helvetica-Bold"),
        ("TEXTCOLOR",      (3, 2), (-1, -1), ORANGE),
        ("LINEABOVE",      (0, 2), (-1, 2),  0.8, colors.HexColor("#E8D8C8")),
        ("BACKGROUND",     (0, -1), (-1, -1), colors.HexColor("#BF360C")),
        ("TEXTCOLOR",      (0, -1), (-1, -1), WHITE),
        ("FONTNAME",       (0, -1), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE",       (0, -1), (-1, -1), 10),
    ]))
    story.append(data_tbl)
    story.append(Spacer(1, 0.8*cm))

    # ── Zone Bon pour accord ─────────────────────────────────────
    sig_tbl = Table(
        [["Signature du client (precedee de la mention « Bon pour accord ») :"]],
        colWidths=[17*cm])
    sig_tbl.setStyle(TableStyle([
        ("FONTNAME",      (0, 0), (-1, -1), "Helvetica"),
        ("FONTSIZE",      (0, 0), (-1, -1), 8.5),
        ("TEXTCOLOR",     (0, 0), (-1, -1), GRAY),
        ("BACKGROUND",    (0, 0), (-1, -1), LGRAY),
        ("TOPPADDING",    (0, 0), (-1, -1), 40),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
        ("LEFTPADDING",   (0, 0), (-1, -1), 10),
        ("BOX",           (0, 0), (-1, -1), 0.5, colors.HexColor("#E8D8C8")),
    ]))
    story.append(sig_tbl)
    story.append(Spacer(1, 0.5*cm))

    # ── Pied de page ─────────────────────────────────────────────
    story.append(HRFlowable(width="100%", thickness=1,
                             color=colors.HexColor("#E8D8C8"), spaceBefore=4))
    story.append(Paragraph(
        f"VentePro  —  Bon de Commande N° {bc.numero_bc or bc.id}"
        f"  —  {str(bc.date_bc)}  —  Valable 30 jours",
        ParagraphStyle("Footer", parent=styles["Normal"],
                       fontSize=7.5,
                       textColor=colors.HexColor("#999999"),
                       alignment=TA_CENTER)))
    doc.build(story)


# ══════════════════════════════════════════════════════════
#   DETAIL BC  (double-clic)
# ══════════════════════════════════════════════════════════
def detail_bc(table):
    sel = table.selection()
    if not sel:
        return

    valeurs = table.item(sel[0])["values"]
    bc      = session.query(BonCommande).filter_by(id=valeurs[0]).first()
    if not bc:
        return

    client  = session.query(Client).filter_by(id=bc.client_id).first()
    produit = session.query(Produit).filter_by(id=bc.produit_id).first()

    t = get_theme()
    win, body, t = _fenetre_modale(
        f"Detail — {bc.numero_bc or bc.id}", 420, 440)

    info_bg = "#1A1A35" if t["bg"] in ("#1A1A2E", "#0F0F1A") else "#FFF0E5"
    frame   = tk.Frame(body, bg=info_bg)
    frame.pack(fill="both", expand=True)

    champs = [
        ("N BC",      bc.numero_bc or "-"),
        ("Date",      str(bc.date_bc)),
        ("Client",    client.nom  if client  else "N/A"),
        ("Produit",   produit.nom if produit else "N/A"),
        ("Categorie", bc.categorie or "-"),
        ("Quantite",  str(bc.quantite)),
        ("Prix HT",   f"{bc.prix_ht:,.2f} MAD"),
        ("Prix TTC",  f"{bc.prix_ttc:,.2f} MAD"),
        ("Total",     f"{bc.prix_total:,.2f} MAD"),
        ("Statut",    bc.statut or "-"),
    ]

    for label, val in champs:
        row = tk.Frame(frame, bg=info_bg)
        row.pack(fill="x", padx=16, pady=4)
        tk.Label(row, text=f"{label} :", bg=info_bg,
                 fg="#FF8F00", font=("Arial", 11, "bold"),
                 width=11, anchor="w").pack(side="left")
        tk.Label(row, text=val, bg=info_bg,
                 fg=t["text"], font=("Arial", 11)).pack(side="left")
        tk.Frame(frame, bg="#333355" if t["bg"] in ("#1A1A2E",)
                 else "#E8D8C8", height=1).pack(fill="x", padx=16)

    btn_row = tk.Frame(body, bg=t["bg"])
    btn_row.pack(pady=(12, 0))

    _btn_ok(btn_row, "📤  Exporter PDF",
            lambda: _exporter_un_pdf(bc, client, produit),
            "#1565C0", "#0D47A1").pack(side="left", padx=6)
    tk.Button(btn_row, text="Fermer", command=win.destroy,
              bg="#E65100", fg="white",
              font=("Arial", 12, "bold"),
              relief="flat", cursor="hand2",
              padx=16, pady=9).pack(side="left", padx=6)


def _imprimer_un(bc, client, produit):
    """Conserve la compatibilite — redirige vers l'export PDF."""
    _exporter_un_pdf(bc, client, produit)